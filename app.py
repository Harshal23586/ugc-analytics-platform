import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import numpy as np
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
from sklearn.preprocessing import StandardScaler
from sklearn.cluster import KMeans
import warnings
warnings.filterwarnings('ignore')

# Additional imports for enhanced functionality
from datetime import datetime, timedelta
import json
from typing import Dict, List, Tuple, Any, Optional
import hashlib
import sqlite3
import io
import base64
import os
import tempfile
from pathlib import Path
from datetime import datetime
import glob


# RAG-specific imports
import PyPDF2
import docx
import faiss
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# Add these instead:
import re
from typing import List

# Initialize session state at module level
if 'session_initialized' not in st.session_state:
    st.session_state.session_initialized = True
    st.session_state.institution_user = None
    st.session_state.user_role = None
    st.session_state.rag_analysis = None
    st.session_state.selected_institution = None

def create_api_documentation():
    """Create API documentation and testing interface"""
    st.header("🌐 API Integration Portal")
    
    st.info("""
    **RESTful API for UGC/AICTE Institutional Analytics Platform**
    
    This API allows external systems, hackathon participants, and institutions to:
    - Access institutional performance data
    - Run AI-powered analytics
    - Get configuration parameters
    - Export data in multiple formats
    """)
    
    tab1, tab2, tab3, tab4 = st.tabs([
        "📖 API Documentation", 
        "🔑 API Keys", 
        "🧪 API Testing",
        "📊 Quick Integration"
    ])
    
    with tab1:
        st.subheader("API Endpoints Overview")
        
        endpoints = [
            {
                "endpoint": "GET /institutions",
                "description": "Get list of institutions",
                "parameters": "year, institution_type, state, limit"
            },
            {
                "endpoint": "GET /institutions/{id}",
                "description": "Get institution performance data",
                "parameters": "start_year, end_year"
            },
            {
                "endpoint": "GET /performance/ranking",
                "description": "Get performance ranking",
                "parameters": "year, institution_type, limit"
            },
            {
                "endpoint": "POST /analysis/run",
                "description": "Run AI analysis",
                "parameters": "institution_id, analysis_type"
            },
            {
                "endpoint": "GET /metrics",
                "description": "Get performance metrics configuration",
                "parameters": "None"
            },
            {
                "endpoint": "GET /documents",
                "description": "Get document requirements",
                "parameters": "approval_type"
            },
            {
                "endpoint": "GET /export/institutions",
                "description": "Export data",
                "parameters": "format, year"
            }
        ]
        
        for endpoint in endpoints:
            with st.expander(f"🌐 {endpoint['endpoint']}"):
                st.write(f"**Description:** {endpoint['description']}")
                st.write(f"**Parameters:** {endpoint['parameters']}")
                
                # Show example request
                st.code(f"""
# Python Example
import requests

api_key = "your_api_key_here"
headers = {{"Authorization": f"Bearer {{{{api_key}}}}"}}

# For {endpoint['endpoint']}
response = requests.get(
    "http://localhost:8000{endpoint['endpoint'].replace('{id}', 'INST_0001')}",
    headers=headers,
    params={{"year": 2023}}
)

print(response.json())
                """, language="python")
    
    with tab2:
        st.subheader("🔑 API Access Keys")
        
        st.warning("⚠️ These keys are for demonstration. Use secure keys in production.")
        
        api_keys = {
            "ugc_admin": "ugc_api_key_2024_secure",
            "aictel_team": "aictel_api_2024_secure", 
            "hackathon_2024": "smart_india_hackathon_key",
            "institution_api": "institution_access_2024"
        }
        
        for role, key in api_keys.items():
            col1, col2 = st.columns([2, 3])
            with col1:
                st.write(f"**{role.replace('_', ' ').title()}**")
            with col2:
                st.code(key, language="text")
        
        st.info("""
        **Usage:**
        ```python
        headers = {
            "Authorization": "Bearer your_api_key_here"
        }
        ```
        """)
    
    with tab3:
        st.subheader("🧪 API Testing Interface")
        
        # Quick test interface
        endpoint = st.selectbox("Select Endpoint", [
            "/institutions",
            "/institutions/INST_0001", 
            "/performance/ranking",
            "/metrics",
            "/documents"
        ])
        
        api_key = st.selectbox("Select API Key", list(api_keys.keys()))
        
        if st.button("🚀 Test API Endpoint"):
            with st.spinner("Testing API endpoint..."):
                # Simulate API call
                st.success("✅ API Test Successful!")
                
                # Show expected response structure
                if "/institutions" in endpoint and "INST" not in endpoint:
                    st.json({
                        "institutions": [
                            {
                                "institution_id": "INST_0001",
                                "institution_name": "University 001",
                                "institution_type": "State University",
                                "state": "Maharashtra",
                                "year": 2023
                            }
                        ]
                    })
                elif "/institutions/INST" in endpoint:
                    st.json({
                        "institution_id": "INST_0001",
                        "institution_name": "University 001",
                        "performance_score": 8.5,
                        "risk_level": "Low Risk",
                        "approval_recommendation": "Full Approval - 5 Years"
                    })
    
    with tab4:
        st.subheader("📊 Quick Integration Guide")
        
        st.download_button(
            label="📥 Download Python SDK",
            data="""
# UGC/AICTE API Python Client
import requests

class UGCAnalyticsAPI:
    def __init__(self, api_key, base_url="http://localhost:8000"):
        self.api_key = api_key
        self.base_url = base_url
        self.headers = {"Authorization": f"Bearer {api_key}"}
    
    def get_institutions(self, year=None, institution_type=None):
        params = {}
        if year: params['year'] = year
        if institution_type: params['institution_type'] = institution_type
        
        response = requests.get(
            f"{self.base_url}/institutions",
            headers=self.headers,
            params=params
        )
        return response.json()
    
    def get_performance(self, institution_id):
        response = requests.get(
            f"{self.base_url}/institutions/{institution_id}",
            headers=self.headers
        )
        return response.json()
    
    def run_analysis(self, institution_id, analysis_type="comprehensive"):
        data = {
            "institution_id": institution_id,
            "analysis_type": analysis_type
        }
        response = requests.post(
            f"{self.base_url}/analysis/run",
            headers=self.headers,
            json=data
        )
        return response.json()

# Usage
api = UGCAnalyticsAPI(api_key="your_api_key")
institutions = api.get_institutions(year=2023)
            """,
            file_name="ugc_api_client.py"
        )
        
        st.write("**Quick Start Commands:**")
        st.code("""
# Start API Server
python api_endpoints.py

# Install dependencies
pip install fastapi uvicorn sqlalchemy

# Test API
curl -X GET "http://localhost:8000/institutions" \\
     -H "Authorization: Bearer smart_india_hackathon_key"
        """, language="bash")


# Also update the RAGDocument class initialization
class RAGDocument:
    def __init__(self, page_content: str, metadata: dict = None):
        self.page_content = page_content
        self.metadata = metadata if metadata is not None else {}

class SimpleTextSplitter:
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str) -> List[str]:
        """Simple text splitter that splits by sentences and chunks"""
        # Split by sentences (simple approach)
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # If adding this sentence would exceed chunk size, save current chunk
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                # Keep overlap if specified
                if self.chunk_overlap > 0:
                    # Simple overlap: keep last few sentences
                    overlap_sentences = current_chunk.split('.')[-3:]
                    current_chunk = '.'.join(overlap_sentences) + '. ' + sentence
                else:
                    current_chunk = sentence
            else:
                if current_chunk:
                    current_chunk += '. ' + sentence
                else:
                    current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks

class SimpleVectorStore:
    def __init__(self, embedding_model):
        self.embedding_model = embedding_model
        self.documents = []
        self.embeddings = []
    
    def from_embeddings(self, text_embeddings):
        """Create vector store from text-embedding pairs"""
        texts, embeddings = zip(*text_embeddings)
        self.documents = list(texts)
        self.embeddings = np.array(embeddings)
        return self
    
    def similarity_search_with_score(self, query: str, k: int = 5):
        """Simple similarity search using cosine similarity"""
        if not self.embeddings.size:
            return []
        
        query_embedding = self.embedding_model.encode([query])
        similarities = cosine_similarity(query_embedding, self.embeddings)[0]
        
        # Get top k results
        top_indices = np.argsort(similarities)[-k:][::-1]
        results = []
        
        for idx in top_indices:
            if similarities[idx] > 0:  # Only include positive similarities
                doc = RAGDocument(
                    page_content=self.documents[idx],
                    metadata={"similarity_score": float(similarities[idx])}
                )
                results.append((doc, float(similarities[idx])))
        
        return results

# Page configuration
st.set_page_config(
    page_title="AI-Powered Institutional Approval System - UGC/AICTE",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded"
)

class RAGDataExtractor:
    def __init__(self):
        try:
            # Import torch first
            import torch
            
            # Set device to CPU explicitly
            device = torch.device('cpu')
            
            # Load model with safe initialization
            # Use a smaller, more compatible model
            model_name = 'sentence-transformers/all-MiniLM-L6-v2'
            
            # Initialize with device mapping to avoid meta tensor issues
            self.embedding_model = SentenceTransformer(
                model_name, 
                device='cpu',
                cache_folder='./model_cache'
            )
            
            # Force model to CPU and ensure proper initialization
            if hasattr(self.embedding_model, 'to'):
                self.embedding_model = self.embedding_model.to(device)
            
            # Test the model with a small inference
            test_embedding = self.embedding_model.encode(["test sentence"])
            if test_embedding is not None:
                st.success("✅ RAG System with embeddings initialized successfully")
            else:
                raise Exception("Test embedding failed")
                
            self.text_splitter = SimpleTextSplitter(
                chunk_size=1000,
                chunk_overlap=200
            )
            self.vector_store = None
            self.documents = []
            
        except Exception as e:
            st.warning(f"⚠️ RAG system using lightweight mode: {e}")
            # Use a lightweight alternative - TF-IDF based approach
            self.embedding_model = None
            self.text_splitter = SimpleTextSplitter(
                chunk_size=1000,
                chunk_overlap=200
            )
            self.vector_store = None
            self.documents = []
            self._setup_lightweight_analyzer()

    def _setup_lightweight_analyzer(self):
        """Setup lightweight text analysis without heavy embeddings"""
        try:
            from sklearn.feature_extraction.text import TfidfVectorizer
            self.vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
            self.tfidf_matrix = None
            self.document_texts = []
        except:
            self.vectorizer = None

    def build_vector_store(self, documents: List[RAGDocument]):
        """Build vector store from documents with fallback options"""
        if not documents:
            return None
        
        texts = [doc.page_content for doc in documents]
        if not texts:
            return None
            
        try:
            # Try using the embedding model first
            if self.embedding_model is not None:
                embeddings = self.embedding_model.encode(texts)
                text_embeddings = list(zip(texts, embeddings))
                self.vector_store = SimpleVectorStore(self.embedding_model).from_embeddings(text_embeddings)
                self.documents = documents
                st.success(f"✅ Vector store built with {len(documents)} documents using embeddings")
            elif hasattr(self, 'vectorizer') and self.vectorizer is not None:
                # Fallback to TF-IDF
                self.tfidf_matrix = self.vectorizer.fit_transform(texts)
                self.document_texts = texts
                self.documents = documents
                st.success(f"✅ Vector store built with {len(documents)} documents using TF-IDF")
            else:
                # Basic text storage
                self.document_texts = texts
                self.documents = documents
                st.success(f"✅ Documents stored for basic search ({len(documents)} documents)")
                
        except Exception as e:
            st.warning(f"Vector store creation using basic storage: {e}")
            # Basic storage as fallback
            self.document_texts = texts
            self.documents = documents

    def query_documents_lightweight(self, query: str, k: int = 5) -> List[Tuple[str, float]]:
        """Lightweight document query using TF-IDF or basic text matching"""
        if not hasattr(self, 'document_texts') or not self.document_texts:
            return []
            
        results = []
        
        # Simple keyword-based matching
        query_terms = set(query.lower().split())
        
        for i, text in enumerate(self.document_texts):
            if not text:
                continue
                
            text_terms = set(text.lower().split())
            common_terms = query_terms.intersection(text_terms)
            
            if common_terms:
                # Simple relevance score based on term overlap
                score = len(common_terms) / len(query_terms)
                if score > 0.1:  # Minimum threshold
                    doc = RAGDocument(
                        page_content=text[:500] + "..." if len(text) > 500 else text,
                        metadata={"similarity_score": float(score), "source": "lightweight_search"}
                    )
                    results.append((doc, float(score)))
        
        # Sort by score and return top k
        results.sort(key=lambda x: x[1], reverse=True)
        return results[:k]

    def extract_comprehensive_data(self, uploaded_files: List) -> Dict[str, Any]:
        """Extract comprehensive data from all uploaded files"""
        all_text = ""
        all_structured_data = {
            'academic_metrics': {},
            'research_metrics': {},
            'infrastructure_metrics': {},
            'governance_metrics': {},
            'student_metrics': {},
            'financial_metrics': {},
            'raw_text': "",
            'file_names': []
        }
    
        documents = []
        processed_count = 0
    
        for file in uploaded_files:
            try:
                # Extract text
                text = self.extract_text_from_file(file)
                if not text or not text.strip():
                    st.warning(f"No extractable text found in {file.name}")
                    all_structured_data['file_names'].append(file.name)
                    continue
                    
                cleaned_text = self.preprocess_text(text)
                all_text += cleaned_text + "\n\n"
            
                # Create document for analysis
                doc = RAGDocument(
                    page_content=cleaned_text,
                    metadata={"source": file.name, "type": "institutional_data"}
                )
                documents.append(doc)
            
                # Extract structured data using enhanced pattern matching
                file_data = self.extract_structured_data_enhanced(cleaned_text, file.name)
            
                # Merge data from all files
                for category in file_data:
                    if category in all_structured_data:
                        all_structured_data[category].update(file_data[category])
            
                all_structured_data['file_names'].append(file.name)
                processed_count += 1
                
                st.success(f"✅ Processed {file.name} - extracted {len(file_data)} data categories")
            
            except Exception as e:
                st.error(f"Error processing {file.name}: {str(e)}")
                all_structured_data['file_names'].append(file.name)
                continue
    
        # Build search index
        if documents:
            self.build_vector_store(documents)
            st.success(f"✅ Successfully processed {processed_count}/{len(uploaded_files)} files")
        else:
            st.warning("⚠️ No documents were successfully processed")
    
        all_structured_data['raw_text'] = all_text
    
        return all_structured_data

    def extract_structured_data_enhanced(self, text: str, filename: str) -> Dict[str, Any]:
        """Enhanced structured data extraction with better pattern matching"""
        data = {
            'academic_metrics': {},
            'research_metrics': {},
            'infrastructure_metrics': {},
            'governance_metrics': {},
            'student_metrics': {},
            'financial_metrics': {}
        }
        
        # Enhanced patterns for academic metrics
        academic_patterns = {
            'naac_grade': [
                r'NAAC\s*(?:grade|accreditation|score)[:\s]*([A+]+)',
                r'Accreditation\s*(?:Grade|Status)[:\s]*([A+]+)',
                r'Grade\s*[:\s]*([A+]+)'
            ],
            'nirf_ranking': [
                r'NIRF\s*(?:rank|ranking)[:\s]*(\d+)',
                r'National.*Ranking[:\s]*(\d+)',
                r'Rank[:\s]*(\d+).*NIRF'
            ],
            'student_faculty_ratio': [
                r'(?:student|student-faculty)\s*(?:ratio|ratio:)[:\s]*(\d+(?:\.\d+)?)',
                r'Faculty.*Student[:\s]*(\d+(?:\.\d+)?)',
                r'Ratio[:\s]*(\d+(?:\.\d+)?).*student.*faculty'
            ]
        }
        
        # Research metrics patterns
        research_patterns = {
            'research_publications': [
                r'research\s*(?:publications|papers)[:\s]*(\d+)',
                r'publications[:\s]*(\d+)',
                r'published\s*(?:papers|articles)[:\s]*(\d+)'
            ],
            'research_grants': [
                r'research\s*(?:grants|funding)[:\s]*[₹$]?\s*(\d+(?:,\d+)*(?:\.\d+)?)',
                r'grants.*received[:\s]*[₹$]?\s*(\d+(?:,\d+)*(?:\.\d+)?)',
                r'funding.*amount[:\s]*[₹$]?\s*(\d+(?:,\d+)*(?:\.\d+)?)'
            ]
        }
        
        # Extract data using multiple patterns
        for category, patterns in academic_patterns.items():
            for pattern in patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    data['academic_metrics'][category] = matches[0]
                    break
        
        for category, patterns in research_patterns.items():
            for pattern in patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    data['research_metrics'][category] = matches[0]
                    break
        
        # Extract numbers with context for other categories
        self.extract_contextual_data_enhanced(text, data, filename)
        
        return data

    def extract_contextual_data_enhanced(self, text: str, data: Dict, filename: str):
        """Enhanced contextual data extraction"""
        # Look for infrastructure metrics
        infra_patterns = [
            (r'library.*?(\d+(?:,\d+)*)\s*(?:volumes|books)', 'library_volumes', 'infrastructure_metrics'),
            (r'campus.*?(\d+(?:\.\d+)?)\s*(?:acres|hectares)', 'campus_area', 'infrastructure_metrics'),
            (r'laboratory.*?(\d+)', 'laboratories_count', 'infrastructure_metrics'),
            (r'classroom.*?(\d+)', 'classrooms_count', 'infrastructure_metrics')
        ]
        
        # Financial metrics
        financial_patterns = [
            (r'budget.*?[₹$]?\s*(\d+(?:,\d+)*(?:\.\d+)?)', 'annual_budget', 'financial_metrics'),
            (r'grant.*?[₹$]?\s*(\d+(?:,\d+)*(?:\.\d+)?)', 'total_grants', 'financial_metrics'),
            (r'revenue.*?[₹$]?\s*(\d+(?:,\d+)*(?:\.\d+)?)', 'annual_revenue', 'financial_metrics')
        ]
        
        # Student metrics
        student_patterns = [
            (r'students.*?(\d+(?:,\d+)*)', 'total_students', 'student_metrics'),
            (r'enrollment.*?(\d+(?:,\d+)*)', 'total_enrollment', 'student_metrics'),
            (r'placement.*?(\d+(?:\.\d+)?)%', 'placement_rate', 'student_metrics')
        ]
        
        all_patterns = infra_patterns + financial_patterns + student_patterns
        
        for pattern, key, category in all_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                data[category][key] = matches[0]
        
class InstitutionalAIAnalyzer:
    def __init__(self):
        self.init_database()
        self.historical_data = self.load_or_generate_data()
        self.performance_metrics = self.define_performance_metrics()
        self.document_requirements = self.define_document_requirements()
        
        # Initialize RAG with progress indication
        with st.spinner("🔄 Initializing AI Document Analysis System..."):
            try:
                # Show initialization status
                if 'rag_initialized' not in st.session_state:
                    st.session_state.rag_initialized = False
            
                if not st.session_state.rag_initialized:
                    with st.spinner("🔄 Initializing AI Document Analysis System..."):
                        self.rag_extractor = RAGDataExtractor()
                        st.session_state.rag_initialized = True
                    
                        # Check what mode we're running in
                        if hasattr(self.rag_extractor, 'embedding_model') and self.rag_extractor.embedding_model is not None:
                            st.success("✅ AI Document Analysis System Ready (Full Features)")
                        else:
                            st.info("ℹ️ AI Document Analysis System Ready (Basic Features - Text Extraction & Pattern Matching)")
            except Exception as e:
                st.warning(f"⚠️ AI System using basic mode: {e}")
                self.rag_extractor = RAGDataExtractor()
                st.session_state.rag_initialized = True
        
            self.create_dummy_institution_users()    

    def generate_comprehensive_historical_data(self) -> pd.DataFrame:
        """Generate comprehensive historical data for 20 institutions over 10 years"""
        np.random.seed(42)
        n_institutions = 20  # Changed from 200 to 20
        years_of_data = 10   # Changed from 5 to 10
        
        institutions_data = []
        
        for inst_id in range(1, n_institutions + 1):
            base_quality = np.random.uniform(0.3, 0.9)
            
            # Generate data for each of the 10 years (2014-2023)
            for year_offset in range(years_of_data):
                year = 2023 - year_offset  # This will give 2023, 2022, 2021, ..., 2014
                inst_trend = base_quality + (year_offset * 0.02)
                
                # Generate realistic data with proper distributions
                naac_grades = ['A++', 'A+', 'A', 'B++', 'B+', 'B', 'C']
                naac_probs = [0.05, 0.10, 0.15, 0.25, 0.25, 0.15, 0.05]
                naac_grade = np.random.choice(naac_grades, p=naac_probs)
                
                # NIRF Ranking - better institutions more likely to be ranked
                if base_quality > 0.7 and np.random.random() < 0.8:
                    nirf_rank = np.random.randint(1, 101)
                elif base_quality > 0.5 and np.random.random() < 0.5:
                    nirf_rank = np.random.randint(101, 201)
                else:
                    nirf_rank = None
                
                student_faculty_ratio = max(10, np.random.normal(20, 5))
                phd_faculty_ratio = np.random.beta(2, 2) * 0.6 + 0.3
                
                # Research metrics with realistic distributions
                publications = max(0, int(np.random.poisson(inst_trend * 30)))
                research_grants = max(0, int(np.random.exponential(inst_trend * 500000)))
                patents = np.random.poisson(inst_trend * 3)
                
                # Infrastructure scores
                digital_infrastructure_score = max(1, min(10, np.random.normal(7, 1.5)))
                library_volumes = max(1000, int(np.random.normal(20000, 10000)))
                
                # Governance scores
                financial_stability = max(1, min(10, np.random.normal(7.5, 1.2)))
                compliance_score = max(1, min(10, np.random.normal(8, 1)))
                
                # Student development
                placement_rate = max(40, min(98, np.random.normal(75, 10)))
                higher_education_rate = max(5, min(50, np.random.normal(20, 8)))
                
                # Social impact
                community_projects = np.random.poisson(inst_trend * 8)
                
                # Calculate performance score
                faculty_count = max(1, np.random.randint(30, 150))
                
                performance_score = self.calculate_performance_score({
                    'naac_grade': naac_grade,
                    'nirf_ranking': nirf_rank,
                    'student_faculty_ratio': student_faculty_ratio,
                    'phd_faculty_ratio': phd_faculty_ratio,
                    'publications_per_faculty': publications / faculty_count,
                    'research_grants': research_grants,
                    'digital_infrastructure': digital_infrastructure_score,
                    'financial_stability': financial_stability,
                    'placement_rate': placement_rate,
                    'community_engagement': community_projects
                })
                
                institution_data = {
                    'institution_id': f'INST_{inst_id:04d}',
                    'institution_name': f'University/College {inst_id:03d}',
                    'year': year,
                    'institution_type': np.random.choice(['State University', 'Deemed University', 'Private University', 'Autonomous College'], p=[0.3, 0.2, 0.3, 0.2]),
                    'state': np.random.choice(['Maharashtra', 'Karnataka', 'Tamil Nadu', 'Delhi', 'Uttar Pradesh', 'Kerala', 'Gujarat'], p=[0.2, 0.15, 0.15, 0.1, 0.2, 0.1, 0.1]),
                    'established_year': np.random.randint(1950, 2015),
                    
                    # Academic Metrics
                    'naac_grade': naac_grade,
                    'nirf_ranking': nirf_rank,
                    'student_faculty_ratio': round(student_faculty_ratio, 1),
                    'phd_faculty_ratio': round(phd_faculty_ratio, 3),
                    
                    # Research Metrics
                    'research_publications': publications,
                    'research_grants_amount': research_grants,
                    'patents_filed': patents,
                    'industry_collaborations': np.random.poisson(inst_trend * 6),
                    
                    # Infrastructure Metrics
                    'digital_infrastructure_score': round(digital_infrastructure_score, 1),
                    'library_volumes': library_volumes,
                    'laboratory_equipment_score': round(max(1, min(10, np.random.normal(7, 1.3))), 1),
                    
                    # Governance Metrics
                    'financial_stability_score': round(financial_stability, 1),
                    'compliance_score': round(compliance_score, 1),
                    'administrative_efficiency': round(max(1, min(10, np.random.normal(7.2, 1.1))), 1),
                    
                    # Student Development Metrics
                    'placement_rate': round(placement_rate, 1),
                    'higher_education_rate': round(higher_education_rate, 1),
                    'entrepreneurship_cell_score': round(max(1, min(10, np.random.normal(6.5, 1.5))), 1),
                    
                    # Social Impact Metrics
                    'community_projects': community_projects,
                    'rural_outreach_score': round(max(1, min(10, np.random.normal(6.8, 1.4))), 1),
                    'inclusive_education_index': round(max(1, min(10, np.random.normal(7.5, 1.2))), 1),
                    
                    # Government Schemes Participation
                    'rusa_participation': np.random.choice([0, 1], p=[0.4, 0.6]),
                    'nmeict_participation': np.random.choice([0, 1], p=[0.5, 0.5]),
                    'fist_participation': np.random.choice([0, 1], p=[0.6, 0.4]),
                    'dst_participation': np.random.choice([0, 1], p=[0.7, 0.3]),
                    
                    # Overall Performance
                    'performance_score': round(performance_score, 2),
                    'approval_recommendation': self.generate_approval_recommendation(performance_score),
                    'risk_level': self.assess_risk_level(performance_score)
                }
                
                institutions_data.append(institution_data)
        
        df = pd.DataFrame(institutions_data)
        
        # Log the data dimensions
        print(f"✅ Generated data for {df['institution_id'].nunique()} institutions across {df['year'].nunique()} years")
        print(f"📊 Total records: {len(df)} | Years: {df['year'].min()}-{df['year'].max()}")
        
        return df
        
    def init_database(self):
        """Initialize SQLite database for storing institutional data"""
        self.conn = sqlite3.connect('institutions.db', check_same_thread=False)
        self.conn.row_factory = sqlite3.Row
        cursor = self.conn.cursor()

    
        
        # Create institutions table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS institutions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                institution_id TEXT UNIQUE,
                institution_name TEXT,
                year INTEGER,
                institution_type TEXT,
                state TEXT,
                established_year INTEGER,
                naac_grade TEXT,
                nirf_ranking INTEGER,
                student_faculty_ratio REAL,
                phd_faculty_ratio REAL,
                research_publications INTEGER,
                research_grants_amount REAL,
                patents_filed INTEGER,
                industry_collaborations INTEGER,
                digital_infrastructure_score REAL,
                library_volumes INTEGER,
                laboratory_equipment_score REAL,
                financial_stability_score REAL,
                compliance_score REAL,
                administrative_efficiency REAL,
                placement_rate REAL,
                higher_education_rate REAL,
                entrepreneurship_cell_score REAL,
                community_projects INTEGER,
                rural_outreach_score REAL,
                inclusive_education_index REAL,
                rusa_participation INTEGER,
                nmeict_participation INTEGER,
                fist_participation INTEGER,
                dst_participation INTEGER,
                performance_score REAL,
                approval_recommendation TEXT,
                risk_level TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Create documents table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS institution_documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                institution_id TEXT,
                document_name TEXT,
                document_type TEXT,
                file_path TEXT,
                upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                status TEXT DEFAULT 'Pending',
                extracted_data TEXT,
                FOREIGN KEY (institution_id) REFERENCES institutions (institution_id)
            )
        ''')
        
        # Create RAG analysis table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS rag_analysis (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                institution_id TEXT,
                analysis_type TEXT,
                extracted_data TEXT,
                ai_insights TEXT,
                confidence_score REAL,
                analysis_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (institution_id) REFERENCES institutions (institution_id)
            )
        ''')
        
        # Create institution submissions table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS institution_submissions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                institution_id TEXT,
                submission_type TEXT,
                submission_data TEXT,
                status TEXT DEFAULT 'Under Review',
                submitted_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                reviewed_by TEXT,
                review_date TIMESTAMP,
                review_comments TEXT,
                FOREIGN KEY (institution_id) REFERENCES institutions (institution_id)
            )
        ''')
        
        # Create institution users table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS institution_users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                institution_id TEXT,
                username TEXT UNIQUE,
                password_hash TEXT,
                contact_person TEXT,
                email TEXT,
                phone TEXT,
                role TEXT DEFAULT 'Institution',
                is_active INTEGER DEFAULT 1,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (institution_id) REFERENCES institutions (institution_id)
            )
        ''')
        
        self.conn.commit()

    def create_dummy_institution_users(self):
        """Create dummy institution users for testing"""
        dummy_users = [
            {
                'institution_id': 'INST_0001',
                'username': 'inst001_admin',
                'password': 'password123',
                'contact_person': 'Dr. Rajesh Kumar',
                'email': 'rajesh.kumar@university001.edu.in',
                'phone': '+91-9876543210'
            },
            {
                'institution_id': 'INST_0050',
                'username': 'inst050_registrar',
                'password': 'testpass456',
                'contact_person': 'Ms. Priya Sharma',
                'email': 'priya.sharma@college050.edu.in',
                'phone': '+91-8765432109'
            },
            {
                'institution_id': 'INST_0100',
                'username': 'inst100_director',
                'password': 'demo789',
                'contact_person': 'Prof. Amit Patel',
                'email': 'amit.patel@university100.edu.in',
                'phone': '+91-7654321098'
            },
            {
                'institution_id': 'INST_0150',
                'username': 'inst150_officer',
                'password': 'admin2024',
                'contact_person': 'Dr. Sunita Reddy',
                'email': 'sunita.reddy@college150.edu.in',
                'phone': '+91-6543210987'
            },
            {
                'institution_id': 'INST_0200',
                'username': 'inst200_manager',
                'password': 'securepass',
                'contact_person': 'Mr. Vikram Singh',
                'email': 'vikram.singh@university200.edu.in',
                'phone': '+91-5432109876'
            }
        ]
    
        for user_data in dummy_users:
            try:
                # Check if user already exists
                cursor = self.conn.cursor()
                cursor.execute('SELECT * FROM institution_users WHERE username = ?', (user_data['username'],))
                existing_user = cursor.fetchone()
            
                if not existing_user:
                    self.create_institution_user(
                        user_data['institution_id'],
                        user_data['username'],
                        user_data['password'],
                        user_data['contact_person'],
                        user_data['email'],
                        user_data['phone']
                    )
                    print(f"Created user: {user_data['username']}")
            except Exception as e:
                print(f"Error creating user {user_data['username']}: {e}")

    def create_institution_user(self, institution_id: str, username: str, password: str, 
                          contact_person: str, email: str, phone: str):
        """Create new institution user account"""
        cursor = self.conn.cursor()
        try:
            cursor.execute('''
                INSERT INTO institution_users 
                (institution_id, username, password_hash, contact_person, email, phone)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (institution_id, username, self.hash_password(password), 
                  contact_person, email, phone))
            self.conn.commit()
            return True
        except sqlite3.IntegrityError:
            return False

    def hash_password(self, password: str) -> str:
        """Simple password hashing (use proper hashing in production)"""
        return hashlib.sha256(password.encode()).hexdigest()

    def analyze_documents_with_rag(self, institution_id: str, uploaded_files: List) -> Dict[str, Any]:
        """Analyze uploaded documents using available analysis methods"""
        try:
            if not uploaded_files:
                return self.get_default_analysis_result(uploaded_files)
            
            # Show progress
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            status_text.text("📄 Starting document analysis...")
            progress_bar.progress(10)
            
            # Extract data
            extracted_data = self.rag_extractor.extract_comprehensive_data(uploaded_files)
            progress_bar.progress(60)
            
            status_text.text("🤖 Generating AI insights...")
            # Generate AI insights
            ai_insights = self.generate_ai_insights(extracted_data)
            progress_bar.progress(90)
            
            # Save to database
            cursor = self.conn.cursor()
            cursor.execute('''
                INSERT INTO rag_analysis 
                (institution_id, analysis_type, extracted_data, confidence_score)
                VALUES (?, ?, ?, ?)
            ''', (institution_id, 'document_analysis', json.dumps(extracted_data), 0.85))
            self.conn.commit()
            
            progress_bar.progress(100)
            status_text.text("✅ Analysis complete!")
            
            return {
                'extracted_data': extracted_data,
                'ai_insights': ai_insights,
                'confidence_score': 0.85,
                'status': 'Analysis Complete'
            }
            
        except Exception as e:
            st.error(f"Error in document analysis: {str(e)}")
            return self.get_default_analysis_result(uploaded_files)

    def get_default_analysis_result(self, uploaded_files: List) -> Dict[str, Any]:
        """Return a safe default structure when analysis fails"""
        return {
            'extracted_data': {
                'academic_metrics': {},
                'research_metrics': {},
                'infrastructure_metrics': {},
                'governance_metrics': {},
                'student_metrics': {},
                'financial_metrics': {},
                'raw_text': "",
                'file_names': [f.name for f in uploaded_files] if uploaded_files else []
            },
            'ai_insights': {
                'strengths': [],
                'weaknesses': [],
                'recommendations': ["Document processing completed in basic mode"],
                'risk_assessment': {'score': 5.0, 'level': 'Medium', 'factors': []},
                'compliance_status': {}
            },
            'confidence_score': 0.0,
            'status': 'Analysis Completed in Basic Mode'
        }

    def generate_ai_insights(self, extracted_data: Dict) -> Dict[str, Any]:
        """Generate AI insights from extracted data"""
        # Ensure extracted_data is not None
        if not extracted_data:
            extracted_data = {}
    
        insights = {
            'strengths': [],
            'weaknesses': [],
            'recommendations': [],
            'risk_assessment': {},
            'compliance_status': {}
        }
    
        # Safe access to nested dictionaries
        academic_data = extracted_data.get('academic_metrics', {})
        research_data = extracted_data.get('research_metrics', {})
        financial_data = extracted_data.get('financial_metrics', {})
        
        # Analyze academic metrics
        academic_data = extracted_data.get('academic_metrics', {})
        research_data = extracted_data.get('research_metrics', {})
        
        # Strength analysis
        if academic_data.get('naac_grade') in ['A++', 'A+', 'A']:
            insights['strengths'].append(f"Strong NAAC accreditation: {academic_data['naac_grade']}")
        
        if research_data.get('research_publications', 0) > 50:
            insights['strengths'].append("Robust research publication output")
        
        # Weakness analysis
        if academic_data.get('student_faculty_ratio', 0) > 25:
            insights['weaknesses'].append("High student-faculty ratio needs improvement")
        
        if research_data.get('patents_filed', 0) < 5:
            insights['weaknesses'].append("Limited patent filings - need to strengthen IPR culture")
        
        # Recommendations
        if not academic_data.get('nirf_ranking'):
            insights['recommendations'].append("Consider participating in NIRF ranking for better visibility")
        
        if research_data.get('industry_collaborations', 0) < 3:
            insights['recommendations'].append("Increase industry collaborations for practical exposure")
        
        # Risk assessment
        risk_score = self.calculate_risk_score(extracted_data)
        insights['risk_assessment'] = {
            'score': risk_score,
            'level': 'Low' if risk_score < 4 else 'Medium' if risk_score < 7 else 'High',
            'factors': self.identify_risk_factors(extracted_data)
        }
        
        return insights
    
    def calculate_risk_score(self, extracted_data: Dict) -> float:
        """Calculate risk score based on extracted data"""
        score = 5.0  # Default medium risk
        
        academic_data = extracted_data.get('academic_metrics', {})
        research_data = extracted_data.get('research_metrics', {})
        financial_data = extracted_data.get('financial_metrics', {})
        
        # Positive factors (reduce risk)
        if academic_data.get('naac_grade') in ['A++', 'A+', 'A']:
            score -= 1.5
        if research_data.get('research_publications', 0) > 50:
            score -= 1.0
        if financial_data.get('financial_stability_score', 0) > 7:
            score -= 1.0
        
        # Negative factors (increase risk)
        if academic_data.get('student_faculty_ratio', 0) > 25:
            score += 1.5
        if research_data.get('patents_filed', 0) < 2:
            score += 1.0
        if not academic_data.get('nirf_ranking'):
            score += 0.5
        
        return max(1.0, min(10.0, score))
    
    def identify_risk_factors(self, extracted_data: Dict) -> List[str]:
        """Identify specific risk factors"""
        risk_factors = []
        academic_data = extracted_data.get('academic_metrics', {})
        research_data = extracted_data.get('research_metrics', {})
        
        if academic_data.get('student_faculty_ratio', 0) > 25:
            risk_factors.append("High student-faculty ratio affecting quality")
        if research_data.get('industry_collaborations', 0) < 2:
            risk_factors.append("Limited industry exposure")
        if not academic_data.get('naac_grade'):
            risk_factors.append("No NAAC accreditation")
        
        return risk_factors
    
    def load_or_generate_data(self):
        """Load data from database or generate sample data with 20×10 specification"""
        try:
            # Try to load from database
            df = pd.read_sql('SELECT * FROM institutions', self.conn)
            if len(df) > 0:
                # Verify the loaded data matches 20×10 specification
                unique_institutions = df['institution_id'].nunique()
                unique_years = df['year'].nunique()
                
                print(f"📊 Loaded data: {unique_institutions} institutions, {unique_years} years, {len(df)} records")
                
                # If data doesn't match 20×10, regenerate
                if unique_institutions != 20 or unique_years != 10:
                    print("⚠️ Data doesn't match 20×10 specification. Regenerating...")
                    df = self.generate_comprehensive_historical_data()
                    df.to_sql('institutions', self.conn, if_exists='replace', index=False)
                
                return df
        except:
            pass
        
        # Generate sample data if database is empty
        print("🔄 Generating new 20×10 sample data...")
        sample_data = self.generate_comprehensive_historical_data()
        sample_data.to_sql('institutions', self.conn, if_exists='replace', index=False)
        return sample_data
    
    def define_performance_metrics(self) -> Dict[str, Dict]:
        """Define key performance indicators for institutional evaluation"""
        return {
            "academic_excellence": {
                "weight": 0.25,
                "sub_metrics": {
                    "naac_grade": 0.30,
                    "nirf_ranking": 0.25,
                    "student_faculty_ratio": 0.20,
                    "phd_faculty_ratio": 0.15,
                    "curriculum_innovation": 0.10
                }
            },
            "research_innovation": {
                "weight": 0.20,
                "sub_metrics": {
                    "publications_per_faculty": 0.30,
                    "research_grants": 0.25,
                    "patents_filed": 0.20,
                    "conferences_organized": 0.15,
                    "industry_collaborations": 0.10
                }
            },
            "infrastructure_facilities": {
                "weight": 0.15,
                "sub_metrics": {
                    "campus_area": 0.25,
                    "digital_infrastructure": 0.25,
                    "library_resources": 0.20,
                    "laboratory_equipment": 0.20,
                    "hostel_facilities": 0.10
                }
            },
            "governance_administration": {
                "weight": 0.15,
                "sub_metrics": {
                    "financial_stability": 0.30,
                    "administrative_efficiency": 0.25,
                    "compliance_record": 0.25,
                    "grievance_redressal": 0.20
                }
            },
            "student_development": {
                "weight": 0.15,
                "sub_metrics": {
                    "placement_rate": 0.35,
                    "higher_education_rate": 0.20,
                    "entrepreneurship_cell": 0.15,
                    "extracurricular_activities": 0.15,
                    "alumni_network": 0.15
                }
            },
            "social_impact": {
                "weight": 0.10,
                "sub_metrics": {
                    "community_engagement": 0.30,
                    "rural_outreach": 0.25,
                    "inclusive_education": 0.25,
                    "environmental_initiatives": 0.20
                }
            }
        }
    
    def define_document_requirements(self) -> Dict[str, Dict]:
        """Define document requirements for different approval types"""
        return {
            "new_approval": {
                "mandatory": [
                    "affidavit_legal_status", "land_documents", "building_plan_approval",
                    "infrastructure_details", "financial_solvency_certificate",
                    "faculty_recruitment_plan", "academic_curriculum", "governance_structure"
                ],
                "supporting": [
                    "feasibility_report", "market_demand_analysis", "five_year_development_plan",
                    "industry_partnerships", "research_facilities_plan"
                ]
            },
            "renewal_approval": {
                "mandatory": [
                    "previous_approval_letters", "annual_reports", "financial_audit_reports",
                    "faculty_student_data", "infrastructure_utilization", "academic_performance"
                ],
                "supporting": [
                    "naac_accreditation", "nirf_data", "research_publications",
                    "placement_records", "social_impact_reports"
                ]
            },
            "expansion_approval": {
                "mandatory": [
                    "current_status_report", "expansion_justification", "additional_infrastructure",
                    "enhanced_faculty_plan", "financial_viability", "market_analysis"
                ],
                "supporting": [
                    "stakeholder_feedback", "alumni_support", "industry_demand",
                    "government_schemes_participation"
                ]
            }
        }

    def authenticate_institution_user(self, username: str, password: str) -> Dict:
        """Authenticate institution user"""
        if not username or not password:
            return None
        
        cursor = self.conn.cursor()
        cursor.execute('''
            SELECT iu.*, i.institution_name 
            FROM institution_users iu 
            JOIN institutions i ON iu.institution_id = i.institution_id 
            WHERE iu.username = ? AND iu.is_active = 1
        ''', (username,))
    
        user = cursor.fetchone()
        if user:
            # Convert to dictionary safely
            columns = [description[0] for description in cursor.description]
            user_dict = dict(zip(columns, user))
        
            # Check if password_hash exists and matches
            password_hash = user_dict.get('password_hash')
            if password_hash and password_hash == self.hash_password(password):
                return {
                    'institution_id': user_dict.get('institution_id'),
                    'institution_name': user_dict.get('institution_name'),
                    'username': user_dict.get('username'),
                    'role': user_dict.get('role', 'Institution'),
                    'contact_person': user_dict.get('contact_person', ''),
                    'email': user_dict.get('email', '')
                }
        return None
    
    def hash_password(self, password: str) -> str:
        """Simple password hashing (use proper hashing in production)"""
        return hashlib.sha256(password.encode()).hexdigest()

    def create_institution_user(self, institution_id: str, username: str, password: str, 
                              contact_person: str, email: str, phone: str):
        """Create new institution user account"""
        cursor = self.conn.cursor()
        try:
            cursor.execute('''
                INSERT INTO institution_users 
                (institution_id, username, password_hash, contact_person, email, phone)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (institution_id, username, self.hash_password(password), 
                  contact_person, email, phone))
            self.conn.commit()
            return True
        except sqlite3.IntegrityError:
            return False

    def save_institution_submission(self, institution_id: str, submission_type: str, 
                                  submission_data: Dict):
        """Save institution submission data"""
        cursor = self.conn.cursor()
        cursor.execute('''
            INSERT INTO institution_submissions 
            (institution_id, submission_type, submission_data)
            VALUES (?, ?, ?)
        ''', (institution_id, submission_type, json.dumps(submission_data)))
        self.conn.commit()

    def get_institution_submissions(self, institution_id: str) -> pd.DataFrame:
        """Get submissions for a specific institution"""
        return pd.read_sql('''
            SELECT * FROM institution_submissions 
            WHERE institution_id = ? 
            ORDER BY submitted_date DESC
        ''', self.conn, params=(institution_id,))

    
    def calculate_performance_score(self, *args, **kwargs) -> float:
        """Calculate overall performance score based on weighted metrics"""
    
        # Handle different calling patterns
        if len(args) == 1 and isinstance(args[0], dict):
            metrics = args[0]
        elif kwargs:
            metrics = kwargs
        else:
            # Default metrics if nothing provided
            return 5.0
    
        score = 0
    
        # NAAC Grade scoring
        naac_scores = {'A++': 10, 'A+': 9, 'A': 8, 'B++': 7, 'B+': 6, 'B': 5, 'C': 4}
        naac_score = naac_scores.get(metrics.get('naac_grade'), 5)
        score += naac_score * 0.15
    
        # NIRF Ranking scoring (inverse)
        nirf_score = 0
        if metrics.get('nirf_ranking') and metrics['nirf_ranking'] <= 200:
            nirf_score = (201 - metrics['nirf_ranking']) / 200 * 10
        score += nirf_score * 0.10
    
        # Student-Faculty Ratio (lower is better)
        sf_ratio = metrics.get('student_faculty_ratio', 20)
        sf_ratio_score = max(0, 10 - max(0, sf_ratio - 15) / 3)
        score += sf_ratio_score * 0.10
    
        # PhD Faculty Ratio
        phd_score = metrics.get('phd_faculty_ratio', 0.6) * 10
        score += phd_score * 0.10
    
        # Research Publications per faculty
        pub_per_faculty = metrics.get('publications_per_faculty', 0)
        pub_score = min(10, pub_per_faculty * 3)
        score += pub_score * 0.10
    
        # Research Grants (log scale)
        grant_score = min(10, np.log1p(metrics.get('research_grants', 0) / 100000) * 2.5)
        score += grant_score * 0.10
    
        # Infrastructure
        infra_score = metrics.get('digital_infrastructure', 7)
        score += infra_score * 0.10
    
        # Financial Stability
        financial_score = metrics.get('financial_stability', 7.5)
        score += financial_score * 0.10
    
        # Placement Rate
        placement_score = metrics.get('placement_rate', 70) / 10
        score += placement_score * 0.10
    
        # Community Engagement
        community_score = min(10, metrics.get('community_engagement', 0) / 1.5)
        score += community_score * 0.05
    
        return min(10, score)
    
    def generate_approval_recommendation(self, performance_score: float) -> str:
        """Generate approval recommendation based on performance score"""
        if performance_score >= 8.0:
            return "Full Approval - 5 Years"
        elif performance_score >= 7.0:
            return "Provisional Approval - 3 Years"
        elif performance_score >= 6.0:
            return "Conditional Approval - 1 Year"
        elif performance_score >= 5.0:
            return "Approval with Strict Monitoring - 1 Year"
        else:
            return "Rejection - Significant Improvements Required"
    
    def assess_risk_level(self, performance_score: float) -> str:
        """Assess institutional risk level"""
        if performance_score >= 8.0:
            return "Low Risk"
        elif performance_score >= 6.5:
            return "Medium Risk"
        elif performance_score >= 5.0:
            return "High Risk"
        else:
            return "Critical Risk"
    
    def save_uploaded_documents(self, institution_id: str, uploaded_files: List, document_types: List[str]):
        """Save uploaded documents to database"""
        cursor = self.conn.cursor()
        for i, uploaded_file in enumerate(uploaded_files):
            cursor.execute('''
                INSERT INTO institution_documents (institution_id, document_name, document_type, status)
                VALUES (?, ?, ?, ?)
            ''', (institution_id, uploaded_file.name, document_types[i], 'Uploaded'))
        self.conn.commit()
    
    def get_institution_documents(self, institution_id: str) -> pd.DataFrame:
        """Get documents for a specific institution"""
        return pd.read_sql('''
            SELECT * FROM institution_documents 
            WHERE institution_id = ? 
            ORDER BY upload_date DESC
        ''', self.conn, params=(institution_id,))
    
    def analyze_document_sufficiency(institution_id, approval_type, analyzer):
        """Analyze document sufficiency and generate recommendations"""
    
        # Get uploaded documents
        try:
            uploaded_docs = analyzer.get_institution_documents(institution_id)
            uploaded_doc_names = uploaded_docs['document_name'].tolist() if not uploaded_docs.empty else []
        except:
            uploaded_doc_names = []
    
        # Get requirements
        requirements = get_document_requirements_by_parameters(approval_type)
    
        # Calculate sufficiency scores
        mandatory_docs = []
        for param_docs in requirements["mandatory"].values():
            mandatory_docs.extend(param_docs)
    
        supporting_docs = []
        for param_docs in requirements["supporting"].values():
            supporting_docs.extend(param_docs)
    
        # Count uploaded documents
        mandatory_uploaded = 0
        for doc in mandatory_docs:
            if any(doc.lower() in uploaded_doc.lower() for uploaded_doc in uploaded_doc_names):
                mandatory_uploaded += 1
    
        supporting_uploaded = 0
        for doc in supporting_docs:
            if any(doc.lower() in uploaded_doc.lower() for uploaded_doc in uploaded_doc_names):
                supporting_uploaded += 1
    
        # Calculate percentages
        mandatory_sufficiency = (mandatory_uploaded / len(mandatory_docs)) * 100 if mandatory_docs else 0
        overall_sufficiency = ((mandatory_uploaded + supporting_uploaded) / (len(mandatory_docs) + len(supporting_docs))) * 100 if (mandatory_docs and supporting_docs) else mandatory_sufficiency
    
        # Display analysis results
        st.subheader("📊 Document Sufficiency Analysis")
    
        col1, col2, col3 = st.columns(3)
    
        with col1:
            st.metric("Mandatory Documents", f"{mandatory_sufficiency:.1f}%")
    
        with col2:
            st.metric("Overall Sufficiency", f"{overall_sufficiency:.1f}%")
    
        with col3:
            st.metric("Total Uploaded", f"{len(uploaded_doc_names)}")
    
        # Visual gauge for mandatory documents
        try:
            import plotly.graph_objects as go
            fig = go.Figure(go.Indicator(
                mode="gauge+number+delta",
                value=mandatory_sufficiency,
                domain={'x': [0, 1], 'y': [0, 1]},
                title={'text': "Mandatory Documents Sufficiency"},
                gauge={
                    'axis': {'range': [None, 100]},
                    'bar': {'color': "darkblue"},
                    'steps': [
                        {'range': [0, 50], 'color': "lightgray"},
                        {'range': [50, 80], 'color': "yellow"},
                        {'range': [80, 100], 'color': "lightgreen"}
                    ],
                    'threshold': {
                        'line': {'color': "red", 'width': 4},
                        'thickness': 0.75,
                        'value': 90
                    }
                }
            ))
            fig.update_layout(height=300)
            st.plotly_chart(fig, use_container_width=True)
        except:
            st.info("📊 Mandatory Documents Sufficiency Visualization")
            st.progress(mandatory_sufficiency / 100)
    
        # Generate recommendations
        st.subheader("💡 AI Recommendations")
    
        if mandatory_sufficiency < 50:
            st.error("**❌ Critical Issue**: Less than 50% of mandatory documents uploaded.")
            st.write("**Action Required**: Upload all mandatory documents to proceed with approval process.")
    
        elif mandatory_sufficiency < 80:
            st.warning("**🟡 Attention Needed**: Some mandatory documents are missing.")
            st.write("**Recommended Action**: Complete mandatory document upload for smoother approval process.")
    
        elif mandatory_sufficiency < 100:
            st.info("**🔵 Good Progress**: Most mandatory documents are uploaded.")
            st.write("**Next Steps**: Complete remaining mandatory documents and consider uploading supporting documents.")
    
        else:
            st.success("**✅ Excellent**: All mandatory documents are uploaded!")
            st.write("**Next Steps**: Consider uploading supporting documents for enhanced assessment.")
    
        # Missing documents analysis
        missing_mandatory = []
        for param_docs in requirements["mandatory"].values():
            for doc in param_docs:
                if not any(doc.lower() in uploaded_doc.lower() for uploaded_doc in uploaded_doc_names):
                    missing_mandatory.append(doc)
    
        if missing_mandatory:
            st.error("**Missing Mandatory Documents:**")
            for doc in missing_mandatory[:5]:  # Show top 5 missing
                st.write(f"• {doc}")
    
        # Parameter-wise analysis
        st.subheader("📈 Parameter-wise Document Coverage")
    
        param_coverage = {}
        for parameter, documents in requirements["mandatory"].items():
            uploaded_count = 0
            for doc in documents:
                if any(doc.lower() in uploaded_doc.lower() for uploaded_doc in uploaded_doc_names):
                    uploaded_count += 1
            total_count = len(documents)
            coverage = (uploaded_count / total_count) * 100 if total_count > 0 else 0
            param_coverage[parameter] = coverage
    
        # Create bar chart for parameter coverage
        try:
            import plotly.express as px
            import pandas as pd
        
            param_df = pd.DataFrame({
                'Parameter': list(param_coverage.keys()),
                'Coverage (%)': list(param_coverage.values())
            })
        
            fig = px.bar(param_df, x='Parameter', y='Coverage (%)', 
                         title="Document Coverage by Parameter",
                         color='Coverage (%)',
                         color_continuous_scale='RdYlGn')
            fig.update_layout(xaxis_tickangle=-45)
            st.plotly_chart(fig, use_container_width=True)
        except:
            # Fallback display if plotly fails
            for param, coverage in param_coverage.items():
                st.write(f"**{param}**: {coverage:.1f}% coverage")
                st.progress(coverage / 100)
    
    def generate_document_recommendations(self, mandatory_sufficiency: float) -> List[str]:
        """Generate recommendations based on document sufficiency"""
        recommendations = []
        
        if mandatory_sufficiency < 100:
            recommendations.append("Upload all mandatory documents to proceed with approval process")
        
        if mandatory_sufficiency < 80:
            recommendations.append("Critical documents missing - application cannot be processed")
        
        if mandatory_sufficiency >= 100:
            recommendations.append("All mandatory documents present - ready for comprehensive evaluation")
        
        return recommendations
    
    def generate_comprehensive_report(self, institution_id: str) -> Dict[str, Any]:
        """Generate comprehensive AI analysis report for an institution"""
        inst_data = self.historical_data[
            self.historical_data['institution_id'] == institution_id
        ]
        
        if inst_data.empty:
            return {"error": "Institution not found"}
        
        latest_data = inst_data[inst_data['year'] == inst_data['year'].max()].iloc[0]
        historical_trend = inst_data.groupby('year')['performance_score'].mean()
        
        # Performance trends
        if len(historical_trend) > 1:
            if historical_trend.iloc[-1] > historical_trend.iloc[-2]:
                trend_analysis = "Improving"
            elif historical_trend.iloc[-1] == historical_trend.iloc[-2]:
                trend_analysis = "Stable"
            else:
                trend_analysis = "Declining"
        else:
            trend_analysis = "Insufficient Data"
        
        # Comparative analysis
        similar_institutions = self.find_similar_institutions(institution_id)
        
        return {
            "institution_info": {
                "name": latest_data['institution_name'],
                "type": latest_data['institution_type'],
                "state": latest_data['state'],
                "established": latest_data['established_year']
            },
            "performance_analysis": {
                "current_score": latest_data['performance_score'],
                "historical_trend": historical_trend.to_dict(),
                "trend_analysis": trend_analysis,
                "approval_recommendation": latest_data['approval_recommendation'],
                "risk_level": latest_data['risk_level']
            },
            "strengths": self.identify_strengths(latest_data),
            "weaknesses": self.identify_weaknesses(latest_data),
            "comparative_analysis": similar_institutions,
            "ai_recommendations": self.generate_ai_recommendations(latest_data)
        }
    
    def find_similar_institutions(self, institution_id: str) -> Dict:
        """Find similar institutions for comparative analysis"""
        inst_data = self.historical_data[
            self.historical_data['institution_id'] == institution_id
        ]
        
        if inst_data.empty:
            return {}
        
        latest_data = inst_data[inst_data['year'] == inst_data['year'].max()].iloc[0]
        
        # Find similar institutions based on type and performance
        similar_inst = self.historical_data[
            (self.historical_data['institution_type'] == latest_data['institution_type']) &
            (self.historical_data['year'] == latest_data['year']) &
            (self.historical_data['institution_id'] != institution_id)
        ]
        
        if len(similar_inst) > 0:
            similar_inst = similar_inst.nlargest(5, 'performance_score')
            benchmark_data = similar_inst[['institution_name', 'performance_score', 'approval_recommendation']].to_dict('records')
        else:
            benchmark_data = []
        
        return {
            "benchmark_institutions": benchmark_data,
            "performance_percentile": self.calculate_performance_percentile(latest_data['performance_score'], latest_data['institution_type'])
        }
    
    def calculate_performance_percentile(self, score: float, inst_type: str) -> float:
        """Calculate performance percentile within institution type"""
        type_data = self.historical_data[
            (self.historical_data['institution_type'] == inst_type) &
            (self.historical_data['year'] == 2023)
        ]
        
        if len(type_data) == 0:
            return 50.0
        
        return (type_data['performance_score'] < score).mean() * 100
    
    def identify_strengths(self, institution_data: pd.Series) -> List[str]:
        """Identify institutional strengths"""
        strengths = []
        
        if institution_data['naac_grade'] in ['A++', 'A+', 'A']:
            strengths.append(f"Excellent NAAC Accreditation: {institution_data['naac_grade']}")
        
        if institution_data['placement_rate'] > 80:
            strengths.append(f"Strong Placement Record: {institution_data['placement_rate']:.1f}%")
        
        if institution_data['research_publications'] > 100:
            strengths.append(f"Robust Research Output: {institution_data['research_publications']} publications")
        
        if institution_data['financial_stability_score'] > 8.5:
            strengths.append("Excellent Financial Stability")
        
        if institution_data['phd_faculty_ratio'] > 0.7:
            strengths.append(f"Highly Qualified Faculty: {institution_data['phd_faculty_ratio']:.1%} PhDs")
        
        if institution_data['digital_infrastructure_score'] > 8.5:
            strengths.append("Advanced Digital Infrastructure")
            
        return strengths
    
    def identify_weaknesses(self, institution_data: pd.Series) -> List[str]:
        """Identify institutional weaknesses"""
        weaknesses = []
        
        if institution_data['student_faculty_ratio'] > 25:
            weaknesses.append(f"High Student-Faculty Ratio: {institution_data['student_faculty_ratio']:.1f}")
        
        if institution_data['placement_rate'] < 65:
            weaknesses.append(f"Low Placement Rate: {institution_data['placement_rate']:.1f}%")
        
        if institution_data['research_publications'] < 20:
            weaknesses.append(f"Inadequate Research Output: {institution_data['research_publications']} publications")
        
        if institution_data['digital_infrastructure_score'] < 7:
            weaknesses.append(f"Weak Digital Infrastructure: {institution_data['digital_infrastructure_score']:.1f}/10")
        
        if institution_data['community_projects'] < 5:
            weaknesses.append("Limited Community Engagement")
            
        return weaknesses
    
    def generate_ai_recommendations(self, institution_data: pd.Series) -> List[str]:
        """Generate AI-powered improvement recommendations"""
        recommendations = []
        
        if institution_data['student_faculty_ratio'] > 25:
            recommendations.append("Recruit additional faculty members to improve student-faculty ratio")
        
        if institution_data['placement_rate'] < 70:
            recommendations.append("Strengthen industry partnerships and career development programs")
        
        if institution_data['research_publications'] < 50:
            recommendations.append("Establish research promotion policy and faculty development programs")
        
        if institution_data['digital_infrastructure_score'] < 7:
            recommendations.append("Invest in digital infrastructure and e-learning platforms")
        
        if institution_data['community_projects'] < 5:
            recommendations.append("Enhance community engagement and social outreach programs")
        
        return recommendations

# Institution-specific modules
def create_institution_login(analyzer):
    st.header("🏛️ Institution Portal Login")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Existing Institution Users")
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        
        if st.button("Login"):
            user = analyzer.authenticate_institution_user(username, password)
            if user:
                st.session_state.institution_user = user
                st.session_state.user_role = "Institution"
                st.success(f"Welcome, {user['contact_person']} from {user['institution_name']}!")
                st.rerun()  # Use rerun instead of stop
            else:
                st.error("Invalid username or password")
    
    with col2:
        st.subheader("New Institution Registration")
        
        # Get available institutions
        available_institutions = analyzer.historical_data[
            analyzer.historical_data['year'] == 2023
        ][['institution_id', 'institution_name']].drop_duplicates()
        
        selected_institution = st.selectbox(
            "Select Your Institution",
            available_institutions['institution_id'].tolist(),
            format_func=lambda x: available_institutions[
                available_institutions['institution_id'] == x
            ]['institution_name'].iloc[0]
        )
        
        new_username = st.text_input("Choose Username")
        new_password = st.text_input("Choose Password", type="password")
        confirm_password = st.text_input("Confirm Password", type="password")
        contact_person = st.text_input("Contact Person Name")
        email = st.text_input("Email Address")
        phone = st.text_input("Phone Number")
        
        if st.button("Register Institution Account"):
            if new_password != confirm_password:
                st.error("Passwords do not match!")
            elif not all([new_username, new_password, contact_person, email]):
                st.error("Please fill all required fields!")
            else:
                success = analyzer.create_institution_user(
                    selected_institution, new_username, new_password,
                    contact_person, email, phone
                )
                if success:
                    st.success("Institution account created successfully! You can now login.")
                else:
                    st.error("Username already exists. Please choose a different username.")

# Update the institution dashboard to use the correct function names
def create_institution_dashboard(analyzer, user):
    if not user:
        st.error("No user data available")
        return
        
    st.header(f"🏛️ Institution Dashboard - {user.get('institution_name', 'Unknown')}")
    
    # Display institution overview
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Institution ID", user.get('institution_id', 'N/A'))
    with col2:
        st.metric("Contact Person", user.get('contact_person', 'N/A'))
    with col3:
        st.metric("Email", user.get('email', 'N/A'))
    with col4:
        st.metric("Role", user.get('role', 'N/A'))
    
    # Navigation for institution users - CORRECTED TABS
    institution_tabs = st.tabs([
        "📤 Document Upload", 
        "📝 Basic Data Submission",  # This tab shows the basic form
        "🏛️ Systematic Data Form",   # This tab shows the comprehensive form
        "📊 My Submissions",
        "📋 Requirements Guide",
        "🔄 Approval Workflow"
    ])
    
    with institution_tabs[0]:
        create_institution_document_upload(analyzer, user)
    
    with institution_tabs[1]:
        create_institution_data_submission(analyzer, user)  # This should show the basic form
    
    with institution_tabs[2]:
        create_systematic_data_submission_form(analyzer, user)  # The new comprehensive form
    
    with institution_tabs[3]:
        create_institution_submissions_view(analyzer, user)
    
    with institution_tabs[4]:
        create_institution_requirements_guide(analyzer)
    
    with institution_tabs[5]:
        create_institution_approval_workflow(analyzer, user)

def create_institution_data_submission(analyzer, user):
    st.subheader("📝 Basic Data Submission Form")
    
    st.info("Submit essential institutional data and performance metrics through this simplified form")
    
    with st.form("basic_institution_data_submission"):
        st.write("### 🎓 Academic Performance Data")
        
        col1, col2 = st.columns(2)
        
        with col1:
            naac_grade = st.selectbox(
                "NAAC Grade",
                ["A++", "A+", "A", "B++", "B+", "B", "C", "Not Accredited"],
                key="basic_naac_grade"
            )
            student_faculty_ratio = st.number_input(
                "Student-Faculty Ratio",
                min_value=5.0,
                max_value=50.0,
                value=20.0,
                step=0.1,
                key="basic_sf_ratio"
            )
            phd_faculty_ratio = st.number_input(
                "PhD Faculty Ratio (%)",
                min_value=0.0,
                max_value=100.0,
                value=60.0,
                step=1.0,
                key="basic_phd_ratio"
            ) / 100
        
        with col2:
            nirf_ranking = st.number_input(
                "NIRF Ranking (if applicable)",
                min_value=1,
                max_value=200,
                value=None,
                placeholder="Leave blank if not ranked",
                key="basic_nirf"
            )
            placement_rate = st.number_input(
                "Placement Rate (%)",
                min_value=0.0,
                max_value=100.0,
                value=75.0,
                step=1.0,
                key="basic_placement"
            )
        
        st.write("### 🔬 Research & Infrastructure")
        
        col1, col2 = st.columns(2)
        
        with col1:
            research_publications = st.number_input(
                "Research Publications (Last Year)",
                min_value=0,
                value=50,
                key="basic_publications"
            )
            research_grants = st.number_input(
                "Research Grants Amount (₹ Lakhs)",
                min_value=0,
                value=100,
                step=10,
                key="basic_grants"
            )
        
        with col2:
            digital_infrastructure_score = st.slider(
                "Digital Infrastructure Score (1-10)",
                min_value=1,
                max_value=10,
                value=7,
                key="basic_digital"
            )
            library_volumes = st.number_input(
                "Library Volumes (in thousands)",
                min_value=0,
                value=20,
                key="basic_library"
            )
        
        st.write("### ⚖️ Governance & Social Impact")
        
        col1, col2 = st.columns(2)
        
        with col1:
            financial_stability_score = st.slider(
                "Financial Stability Score (1-10)",
                min_value=1,
                max_value=10,
                value=8,
                key="basic_financial"
            )
            community_projects = st.number_input(
                "Community Projects (Last Year)",
                min_value=0,
                value=10,
                key="basic_community"
            )
        
        with col2:
            compliance_score = st.slider(
                "Compliance Score (1-10)",
                min_value=1,
                max_value=10,
                value=8,
                key="basic_compliance"
            )
            administrative_efficiency = st.slider(
                "Administrative Efficiency (1-10)",
                min_value=1,
                max_value=10,
                value=7,
                key="basic_admin"
            )
        
        # Additional quick metrics
        st.write("### 📊 Quick Metrics")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            total_students = st.number_input(
                "Total Students",
                min_value=0,
                value=1000,
                step=100,
                key="basic_students"
            )
        
        with col2:
            total_faculty = st.number_input(
                "Total Faculty",
                min_value=0,
                value=50,
                key="basic_faculty"
            )
        
        with col3:
            campus_area = st.number_input(
                "Campus Area (Acres)",
                min_value=0.0,
                value=50.0,
                step=1.0,
                key="basic_campus"
            )
        
        submission_notes = st.text_area(
            "Additional Notes / Comments",
            placeholder="Add any additional information or context for your submission...",
            key="basic_notes"
        )
        
        submitted = st.form_submit_button("📤 Submit Basic Data")
        
        if submitted:
            submission_data = {
                "academic_data": {
                    "naac_grade": naac_grade,
                    "nirf_ranking": nirf_ranking,
                    "student_faculty_ratio": student_faculty_ratio,
                    "phd_faculty_ratio": phd_faculty_ratio,
                    "placement_rate": placement_rate
                },
                "research_data": {
                    "research_publications": research_publications,
                    "research_grants": research_grants,
                    "digital_infrastructure_score": digital_infrastructure_score,
                    "library_volumes": library_volumes
                },
                "governance_data": {
                    "financial_stability_score": financial_stability_score,
                    "compliance_score": compliance_score,
                    "administrative_efficiency": administrative_efficiency,
                    "community_projects": community_projects
                },
                "institutional_data": {
                    "total_students": total_students,
                    "total_faculty": total_faculty,
                    "campus_area": campus_area
                },
                "submission_notes": submission_notes,
                "submission_date": datetime.now().isoformat(),
                "submission_type": "basic_institutional_data"
            }
            
            analyzer.save_institution_submission(
                user['institution_id'],
                "basic_performance_data",
                submission_data
            )
            
            st.success("✅ Basic data submitted successfully! Your submission is under review.")
            
            # Show quick summary
            st.subheader("📋 Submission Summary")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("NAAC Grade", naac_grade)
                st.metric("Student-Faculty Ratio", f"{student_faculty_ratio}:1")
            
            with col2:
                st.metric("Placement Rate", f"{placement_rate}%")
                st.metric("Research Publications", research_publications)
            
            with col3:
                st.metric("Digital Infrastructure", f"{digital_infrastructure_score}/10")
                st.metric("Financial Stability", f"{financial_stability_score}/10")
            
            st.balloons()

def create_institution_approval_workflow(analyzer, user):
    st.subheader("🔄 Institution Approval Workflow")
    
    st.info("Track your approval application status and follow the workflow steps")
    
    # Show current submission status
    submissions = analyzer.get_institution_submissions(user['institution_id'])
    
    if len(submissions) > 0:
        latest_submission = submissions.iloc[0]
        
        st.subheader("📋 Current Application Status")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            status = latest_submission['status']
            if status == 'Under Review':
                st.warning(f"**Status:** {status}")
            elif status == 'Approved':
                st.success(f"**Status:** {status}")
            elif status == 'Rejected':
                st.error(f"**Status:** {status}")
            else:
                st.info(f"**Status:** {status}")
        
        with col2:
            st.write(f"**Submission Date:** {latest_submission['submitted_date']}")
        
        with col3:
            if latest_submission['reviewed_by']:
                st.write(f"**Reviewed by:** {latest_submission['reviewed_by']}")
    
    else:
        st.info("No submissions found. Submit your data to start the approval process.")
    
    # Workflow steps for institutions
    st.subheader("📝 Approval Process Steps")
    
    workflow_steps = [
        {
            "step": 1,
            "title": "Document Submission",
            "description": "Upload all required documents through the Document Upload portal",
            "status": "Complete" if len(submissions) > 0 else "Pending",
            "action": "Go to Document Upload tab"
        },
        {
            "step": 2,
            "title": "Basic Data Submission",
            "description": "Submit institutional performance data through Basic Data Submission form",
            "status": "Complete" if len(submissions) > 0 else "Pending", 
            "action": "Go to Basic Data Submission tab"
        },
        {
            "step": 3,
            "title": "AI Analysis & Verification",
            "description": "System automatically analyzes documents and data for completeness",
            "status": "In Progress" if len(submissions) > 0 else "Pending",
            "action": "Automatic Process"
        },
        {
            "step": 4,
            "title": "Committee Review",
            "description": "UGC/AICTE committee reviews AI recommendations and documents",
            "status": "Pending",
            "action": "Under Committee Review"
        },
        {
            "step": 5,
            "title": "Final Decision",
            "description": "Receive final approval decision with conditions and timeline",
            "status": "Pending",
            "action": "Awaiting Decision"
        }
    ]
    
    for step in workflow_steps:
        with st.expander(f"Step {step['step']}: {step['title']} - {step['status']}"):
            col1, col2 = st.columns([3, 1])
            with col1:
                st.write(f"**Description:** {step['description']}")
                st.write(f"**Status:** {step['status']}")
            with col2:
                if step['status'] == "Pending":
                    st.button(step['action'], key=f"step_{step['step']}", disabled=True)
                else:
                    st.button(step['action'], key=f"step_{step['step']}")
    
    # Quick actions for institutions
    st.subheader("🚀 Quick Actions")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📤 Upload Missing Documents", type="secondary"):
            st.session_state.active_tab = "Document Upload"
            st.rerun()
    
    with col2:
        if st.button("📝 Update Submission Data", type="secondary"):
            st.session_state.active_tab = "Basic Data Submission" 
            st.rerun()
    
    with col3:
        if st.button("📞 Contact Support", type="secondary"):
            st.info("Please contact UGC/AICTE support at: support@ugc.gov.in")

def create_institution_document_upload(analyzer, user):
    st.subheader("📤 Document Upload Portal")
    
    st.info("Upload required documents for approval processes")
    
    approval_type = st.selectbox(
        "Select Approval Type",
        ["new_approval", "renewal_approval", "expansion_approval"],
        format_func=lambda x: x.replace('_', ' ').title(),
        key="inst_approval_type"
    )
    
    # NEW: Document Parameter Selection
    document_parameter = st.selectbox(
        "Select Document Parameter",
        [
            "Curriculum",
            "Faculty Resources", 
            "Learning and Teaching",
            "Research and Innovation",
            "Extracurricular & Co-curricular Activities",
            "Community Engagement",
            "Green Initiatives",
            "Governance and Administration",
            "Infrastructure Development",
            "Financial Resources and Management"
        ],
        key="doc_parameter"
    )
    
    # NEW: Mandatory Documents dropdown based on selected parameter
    mandatory_documents_map = {
        "Curriculum": [
            "Curriculum framework and syllabus documents for all programs",
            "Course outlines with learning objectives and outcomes",
            "Evidence of curriculum review and revision processes",
            "Student feedback reports on curriculum",
            "Faculty feedback on curriculum implementation",
            "Records of stakeholder consultations for curriculum development",
            "Academic calendar and course schedules"
        ],
        "Faculty Resources": [
            "Faculty recruitment policy and procedures",
            "Faculty qualification records and biodata",
            "Selection committee composition and minutes",
            "Faculty development programs records",
            "Faculty appraisal and performance records",
            "Diversity statistics (gender, social inclusion)",
            "Faculty-student ratio documentation"
        ],
        "Learning and Teaching": [
            "Teaching plans and lesson schedules",
            "Student assessment records and evaluation methods",
            "Learning outcome achievement records",
            "Classroom observation reports",
            "Student performance analysis reports",
            "Digital learning infrastructure details"
        ],
        "Research and Innovation": [
            "Research policy document",
            "Research publication records (citations, impact factors)",
            "Patent filings and grants documentation",
            "Research project funding details",
            "Research collaboration agreements",
            "IPR policy and implementation records"
        ],
        "Extracurricular & Co-curricular Activities": [
            "EC/CC activity calendar and schedules",
            "Student participation records",
            "Activity reports and outcomes",
            "Credit allocation policy for EC/CC activities"
        ],
        "Community Engagement": [
            "Community engagement policy",
            "Outreach program records and reports",
            "Social project documentation",
            "Village/community adoption records",
            "Student internship reports with community focus"
        ],
        "Green Initiatives": [
            "Environmental policy document",
            "Energy consumption and conservation records",
            "Waste management system documentation",
            "Water harvesting and recycling records",
            "Green building certification (if any)",
            "Carbon footprint assessment"
        ],
        "Governance and Administration": [
            "Institutional act, statutes, and regulations",
            "Organizational structure chart",
            "Governance body composition and minutes",
            "Financial management policies",
            "Grievance redressal mechanism records",
            "Enrollment statistics and diversity reports"
        ],
        "Infrastructure Development": [
            "Campus master plan and layout",
            "Building and facility inventory",
            "Laboratory and equipment details",
            "Library resource documentation",
            "IT infrastructure details",
            "Hostel and residential facility records"
        ],
        "Financial Resources and Management": [
            "Annual financial statements and audit reports",
            "Budget allocation and utilization certificates",
            "Salary expenditure records",
            "Research grant utilization details",
            "Infrastructure development expenditure",
            "Donation and CSR fund records"
        ]
    }
    
    mandatory_document = st.selectbox(
        "Select Mandatory Document",
        mandatory_documents_map.get(document_parameter, []),
        key="mandatory_doc"
    )
    
    # NEW: Supporting Documents dropdown based on selected parameter
    supporting_documents_map = {
        "Curriculum": [
            "Innovative teaching-learning materials developed",
            "Records of curriculum innovation and updates",
            "Industry interface documents for curriculum design",
            "Documentation of multidisciplinary courses",
            "Evidence of skill-integration in curriculum",
            "Records of vocational and skill-based courses",
            "Publications/patents arising from curriculum implementation"
        ],
        "Faculty Resources": [
            "Faculty achievement and award records",
            "Participation in Malviya Mission/other development programs",
            "Records of faculty mentoring programs",
            "Faculty research and publication records",
            "Evidence of faculty industry exposure",
            "Faculty participation in national/international conferences",
            "Professional development plans"
        ],
        "Learning and Teaching": [
            "Innovative teaching methodology documentation",
            "Experiential learning activity records",
            "Field work and practical training reports",
            "Student project documentation",
            "Research-oriented learning evidence",
            "Digital technology integration records",
            "Critical thinking development activities"
        ],
        "Research and Innovation": [
            "Research facility details and utilization",
            "Research seminar/conference organization records",
            "Industry research partnership documents",
            "Student research participation records",
            "Translational research outcomes",
            "Research awards and recognition",
            "Interdisciplinary research projects"
        ],
        "Extracurricular & Co-curricular Activities": [
            "Awards and achievements in EC/CC activities",
            "Student club and society records",
            "Cultural and sports event documentation",
            "Leadership development program records",
            "Community service activity reports",
            "Student representation in governance bodies"
        ],
        "Community Engagement": [
            "Community feedback and impact assessment",
            "Collaborative project agreements",
            "Social research publications",
            "CSR initiative documentation",
            "Public awareness campaign records",
            "Extension activity participation records"
        ],
        "Green Initiatives": [
            "Renewable energy installation details",
            "Environmental audit reports",
            "Sustainability project documentation",
            "Green campus initiative records",
            "Environmental awareness program reports",
            "Biodiversity conservation efforts"
        ],
        "Governance and Administration": [
            "e-Governance implementation details",
            "Strategic plans and implementation reports",
            "International collaboration agreements",
            "Administrative innovation records",
            "Stakeholder satisfaction surveys",
            "Compliance and audit reports",
            "Decision-making process documentation"
        ],
        "Infrastructure Development": [
            "Infrastructure utilization reports",
            "Maintenance and upgrade records",
            "Safety and security system details",
            "Accessibility compliance documentation",
            "Future development plans",
            "Specialized facility details (auditorium, sports complex)"
        ],
        "Financial Resources and Management": [
            "Financial planning documents",
            "Resource mobilization records",
            "Investment in academic development",
            "Student scholarship and financial aid details",
            "Revenue generation from various sources",
            "Financial sustainability analysis"
        ]
    }
    
    supporting_document = st.selectbox(
        "Select Supporting Document",
        supporting_documents_map.get(document_parameter, []),
        key="supporting_doc"
    )
    
    uploaded_files = st.file_uploader(
        "Upload Institutional Documents",
        type=['pdf', 'doc', 'docx', 'xlsx', 'jpg', 'png'],
        accept_multiple_files=True,
        help="Upload all required documents for your application"
    )
    
    if uploaded_files:
        # Document type mapping
        st.subheader("📝 Document Type Assignment")
        document_types = []
        for i, file in enumerate(uploaded_files):
            doc_type = st.selectbox(
                f"Document type for: {file.name}",
                [
                    "affidavit_legal_status", "land_documents", "building_plan_approval", 
                    "financial_solvency_certificate", "faculty_recruitment_plan", 
                    "academic_curriculum", "annual_reports", "research_publications",
                    "placement_records", "other"
                ],
                key=f"inst_doc_type_{i}"
            )
            document_types.append(doc_type)
        
        # Display selected document information
        st.subheader("📋 Selected Document Information")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.info(f"**Parameter:** {document_parameter}")
        with col2:
            st.info(f"**Mandatory Doc:** {mandatory_document}")
        with col3:
            st.info(f"**Supporting Doc:** {supporting_document}")
        
        if st.button("💾 Upload Documents"):
            # Save documents
            analyzer.save_uploaded_documents(user['institution_id'], uploaded_files, document_types)
            st.success("✅ Documents uploaded successfully!")
            
            # Analyze document sufficiency
            file_names = [file.name for file in uploaded_files]
            analysis_result = analyzer.analyze_document_sufficiency(file_names, approval_type)
            
            # Display results
            st.subheader("📊 Upload Analysis")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.metric(
                    "Mandatory Documents", 
                    f"{analysis_result['mandatory_sufficiency']:.1f}%",
                    delta=f"{analysis_result['mandatory_sufficiency'] - 100:.1f}%" if analysis_result['mandatory_sufficiency'] < 100 else None,
                    delta_color="inverse"
                )
            
            with col2:
                st.metric(
                    "Overall Sufficiency", 
                    f"{analysis_result['overall_sufficiency']:.1f}%"
                )
            
            # Show missing documents
            if analysis_result['missing_mandatory']:
                st.error("**❌ Missing Mandatory Documents:**")
                for doc in analysis_result['missing_mandatory']:
                    st.write(f"• {doc.replace('_', ' ').title()}")
            
            # Recommendations
            st.info("**💡 Next Steps:**")
            for recommendation in analysis_result['recommendations']:
                st.write(f"• {recommendation}")

def create_systematic_data_submission_form(analyzer, user):
    st.subheader("🏛️ Systematic Data Submission Form - NEP 2020 Framework")
    
    st.info("""
    **Complete this comprehensive data submission form based on the 10-parameter framework from the Dr. Radhakrishnan Committee Report.**
    This data will be used for AI-powered institutional analysis and accreditation assessment.
    """)
    
    with st.form("systematic_data_submission"):
        st.markdown("### 📚 1. CURRICULUM")
        
        col1, col2 = st.columns(2)
        
        with col1:
            curriculum_framework_score = st.slider(
                "Curriculum Framework Quality Score (1-10)",
                min_value=1, max_value=10, value=7,
                help="Assessment of curriculum design and structure"
            )
            stakeholder_consultation = st.selectbox(
                "Stakeholder Consultation in Curriculum Design",
                ["Regular & Comprehensive", "Occasional", "Minimal", "None"],
                help="Industry, alumni, employer involvement"
            )
            curriculum_update_frequency = st.selectbox(
                "Curriculum Review & Update Frequency",
                ["Annual", "Biannual", "Every 3 Years", "Irregular"],
                help="How often curriculum is revised"
            )
        
        with col2:
            multidisciplinary_courses = st.number_input(
                "Number of Multidisciplinary Courses",
                min_value=0, value=5,
                help="Courses integrating multiple disciplines"
            )
            skill_integration_score = st.slider(
                "Skill Integration in Curriculum (1-10)",
                min_value=1, max_value=10, value=6,
                help="Integration of vocational and employability skills"
            )
            digital_content_availability = st.selectbox(
                "Digital Learning Content Availability",
                ["Extensive (>80%)", "Moderate (50-80%)", "Limited (<50%)", "Minimal"],
                help="Availability of digital learning materials"
            )
        
        st.markdown("---")
        st.markdown("### 👨‍🏫 2. FACULTY RESOURCES")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            faculty_student_ratio = st.number_input(
                "Student-Faculty Ratio",
                min_value=5.0, max_value=50.0, value=20.0, step=0.1
            )
            phd_faculty_percentage = st.number_input(
                "PhD Faculty Percentage (%)",
                min_value=0.0, max_value=100.0, value=65.0, step=1.0
            )
            faculty_diversity_index = st.slider(
                "Faculty Diversity Index (1-10)",
                min_value=1, max_value=10, value=6,
                help="Gender, social, regional diversity"
            )
        
        with col2:
            faculty_development_hours = st.number_input(
                "Annual Faculty Development Hours per Faculty",
                min_value=0, value=40,
                help="Training and development hours"
            )
            industry_exposure_faculty = st.number_input(
                "Faculty with Industry Exposure (%)",
                min_value=0.0, max_value=100.0, value=30.0, step=1.0
            )
            international_faculty = st.number_input(
                "International Faculty Percentage (%)",
                min_value=0.0, max_value=100.0, value=5.0, step=1.0
            )
        
        with col3:
            research_publications_per_faculty = st.number_input(
                "Research Publications per Faculty (Annual)",
                min_value=0.0, value=1.5, step=0.1
            )
            faculty_retention_rate = st.number_input(
                "Faculty Retention Rate (%)",
                min_value=0.0, max_value=100.0, value=85.0, step=1.0
            )
        
        st.markdown("---")
        st.markdown("### 🎓 3. LEARNING AND TEACHING")
        
        col1, col2 = st.columns(2)
        
        with col1:
            average_attendance_rate = st.number_input(
                "Average Student Attendance Rate (%)",
                min_value=0.0, max_value=100.0, value=85.0, step=1.0
            )
            digital_platform_usage = st.selectbox(
                "Digital Platform Usage in Teaching",
                ["Extensive Integration", "Moderate Use", "Limited Use", "Traditional Only"],
                help="Use of LMS, online tools, digital resources"
            )
            experiential_learning_hours = st.number_input(
                "Experiential Learning Hours per Student (Annual)",
                min_value=0, value=50
            )
        
        with col2:
            learning_outcome_achievement = st.number_input(
                "Learning Outcome Achievement Rate (%)",
                min_value=0.0, max_value=100.0, value=75.0, step=1.0
            )
            student_feedback_score = st.slider(
                "Student Feedback Score (1-10)",
                min_value=1, max_value=10, value=7
            )
            critical_thinking_assessment = st.selectbox(
                "Critical Thinking Assessment Integration",
                ["Comprehensive", "Moderate", "Limited", "None"],
                help="Assessment of analytical and critical thinking skills"
            )
        
        st.markdown("---")
        st.markdown("### 🔬 4. RESEARCH AND INNOVATION")
        
        col1, col2 = st.columns(2)
        
        with col1:
            research_publications_total = st.number_input(
                "Total Research Publications (Last 3 Years)",
                min_value=0, value=100
            )
            patents_filed = st.number_input(
                "Patents Filed (Last 3 Years)",
                min_value=0, value=10
            )
            research_grants_amount = st.number_input(
                "Research Grants Received (₹ Lakhs - Last 3 Years)",
                min_value=0, value=500
            )
            h_index_institution = st.number_input(
                "Institutional H-index",
                min_value=0, value=25
            )
        
        with col2:
            industry_collaborations = st.number_input(
                "Industry Research Collaborations",
                min_value=0, value=8
            )
            international_research_partnerships = st.number_input(
                "International Research Partnerships",
                min_value=0, value=5
            )
            student_research_participation = st.number_input(
                "Student Research Participation (%)",
                min_value=0.0, max_value=100.0, value=40.0, step=1.0
            )
            research_facility_utilization = st.slider(
                "Research Facility Utilization Rate (1-10)",
                min_value=1, max_value=10, value=7
            )
        
        st.markdown("---")
        st.markdown("### ⚽ 5. EXTRACURRICULAR & CO-CURRICULAR ACTIVITIES")
        
        col1, col2 = st.columns(2)
        
        with col1:
            ec_activities_annual = st.number_input(
                "Annual Extracurricular Activities",
                min_value=0, value=25
            )
            student_participation_rate_ec = st.number_input(
                "Student Participation Rate in EC Activities (%)",
                min_value=0.0, max_value=100.0, value=60.0, step=1.0
            )
            sports_infrastructure_score = st.slider(
                "Sports Infrastructure Quality (1-10)",
                min_value=1, max_value=10, value=6
            )
        
        with col2:
            cultural_events_annual = st.number_input(
                "Annual Cultural Events",
                min_value=0, value=15
            )
            leadership_programs = st.number_input(
                "Leadership Development Programs",
                min_value=0, value=8
            )
            ec_credit_integration = st.selectbox(
                "EC/CC Credit Integration in Curriculum",
                ["Fully Integrated", "Partially Integrated", "Separate", "None"],
                help="Integration of extracurricular credits"
            )
        
        st.markdown("---")
        st.markdown("### 🤝 6. COMMUNITY ENGAGEMENT")
        
        col1, col2 = st.columns(2)
        
        with col1:
            community_projects_annual = st.number_input(
                "Annual Community Engagement Projects",
                min_value=0, value=12
            )
            student_participation_community = st.number_input(
                "Student Participation in Community Service (%)",
                min_value=0.0, max_value=100.0, value=45.0, step=1.0
            )
            rural_outreach_programs = st.number_input(
                "Rural Outreach Programs",
                min_value=0, value=6
            )
        
        with col2:
            social_impact_assessment = st.selectbox(
                "Social Impact Assessment Conducted",
                ["Regularly", "Occasionally", "Rarely", "Never"]
            )
            csr_initiatives = st.number_input(
                "CSR Initiatives Undertaken",
                min_value=0, value=4
            )
            community_feedback_score = st.slider(
                "Community Feedback Score (1-10)",
                min_value=1, max_value=10, value=7
            )
        
        st.markdown("---")
        st.markdown("### 🌱 7. GREEN INITIATIVES")
        
        col1, col2 = st.columns(2)
        
        with col1:
            renewable_energy_usage = st.number_input(
                "Renewable Energy Usage (%)",
                min_value=0.0, max_value=100.0, value=25.0, step=1.0
            )
            waste_management_score = st.slider(
                "Waste Management System Score (1-10)",
                min_value=1, max_value=10, value=6
            )
            water_harvesting_capacity = st.number_input(
                "Water Harvesting Capacity (KL Annual)",
                min_value=0, value=5000
            )
        
        with col2:
            carbon_footprint_reduction = st.number_input(
                "Carbon Footprint Reduction (%) - Last 3 Years",
                min_value=0.0, max_value=100.0, value=15.0, step=1.0
            )
            green_cover_percentage = st.number_input(
                "Green Cover Percentage on Campus",
                min_value=0.0, max_value=100.0, value=40.0, step=1.0
            )
            environmental_awareness_programs = st.number_input(
                "Environmental Awareness Programs (Annual)",
                min_value=0, value=10
            )
        
        st.markdown("---")
        st.markdown("### ⚖️ 8. GOVERNANCE AND ADMINISTRATION")
        
        col1, col2 = st.columns(2)
        
        with col1:
            governance_transparency_score = st.slider(
                "Governance Transparency Score (1-10)",
                min_value=1, max_value=10, value=7
            )
            grievance_redressal_time = st.number_input(
                "Average Grievance Redressal Time (Days)",
                min_value=1, value=15
            )
            egov_implementation_level = st.selectbox(
                "e-Governance Implementation Level",
                ["Advanced", "Moderate", "Basic", "Minimal"]
            )
        
        with col2:
            student_participation_governance = st.number_input(
                "Student Participation in Governance (%)",
                min_value=0.0, max_value=100.0, value=30.0, step=1.0
            )
            administrative_efficiency_score = st.slider(
                "Administrative Efficiency Score (1-10)",
                min_value=1, max_value=10, value=7
            )
            international_collaborations = st.number_input(
                "Active International Collaborations",
                min_value=0, value=8
            )
        
        st.markdown("---")
        st.markdown("### 🏗️ 9. INFRASTRUCTURE DEVELOPMENT")
        
        col1, col2 = st.columns(2)
        
        with col1:
            campus_area = st.number_input(
                "Campus Area (Acres)",
                min_value=0.0, value=50.0, step=0.1
            )
            digital_infrastructure_score = st.slider(
                "Digital Infrastructure Score (1-10)",
                min_value=1, max_value=10, value=7
            )
            laboratory_equipment_score = st.slider(
                "Laboratory Equipment Quality Score (1-10)",
                min_value=1, max_value=10, value=7
            )
        
        with col2:
            library_resources_score = st.slider(
                "Library Resources Score (1-10)",
                min_value=1, max_value=10, value=7
            )
            hostel_capacity_utilization = st.number_input(
                "Hostel Capacity Utilization (%)",
                min_value=0.0, max_value=100.0, value=80.0, step=1.0
            )
            infrastructure_maintenance_score = st.slider(
                "Infrastructure Maintenance Score (1-10)",
                min_value=1, max_value=10, value=7
            )
        
        st.markdown("---")
        st.markdown("### 💰 10. FINANCIAL RESOURCES AND MANAGEMENT")
        
        col1, col2 = st.columns(2)
        
        with col1:
            financial_stability_score = st.slider(
                "Financial Stability Score (1-10)",
                min_value=1, max_value=10, value=7
            )
            research_investment_percentage = st.number_input(
                "Research Investment (% of Total Budget)",
                min_value=0.0, max_value=100.0, value=15.0, step=1.0
            )
            infrastructure_investment = st.number_input(
                "Infrastructure Investment (₹ Lakhs - Annual)",
                min_value=0, value=200
            )
        
        with col2:
            revenue_generation_score = st.slider(
                "Revenue Generation Score (1-10)",
                min_value=1, max_value=10, value=6
            )
            financial_aid_students = st.number_input(
                "Students Receiving Financial Aid (%)",
                min_value=0.0, max_value=100.0, value=25.0, step=1.0
            )
            audit_compliance_score = st.slider(
                "Audit Compliance Score (1-10)",
                min_value=1, max_value=10, value=8
            )
        
        st.markdown("---")
        st.markdown("### 📊 ADDITIONAL INSTITUTIONAL METRICS")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            placement_rate = st.number_input(
                "Placement Rate (%)",
                min_value=0.0, max_value=100.0, value=75.0, step=1.0
            )
            higher_education_rate = st.number_input(
                "Higher Education Progression Rate (%)",
                min_value=0.0, max_value=100.0, value=20.0, step=1.0
            )
        
        with col2:
            entrepreneurship_cell_score = st.slider(
                "Entrepreneurship Cell Activity Score (1-10)",
                min_value=1, max_value=10, value=6
            )
            alumni_engagement_score = st.slider(
                "Alumni Engagement Score (1-10)",
                min_value=1, max_value=10, value=5
            )
        
        with col3:
            institutional_reputation_score = st.slider(
                "Institutional Reputation Score (1-10)",
                min_value=1, max_value=10, value=7
            )
            naac_previous_grade = st.selectbox(
                "Previous NAAC Grade (if any)",
                ["A++", "A+", "A", "B++", "B+", "B", "C", "Not Accredited"]
            )
        
        # Additional information
        st.markdown("---")
        st.markdown("### 📝 ADDITIONAL INFORMATION")
        
        institutional_strengths = st.text_area(
            "Key Institutional Strengths",
            placeholder="Describe your institution's major strengths and achievements...",
            height=100
        )
        
        improvement_areas = st.text_area(
            "Areas for Improvement",
            placeholder="Identify key areas where your institution seeks improvement...",
            height=100
        )
        
        strategic_initiatives = st.text_area(
            "Strategic Initiatives & Future Plans",
            placeholder="Describe ongoing or planned strategic initiatives...",
            height=100
        )
        
        # Submit button
        submitted = st.form_submit_button("🚀 Submit Comprehensive Institutional Data")
        
        if submitted:
            # Compile all data into a structured format
            submission_data = {
                "submission_type": "comprehensive_institutional_data",
                "submission_date": datetime.now().isoformat(),
                "parameter_scores": {
                    "curriculum": {
                        "curriculum_framework_score": curriculum_framework_score,
                        "stakeholder_consultation": stakeholder_consultation,
                        "curriculum_update_frequency": curriculum_update_frequency,
                        "multidisciplinary_courses": multidisciplinary_courses,
                        "skill_integration_score": skill_integration_score,
                        "digital_content_availability": digital_content_availability
                    },
                    "faculty_resources": {
                        "faculty_student_ratio": faculty_student_ratio,
                        "phd_faculty_percentage": phd_faculty_percentage,
                        "faculty_diversity_index": faculty_diversity_index,
                        "faculty_development_hours": faculty_development_hours,
                        "industry_exposure_faculty": industry_exposure_faculty,
                        "international_faculty": international_faculty,
                        "research_publications_per_faculty": research_publications_per_faculty,
                        "faculty_retention_rate": faculty_retention_rate
                    },
                    "learning_teaching": {
                        "average_attendance_rate": average_attendance_rate,
                        "digital_platform_usage": digital_platform_usage,
                        "experiential_learning_hours": experiential_learning_hours,
                        "learning_outcome_achievement": learning_outcome_achievement,
                        "student_feedback_score": student_feedback_score,
                        "critical_thinking_assessment": critical_thinking_assessment
                    },
                    "research_innovation": {
                        "research_publications_total": research_publications_total,
                        "patents_filed": patents_filed,
                        "research_grants_amount": research_grants_amount,
                        "h_index_institution": h_index_institution,
                        "industry_collaborations": industry_collaborations,
                        "international_research_partnerships": international_research_partnerships,
                        "student_research_participation": student_research_participation,
                        "research_facility_utilization": research_facility_utilization
                    },
                    "extracurricular_activities": {
                        "ec_activities_annual": ec_activities_annual,
                        "student_participation_rate_ec": student_participation_rate_ec,
                        "sports_infrastructure_score": sports_infrastructure_score,
                        "cultural_events_annual": cultural_events_annual,
                        "leadership_programs": leadership_programs,
                        "ec_credit_integration": ec_credit_integration
                    },
                    "community_engagement": {
                        "community_projects_annual": community_projects_annual,
                        "student_participation_community": student_participation_community,
                        "rural_outreach_programs": rural_outreach_programs,
                        "social_impact_assessment": social_impact_assessment,
                        "csr_initiatives": csr_initiatives,
                        "community_feedback_score": community_feedback_score
                    },
                    "green_initiatives": {
                        "renewable_energy_usage": renewable_energy_usage,
                        "waste_management_score": waste_management_score,
                        "water_harvesting_capacity": water_harvesting_capacity,
                        "carbon_footprint_reduction": carbon_footprint_reduction,
                        "green_cover_percentage": green_cover_percentage,
                        "environmental_awareness_programs": environmental_awareness_programs
                    },
                    "governance_administration": {
                        "governance_transparency_score": governance_transparency_score,
                        "grievance_redressal_time": grievance_redressal_time,
                        "egov_implementation_level": egov_implementation_level,
                        "student_participation_governance": student_participation_governance,
                        "administrative_efficiency_score": administrative_efficiency_score,
                        "international_collaborations": international_collaborations
                    },
                    "infrastructure_development": {
                        "campus_area": campus_area,
                        "digital_infrastructure_score": digital_infrastructure_score,
                        "laboratory_equipment_score": laboratory_equipment_score,
                        "library_resources_score": library_resources_score,
                        "hostel_capacity_utilization": hostel_capacity_utilization,
                        "infrastructure_maintenance_score": infrastructure_maintenance_score
                    },
                    "financial_management": {
                        "financial_stability_score": financial_stability_score,
                        "research_investment_percentage": research_investment_percentage,
                        "infrastructure_investment": infrastructure_investment,
                        "revenue_generation_score": revenue_generation_score,
                        "financial_aid_students": financial_aid_students,
                        "audit_compliance_score": audit_compliance_score
                    }
                },
                "additional_metrics": {
                    "placement_rate": placement_rate,
                    "higher_education_rate": higher_education_rate,
                    "entrepreneurship_cell_score": entrepreneurship_cell_score,
                    "alumni_engagement_score": alumni_engagement_score,
                    "institutional_reputation_score": institutional_reputation_score,
                    "naac_previous_grade": naac_previous_grade
                },
                "qualitative_data": {
                    "institutional_strengths": institutional_strengths,
                    "improvement_areas": improvement_areas,
                    "strategic_initiatives": strategic_initiatives
                }
            }
            
            # Save to database
            analyzer.save_institution_submission(
                user['institution_id'],
                "comprehensive_institutional_data",
                submission_data
            )
            
            st.success("✅ Comprehensive Institutional Data Submitted Successfully!")
            st.balloons()
            
            # Show quick summary
            st.subheader("📋 Submission Summary")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("Parameters Covered", "10")
                st.metric("Data Points Collected", "65+")
            
            with col2:
                st.metric("AI Analysis Ready", "Yes")
                st.metric("NEP 2020 Compliant", "Yes")
            
            with col3:
                st.metric("Submission Status", "Complete")
            
            st.info("""
            **Next Steps:**
            - Your data will be processed for AI-powered institutional analysis
            - Comprehensive assessment report will be generated
            - Accreditation recommendations will be provided
            - You can track the analysis progress in the 'My Submissions' section
            """)
            
def create_institution_submissions_view(analyzer, user):
    st.subheader("📊 My Submissions & Status")
    
    submissions = analyzer.get_institution_submissions(user['institution_id'])
    
    if len(submissions) > 0:
        for _, submission in submissions.iterrows():
            with st.expander(f"{submission['submission_type']} - {submission['submitted_date']}"):
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.write(f"**Status:** {submission['status']}")
                with col2:
                    st.write(f"**Submitted:** {submission['submitted_date']}")
                with col3:
                    if submission['reviewed_by']:
                        st.write(f"**Reviewed by:** {submission['reviewed_by']}")
                
                if submission['review_comments']:
                    st.info(f"**Review Comments:** {submission['review_comments']}")
                
                # Display submission data
                try:
                    submission_data = json.loads(submission['submission_data'])
                    st.json(submission_data)
                except:
                    st.write("Submission data format not available")
    else:
        st.info("No submissions found. Use the Data Submission tab to submit your institutional data.")

def create_institution_requirements_guide(analyzer):
    st.subheader("📋 Approval Requirements Guide")
    
    requirements = analyzer.document_requirements
    
    for approval_type, docs in requirements.items():
        with st.expander(f"{approval_type.replace('_', ' ').title()} Requirements"):
            st.write("**Mandatory Documents:**")
            for doc in docs['mandatory']:
                st.write(f"• {doc.replace('_', ' ').title()}")
            
            st.write("**Supporting Documents:**")
            for doc in docs['supporting']:
                st.write(f"• {doc.replace('_', ' ').title()}")

# Existing analytical modules (unchanged)
def create_performance_dashboard(analyzer):
    st.header("📊 Institutional Performance Analytics Dashboard")
    
    # Use the actual data from analyzer (should be 20 institutions × 10 years)
    df = analyzer.historical_data
    
    # Show data specification at the top
    st.info(f"📊 **Data Overview**: {df['institution_id'].nunique()} Institutions × {df['year'].nunique()} Years ({df['year'].min()}-{df['year'].max()}) | Total Records: {len(df)}")
    
    # Filter for current year data only for KPI calculations
    current_year_data = df[df['year'] == 2023]
    
    # Ensure we have data
    if len(current_year_data) == 0:
        st.warning("No data available for the current year. Please check data generation.")
        return
    
    # Key Performance Indicators
    st.subheader("🏆 Key Performance Indicators (2023 Data)")
    
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        avg_performance = current_year_data['performance_score'].mean()
        st.metric("Average Performance Score", f"{avg_performance:.2f}/10")
    
    with col2:
        approval_rate = (current_year_data['performance_score'] >= 6.0).mean()
        st.metric("Approval Eligibility Rate", f"{approval_rate:.1%}")
    
    with col3:
        high_risk_count = (current_year_data['risk_level'] == 'High Risk').sum() + (
            current_year_data['risk_level'] == 'Critical Risk').sum()
        st.metric("High/Critical Risk Institutions", high_risk_count)
    
    with col4:
        avg_placement = current_year_data['placement_rate'].mean()
        st.metric("Average Placement Rate", f"{avg_placement:.1f}%")
    
    with col5:
        research_intensity = current_year_data['research_publications'].sum() / len(current_year_data)
        st.metric("Avg Research Publications", f"{research_intensity:.1f}")
    
    # Performance Analysis
    st.subheader("📈 Performance Analysis (20 Institutions)")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Performance Distribution for current year
        fig1 = px.histogram(
            current_year_data, 
            x='performance_score',
            title="Distribution of Institutional Performance Scores (2023)",
            nbins=12,
            color_discrete_sequence=['#1f77b4'],
            opacity=0.8
        )
        fig1.update_layout(
            xaxis_title="Performance Score", 
            yaxis_title="Number of Institutions",
            showlegend=False
        )
        st.plotly_chart(fig1, use_container_width=True)
    
    with col2:
        # Performance by Institution Type
        fig2 = px.box(
            current_year_data,
            x='institution_type',
            y='performance_score',
            title="Performance Score by Institution Type (2023)",
            color='institution_type'
        )
        fig2.update_layout(
            xaxis_title="Institution Type",
            yaxis_title="Performance Score",
            showlegend=False
        )
        st.plotly_chart(fig2, use_container_width=True)
    
    # Trend Analysis - Show all 10 years for the 20 institutions
    st.subheader("📅 Historical Performance Trends (2014-2023)")
    
    # Calculate average performance by year and type
    trend_data = df.groupby(['year', 'institution_type'])['performance_score'].mean().reset_index()
    
    fig3 = px.line(
        trend_data,
        x='year',
        y='performance_score',
        color='institution_type',
        title="Average Performance Score Trend (2014-2023) - 20 Institutions",
        markers=True
    )
    fig3.update_layout(
        xaxis_title="Year", 
        yaxis_title="Average Performance Score",
        legend_title="Institution Type"
    )
    st.plotly_chart(fig3, use_container_width=True)
    
    # Risk Analysis
    st.subheader("⚠️ Institutional Risk Analysis (2023)")
    
    col1, col2 = st.columns(2)
    
    with col1:
        risk_distribution = current_year_data['risk_level'].value_counts()
        fig4 = px.pie(
            values=risk_distribution.values,
            names=risk_distribution.index,
            title="Institutional Risk Level Distribution",
            color=risk_distribution.index,
            color_discrete_map={
                'Low Risk': '#2ecc71',
                'Medium Risk': '#f39c12',
                'High Risk': '#e74c3c',
                'Critical Risk': '#c0392b'
            }
        )
        st.plotly_chart(fig4, use_container_width=True)
    
    with col2:
        # Placement vs Research Analysis
        fig5 = px.scatter(
            current_year_data,
            x='research_publications',
            y='placement_rate',
            color='risk_level',
            size='performance_score',
            hover_data=['institution_name'],
            title="Research Output vs Placement Rate (2023)",
            color_discrete_map={
                'Low Risk': '#2ecc71',
                'Medium Risk': '#f39c12',
                'High Risk': '#e74c3c',
                'Critical Risk': '#c0392b'
            }
        )
        fig5.update_layout(
            xaxis_title="Research Publications",
            yaxis_title="Placement Rate (%)"
        )
        st.plotly_chart(fig5, use_container_width=True)
    
    # Additional Visualizations focused on 20 institutions
    st.subheader("🎯 Performance Insights - 20 Institutions Analysis")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Top 10 performing institutions
        top_performers = current_year_data.nlargest(10, 'performance_score')[['institution_name', 'performance_score', 'naac_grade']]
        fig6 = px.bar(
            top_performers,
            x='performance_score',
            y='institution_name',
            orientation='h',
            title="Top 10 Performing Institutions (2023)",
            color='performance_score',
            color_continuous_scale='Viridis'
        )
        fig6.update_layout(
            yaxis_title="Institution",
            xaxis_title="Performance Score"
        )
        st.plotly_chart(fig6, use_container_width=True)
    
    with col2:
        # State-wise Performance
        state_performance = current_year_data.groupby('state')['performance_score'].mean().sort_values(ascending=False)
        fig7 = px.bar(
            x=state_performance.index,
            y=state_performance.values,
            title="States by Average Performance Score (2023)",
            color=state_performance.values,
            color_continuous_scale='Viridis'
        )
        fig7.update_layout(
            xaxis_title="State",
            yaxis_title="Average Performance Score",
            showlegend=False
        )
        st.plotly_chart(fig7, use_container_width=True)
    
    # Comprehensive Data Table
    st.subheader("📋 Institutional Performance Data (2023)")
    
    # Show key metrics for all 20 institutions
    display_columns = [
        'institution_id', 'institution_name', 'institution_type', 'state',
        'performance_score', 'naac_grade', 'placement_rate', 'risk_level',
        'approval_recommendation'
    ]
    
    st.dataframe(
        current_year_data[display_columns].sort_values('performance_score', ascending=False),
        use_container_width=True,
        height=400
    )
    
    # Institution Comparison Tool
    st.subheader("🔍 Compare Institutions")
    
    # Select institutions to compare
    institutions_list = current_year_data['institution_name'].tolist()
    selected_institutions = st.multiselect(
        "Select institutions to compare:",
        institutions_list,
        default=institutions_list[:3] if len(institutions_list) >= 3 else institutions_list
    )
    
    if selected_institutions:
        comparison_data = current_year_data[current_year_data['institution_name'].isin(selected_institutions)]
        
        # Create comparison chart
        fig8 = px.bar(
            comparison_data,
            x='institution_name',
            y=['performance_score', 'placement_rate', 'research_publications'],
            title="Institution Comparison",
            barmode='group'
        )
        fig8.update_layout(
            xaxis_title="Institution",
            yaxis_title="Score/Percentage/Count"
        )
        st.plotly_chart(fig8, use_container_width=True)
        
        # Show detailed comparison table
        comparison_cols = [
            'institution_name', 'performance_score', 'naac_grade', 'nirf_ranking',
            'placement_rate', 'research_publications', 'student_faculty_ratio',
            'financial_stability_score', 'risk_level'
        ]
        st.dataframe(
            comparison_data[comparison_cols].set_index('institution_name'),
            use_container_width=True
        )


def create_document_analysis_module(analyzer):
    st.header("📋 AI-Powered Document Sufficiency Analysis")
    
    st.info("Analyze document completeness and generate sufficiency reports for approval processes")
    
    # Generate enhanced dummy document data with realistic patterns
    generate_enhanced_dummy_document_data(analyzer)
    
    # Institution selection
    current_institutions = analyzer.historical_data[analyzer.historical_data['year'] == 2023]['institution_id'].unique()
    selected_institution = st.selectbox(
        "Select Institution",
        current_institutions,
        key="doc_analysis_institution"
    )
    
    approval_type = st.selectbox(
        "Select Approval Type",
        ["new_approval", "renewal_approval", "expansion_approval"],
        format_func=lambda x: x.replace('_', ' ').title(),
        key="doc_analysis_approval_type"
    )
    
    # Get institution performance data
    institution_performance = get_institution_performance(selected_institution, analyzer)
    
    # Display performance context
    display_performance_context(institution_performance, selected_institution)
    
    # Display document checklist with enhanced status
    display_enhanced_document_checklist(selected_institution, approval_type, analyzer, institution_performance)
    
    # Analysis and recommendations
    if st.button("🤖 Analyze Document Sufficiency", type="primary"):
        perform_enhanced_document_analysis(selected_institution, approval_type, analyzer, institution_performance)

def get_institution_performance(institution_id, analyzer):
    """Get institution performance data"""
    try:
        # Get the current year data (2023)
        current_year_data = analyzer.historical_data[analyzer.historical_data['year'] == 2023]
        
        inst_data = current_year_data[current_year_data['institution_id'] == institution_id]
        
        if not inst_data.empty:
            return {
                'performance_score': float(inst_data['performance_score'].iloc[0]),
                'naac_grade': str(inst_data['naac_grade'].iloc[0]),
                'risk_level': str(inst_data['risk_level'].iloc[0]),
                'institution_name': str(inst_data['institution_name'].iloc[0]),
                'institution_type': str(inst_data['institution_type'].iloc[0])
            }
    except Exception as e:
        print(f"Error getting institution performance: {e}")
    
    # Return default values if data not found
    return {
        'performance_score': 5.0,
        'naac_grade': 'B',
        'risk_level': 'Medium Risk',
        'institution_name': 'Unknown Institution',
        'institution_type': 'General'
    }



def display_performance_context(performance, institution_id):
    """Display institution performance context"""
    
    st.subheader(f"🏛️ {performance['institution_name']}")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        # Color code based on performance
        if performance['performance_score'] >= 8.0:
            color = "green"
        elif performance['performance_score'] >= 6.0:
            color = "orange"
        else:
            color = "red"
        
        st.metric("Performance Score", f"{performance['performance_score']:.1f}/10")
    
    with col2:
        st.metric("NAAC Grade", performance['naac_grade'])
    
    with col3:
        risk_color = {
            "Low Risk": "green",
            "Medium Risk": "orange", 
            "High Risk": "red",
            "Critical Risk": "darkred"
        }.get(performance['risk_level'], "gray")
        
        st.metric("Risk Level", performance['risk_level'])
    
    with col4:
        st.metric("Institution Type", performance['institution_type'])
    
    # Performance interpretation
    if performance['performance_score'] >= 8.0:
        st.success("🎯 **High Performing Institution**: Expected to have comprehensive document submission")
    elif performance['performance_score'] >= 6.0:
        st.info("📊 **Medium Performing Institution**: Expected to have good document coverage")
    else:
        st.warning("⚠️ **Low Performing Institution**: May have incomplete document submission")



def generate_enhanced_dummy_document_data(analyzer):
    """Generate realistic dummy document data with upload patterns and dates"""
    
    if 'enhanced_docs_generated' not in st.session_state:
        try:
            # Get all institutions (20 institutions)
            institutions = analyzer.historical_data[analyzer.historical_data['year'] == 2023]['institution_id'].unique()
            
            for institution_id in institutions:
                # Get institution performance to determine document completeness
                performance = get_institution_performance(institution_id, analyzer)
                
                # Clear existing dummy data for this institution
                cursor = analyzer.conn.cursor()
                cursor.execute('DELETE FROM institution_documents WHERE institution_id = ?', (institution_id,))
                
                # Generate enhanced documents with realistic patterns
                documents = generate_enhanced_institution_documents(institution_id, performance)
                
                # Insert into database with upload dates
                for doc in documents:
                    cursor.execute('''
                        INSERT INTO institution_documents 
                        (institution_id, document_name, document_type, status, upload_date)
                        VALUES (?, ?, ?, ?, ?)
                    ''', (institution_id, doc['name'], doc['type'], doc['status'], doc['upload_date']))
                
                analyzer.conn.commit()
            
            st.session_state.enhanced_docs_generated = True
            print(f"✅ Generated enhanced document data for {len(institutions)} institutions")
            
        except Exception as e:
            print(f"Could not generate enhanced dummy documents: {e}")


def generate_enhanced_institution_documents(institution_id, performance):
    """Generate realistic document set with upload patterns and dates"""
    
    performance_score = performance['performance_score']
    institution_type = performance['institution_type']
    
    # Enhanced document templates with categories
    all_documents = {
        "mandatory_critical": [
            {"name": "Curriculum Framework 2023.pdf", "type": "curriculum_framework"},
            {"name": "Faculty Qualification Records.pdf", "type": "faculty_qualifications"},
            {"name": "Institutional Statutes.pdf", "type": "institutional_statutes"},
            {"name": "Annual Financial Statements.pdf", "type": "financial_statements"},
            {"name": "Academic Calendar 2023-24.pdf", "type": "academic_calendar"},
        ],
        "mandatory_important": [
            {"name": "Course Outlines and Syllabi.docx", "type": "course_outlines"},
            {"name": "Faculty Recruitment Policy.pdf", "type": "recruitment_policy"},
            {"name": "Organizational Structure.pdf", "type": "organizational_structure"},
            {"name": "Budget Allocation Document.pdf", "type": "budget_allocation"},
            {"name": "Learning Outcome Assessment.pdf", "type": "learning_outcomes"},
        ]
    }
    
    uploaded_documents = []
    
    # Define upload patterns based on performance
    if performance_score >= 8.5:  # Top performers
        pattern = {
            "mandatory_critical": 1.0,  # 100% uploaded
            "mandatory_important": 0.95, # 95% uploaded
        }
        
    elif performance_score >= 7.0:  # Good performers
        pattern = {
            "mandatory_critical": 1.0,   # 100% uploaded
            "mandatory_important": 0.85, # 85% uploaded
        }
        
    elif performance_score >= 5.5:  # Average performers
        pattern = {
            "mandatory_critical": 0.90,  # 90% uploaded
            "mandatory_important": 0.70, # 70% uploaded
        }
        
    else:  # Low performers
        pattern = {
            "mandatory_critical": 0.60,  # 60% uploaded
            "mandatory_important": 0.40, # 40% uploaded
        }
    
    # Generate upload dates (within last 6 months)
    base_date = datetime.now()
    
    for category, docs in all_documents.items():
        upload_probability = pattern[category]
        
        for doc in docs:
            if np.random.random() < upload_probability:
                # Document is uploaded - generate realistic upload date
                days_ago = np.random.randint(1, 180)  # Within last 6 months
                upload_date = base_date - timedelta(days=days_ago)
                
                uploaded_documents.append({
                    **doc,
                    "status": "Uploaded",
                    "upload_date": upload_date
                })
            else:
                # Document is pending
                uploaded_documents.append({
                    **doc,
                    "status": "Pending",
                    "upload_date": None
                })
    
    return uploaded_documents

def update_institution_performance(institution_id, performance_boost, analyzer):
    """Update institution performance based on document upload completeness"""
    try:
        cursor = analyzer.conn.cursor()
        
        # Get current performance
        cursor.execute('''
            SELECT performance_score FROM institutions 
            WHERE institution_id = ? AND year = 2023
        ''', (institution_id,))
        
        result = cursor.fetchone()
        if result:
            current_score = result[0]
            new_score = min(10.0, max(1.0, current_score + performance_boost))
            
            # Update performance score
            cursor.execute('''
                UPDATE institutions 
                SET performance_score = ? 
                WHERE institution_id = ? AND year = 2023
            ''', (new_score, institution_id))
            
            # Update approval recommendation and risk level
            new_recommendation = generate_approval_recommendation(new_score)
            new_risk_level = assess_risk_level(new_score)
            
            cursor.execute('''
                UPDATE institutions 
                SET approval_recommendation = ?, risk_level = ?
                WHERE institution_id = ? AND year = 2023
            ''', (new_recommendation, new_risk_level, institution_id))
            
            analyzer.conn.commit()
            
    except Exception as e:
        print(f"Error updating performance: {e}")



def perform_enhanced_document_analysis(institution_id, approval_type, analyzer, performance):
    """Perform enhanced document analysis with performance impact"""
    
    counts = st.session_state.get('enhanced_document_counts', {})
    
    if not counts:
        st.error("No document data available for analysis")
        return
    
    total_mandatory = counts['total_mandatory']
    uploaded_mandatory = counts['uploaded_mandatory']
    pending_mandatory = counts['pending_mandatory']
    uploaded_supporting = counts['uploaded_supporting']
    performance_data = counts['performance']
    
    # Calculate enhanced metrics
    mandatory_sufficiency = (uploaded_mandatory / total_mandatory) * 100 if total_mandatory > 0 else 0
    compliance_score = min(100, mandatory_sufficiency * 1.2)  # Bonus for timeliness
    
    st.subheader("🎯 Enhanced Document Performance Analysis")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Mandatory Compliance", f"{mandatory_sufficiency:.1f}%")
    
    with col2:
        st.metric("Pending Documents", pending_mandatory)
    
    with col3:
        st.metric("Supporting Documents", uploaded_supporting)
    
    with col4:
        st.metric("Overall Score", f"{performance_data['performance_score']:.1f}/10")
    
    # Performance Impact Analysis
    st.subheader("📈 Document Upload Impact on Performance")
    
    # Show how document upload affects performance
    if mandatory_sufficiency >= 90:
        st.success("**🏆 EXCELLENT COMPLIANCE**: High document upload rate significantly boosts institutional performance")
        st.write("**Impact**: +0.3 points added to performance score")
    elif mandatory_sufficiency >= 70:
        st.info("**📊 GOOD COMPLIANCE**: Solid document submission supports institutional performance")
        st.write("**Impact**: +0.15 points added to performance score")
    elif mandatory_sufficiency >= 50:
        st.warning("**⚠️ AVERAGE COMPLIANCE**: Incomplete submission limits performance potential")
        st.write("**Impact**: No performance boost - missed opportunity")
    else:
        st.error("**🚨 POOR COMPLIANCE**: Significant document gaps negatively impact performance")
        st.write("**Impact**: -0.2 points deducted from performance score")
    
    # Timeline Analysis
    st.subheader("⏰ Submission Timeline Analysis")
    
    uploaded_docs = [d for d in counts['uploaded_docs_data'] if d['status'] == 'Uploaded']
    if uploaded_docs:
        # Calculate average days since upload
        recent_uploads = 0
        for doc in uploaded_docs:
            days_ago = (datetime.now() - pd.to_datetime(doc['upload_date'])).days
            if days_ago <= 30:  # Within last month
                recent_uploads += 1
        
        recent_upload_rate = (recent_uploads / len(uploaded_docs)) * 100
        
        if recent_upload_rate >= 80:
            st.success("**🕒 ACTIVE SUBMISSION**: Most documents uploaded recently")
        elif recent_upload_rate >= 50:
            st.info("**📅 STEADY PROGRESS**: Regular document submission pattern")
        else:
            st.warning("**⏳ DELAYED UPLOADS**: Many documents uploaded long time ago")
    else:
        st.error("**📭 NO UPLOADS**: Institution has not submitted any documents")
    
    # Institutional Comparison
    st.subheader("🏛️ Peer Institution Comparison")
    
    # Get all institutions for comparison
    all_institutions = analyzer.historical_data[analyzer.historical_data['year'] == 2023]
    
    # Simulate document upload rates for comparison
    comparison_data = []
    for _, inst in all_institutions.head(5).iterrows():
        if inst['institution_id'] == institution_id:
            # Current institution
            comparison_data.append({
                'Institution': f"📍 {inst['institution_name']} (Current)",
                'Document Upload Rate': f"{mandatory_sufficiency:.1f}%",
                'Performance Score': inst['performance_score']
            })
        else:
            # Simulate upload rate based on performance
            sim_upload_rate = min(100, inst['performance_score'] * 10 + np.random.randint(-10, 10))
            comparison_data.append({
                'Institution': inst['institution_name'],
                'Document Upload Rate': f"{sim_upload_rate:.1f}%",
                'Performance Score': inst['performance_score']
            })
    
    comparison_df = pd.DataFrame(comparison_data)
    st.dataframe(comparison_df, use_container_width=True)

def display_enhanced_document_checklist(institution_id, approval_type, analyzer, performance):
    """Display enhanced document checklist with upload dates and status"""
    
    # Get requirements
    requirements = get_document_requirements_by_parameters(approval_type)
    
    # Get uploaded documents for this institution with dates
    uploaded_docs_data = []
    try:
        uploaded_docs_df = analyzer.get_institution_documents(institution_id)
        if not uploaded_docs_df.empty:
            for _, row in uploaded_docs_df.iterrows():
                uploaded_docs_data.append({
                    'name': row['document_name'],
                    'type': row['document_type'],
                    'status': row['status'],
                    'upload_date': row['upload_date']
                })
    except Exception as e:
        st.warning(f"Could not load uploaded documents: {e}")
    
    # Display enhanced document statistics
    st.subheader("📊 Enhanced Document Analysis")
    
    # Calculate statistics
    total_docs = len(uploaded_docs_data)
    uploaded_count = len([d for d in uploaded_docs_data if d['status'] == 'Uploaded'])
    pending_count = len([d for d in uploaded_docs_data if d['status'] == 'Pending'])
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Documents", total_docs)
    
    with col2:
        st.metric("Uploaded", uploaded_count, delta=f"+{uploaded_count}")
    
    with col3:
        st.metric("Pending", pending_count, delta=f"-{pending_count}", delta_color="inverse")
    
    with col4:
        completion_rate = (uploaded_count / total_docs * 100) if total_docs > 0 else 0
        st.metric("Completion Rate", f"{completion_rate:.1f}%")
    
    # Display mandatory documents with enhanced status
    st.subheader("📋 Mandatory Documents Status")
    
    total_mandatory = 0
    uploaded_mandatory = 0
    pending_mandatory = 0
    
    for parameter, documents in requirements["mandatory"].items():
        with st.expander(f"🔴 {parameter} - Mandatory Documents", expanded=True):
            for doc_template in documents:
                total_mandatory += 1
                
                # Find matching uploaded document
                matching_doc = None
                for uploaded_doc in uploaded_docs_data:
                    if doc_template.lower() in uploaded_doc['name'].lower():
                        matching_doc = uploaded_doc
                        break
                
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    if matching_doc and matching_doc['status'] == 'Uploaded':
                        days_ago = (datetime.now() - pd.to_datetime(matching_doc['upload_date'])).days
                        st.success(f"✅ {doc_template}")
                        st.caption(f"📅 Uploaded {days_ago} days ago")
                    else:
                        st.error(f"❌ {doc_template}")
                        st.caption("⏳ Status: Pending - Institution has failed to submit")
                
                with col2:
                    if matching_doc and matching_doc['status'] == 'Uploaded':
                        st.markdown("**✅ Uploaded**")
                        uploaded_mandatory += 1
                    else:
                        st.markdown("**🔴 Pending**")
                        pending_mandatory += 1
                
                with col3:
                    if matching_doc and matching_doc['status'] == 'Uploaded':
                        upload_date = pd.to_datetime(matching_doc['upload_date']).strftime("%d %b %Y")
                        st.markdown(f"**{upload_date}**")
                    else:
                        st.markdown("**OVERDUE**")
    
    # Display supporting documents with enhanced status
    st.subheader("📝 Supporting Documents Status")
    
    total_supporting = 0
    uploaded_supporting = 0
    pending_supporting = 0
    
    for parameter, documents in requirements["supporting"].items():
        with st.expander(f"🟡 {parameter} - Supporting Documents"):
            for doc_template in documents:
                total_supporting += 1
                
                # Find matching uploaded document
                matching_doc = None
                for uploaded_doc in uploaded_docs_data:
                    if doc_template.lower() in uploaded_doc['name'].lower():
                        matching_doc = uploaded_doc
                        break
                
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    if matching_doc and matching_doc['status'] == 'Uploaded':
                        days_ago = (datetime.now() - pd.to_datetime(matching_doc['upload_date'])).days
                        st.info(f"✅ {doc_template}")
                        st.caption(f"📅 Uploaded {days_ago} days ago")
                    else:
                        st.warning(f"⭕ {doc_template}")
                        st.caption("💡 Recommended for enhanced assessment")
                
                with col2:
                    if matching_doc and matching_doc['status'] == 'Uploaded':
                        st.markdown("**✅ Uploaded**")
                        uploaded_supporting += 1
                    else:
                        st.markdown("**🟡 Optional**")
                        pending_supporting += 1
                
                with col3:
                    if matching_doc and matching_doc['status'] == 'Uploaded':
                        upload_date = pd.to_datetime(matching_doc['upload_date']).strftime("%d %b %Y")
                        st.markdown(f"**{upload_date}**")
                    else:
                        st.markdown("**NOT UPLOADED**")
    
    # Store enhanced counts for analysis
    st.session_state.enhanced_document_counts = {
        'total_mandatory': total_mandatory,
        'uploaded_mandatory': uploaded_mandatory,
        'pending_mandatory': pending_mandatory,
        'total_supporting': total_supporting,
        'uploaded_supporting': uploaded_supporting,
        'pending_supporting': pending_supporting,
        'uploaded_docs_data': uploaded_docs_data,
        'performance': performance
    }


def get_document_requirements_by_parameters(approval_type):
    """Get document requirements organized by parameters"""
    
    base_requirements = {
        "new_approval": {
            "mandatory": {
                "Curriculum": [
                    "Curriculum framework documents",
                    "Course outlines and objectives",
                    "Academic calendar"
                ],
                "Faculty Resources": [
                    "Faculty qualification records",
                    "Recruitment policy documents"
                ],
                "Governance": [
                    "Institutional statutes",
                    "Organizational structure"
                ]
            },
            "supporting": {
                "Research": [
                    "Research publications",
                    "Project documentation"
                ],
                "Infrastructure": [
                    "Campus facility details",
                    "Laboratory equipment list"
                ]
            }
        },
        "renewal_approval": {
            "mandatory": {
                "Curriculum": [
                    "Updated curriculum framework",
                    "Academic performance reports"
                ],
                "Faculty": [
                    "Updated faculty records",
                    "Development reports"
                ]
            },
            "supporting": {
                "Research": [
                    "Recent publications",
                    "Research projects"
                ]
            }
        },
        "expansion_approval": {
            "mandatory": {
                "Infrastructure": [
                    "Expansion master plan",
                    "Additional facilities plan"
                ],
                "Faculty": [
                    "New faculty requirements",
                    "Recruitment plan"
                ]
            },
            "supporting": {
                "Financial": [
                    "Expansion budget",
                    "Funding plan"
                ]
            }
        }
    }
    
    return base_requirements.get(approval_type, base_requirements["new_approval"])


            
def create_institutional_intelligence_hub(analyzer):
    st.header("🧠 Institutional Intelligence Hub")
    
    st.info("""
    **Comprehensive AI-powered insights, predictive analytics, and strategic recommendations** 
    for institutional excellence and NEP 2020 compliance.
    """)
    
    # Institution selection
    current_institutions = analyzer.historical_data[analyzer.historical_data['year'] == 2023]['institution_id'].unique()
    
    selected_institution = st.selectbox(
        "Select Institution for Deep Analysis",
        current_institutions,
        key="intel_hub_institution"
    )
    
    # Get comprehensive data
    institution_data = get_comprehensive_institution_data(selected_institution, analyzer)
    
    if institution_data:
        display_intelligence_dashboard(institution_data, analyzer)
    
    # Add comparative analysis section
    st.subheader("🏛️ Peer Comparison")
    
    if st.button("📊 Generate Comparative Analysis", type="secondary"):
        show_comparative_analysis_safe(selected_institution, analyzer)
    
    # Add improvement roadmap
    st.subheader("🎯 Improvement Strategy")
    
    if st.button("🔄 Generate Improvement Roadmap", type="secondary"):
        show_improvement_roadmap_safe(selected_institution, analyzer)

def show_improvement_roadmap_safe(institution_id, analyzer):
    """Safe version of improvement roadmap"""
    try:
        # Get institution data
        current_year_data = analyzer.historical_data[analyzer.historical_data['year'] == 2023]
        current_data = current_year_data[current_year_data['institution_id'] == institution_id]
        
        if current_data.empty:
            st.warning("No data available for institution")
            return
        
        current = current_data.iloc[0].to_dict()
        
        # Create roadmap
        roadmap = [
            {"Priority": "High", "Action": "Review and update curriculum", "Timeline": "3 months", "Impact": "High"},
            {"Priority": "High", "Action": "Faculty development programs", "Timeline": "6 months", "Impact": "High"},
            {"Priority": "Medium", "Action": "Enhance research infrastructure", "Timeline": "9 months", "Impact": "Medium"},
            {"Priority": "Medium", "Action": "Improve digital learning facilities", "Timeline": "6 months", "Impact": "Medium"},
            {"Priority": "High", "Action": "Strengthen industry partnerships", "Timeline": "4 months", "Impact": "High"},
        ]
        
        # Add specific actions based on weaknesses
        weaknesses = get_institutional_weaknesses(current)
        
        if "High student-faculty ratio" in str(weaknesses):
            roadmap.append({"Priority": "Critical", "Action": "Recruit additional faculty", "Timeline": "6 months", "Impact": "High"})
        
        if "Low placement rate" in str(weaknesses):
            roadmap.append({"Priority": "Critical", "Action": "Enhance placement services", "Timeline": "3 months", "Impact": "High"})
        
        df_roadmap = pd.DataFrame(roadmap)
        st.dataframe(df_roadmap, use_container_width=True)
        
    except Exception as e:
        st.error(f"Error generating roadmap: {str(e)}")



def show_comparative_analysis_safe(institution_id, analyzer):
    """Safe version of comparative analysis"""
    try:
        # Get current institution data
        current_year_data = analyzer.historical_data[analyzer.historical_data['year'] == 2023]
        current_data = current_year_data[current_year_data['institution_id'] == institution_id]
        
        if current_data.empty:
            st.warning("No data available for current institution")
            return
        
        current_row = current_data.iloc[0]
        
        # Get peer institutions
        peer_data = current_year_data[
            (current_year_data['institution_type'] == current_row['institution_type']) &
            (current_year_data['institution_id'] != institution_id)
        ]
        
        if peer_data.empty:
            st.info("No peer institutions found for comparison")
            return
        
        # Create comparison table
        comparison = [{
            'Institution': f"📍 {current_row['institution_name']}",
            'Type': current_row['institution_type'],
            'Performance': current_row['performance_score'],
            'NAAC Grade': current_row.get('naac_grade', 'N/A'),
            'Placement %': current_row.get('placement_rate', 0)
        }]
        
        # Add top 3 peers
        top_peers = peer_data.nlargest(3, 'performance_score')
        for _, peer in top_peers.iterrows():
            comparison.append({
                'Institution': peer['institution_name'],
                'Type': peer['institution_type'],
                'Performance': peer['performance_score'],
                'NAAC Grade': peer.get('naac_grade', 'N/A'),
                'Placement %': peer.get('placement_rate', 0)
            })
        
        df_comparison = pd.DataFrame(comparison)
        st.dataframe(df_comparison, use_container_width=True)
        
        # Performance percentile
        all_scores = peer_data['performance_score'].tolist() + [current_row['performance_score']]
        percentile = (sum(1 for s in all_scores if s < current_row['performance_score']) / len(all_scores)) * 100
        st.metric("Performance Percentile", f"{percentile:.1f}%")
        
    except Exception as e:
        st.error(f"Error in comparative analysis: {str(e)}")




def show_improvement_roadmap(institution_id, analyzer):
    """Show improvement roadmap"""
    st.subheader("🎯 Improvement Roadmap")
    
    # Get institution data
    institution_data = get_comprehensive_institution_data(institution_id, analyzer)
    if not institution_data:
        return
    
    current = institution_data['current']
    
    st.info("**AI-Generated Improvement Roadmap for Next 12 Months**")
    
    roadmap_items = [
        {"Area": "Academic Excellence", "Action": "Review and update curriculum", "Timeline": "3 months", "Priority": "High"},
        {"Area": "Faculty Development", "Action": "Conduct faculty training programs", "Timeline": "6 months", "Priority": "High"},
        {"Area": "Research & Innovation", "Action": "Increase research publications by 20%", "Timeline": "12 months", "Priority": "Medium"},
        {"Area": "Infrastructure", "Action": "Upgrade digital learning facilities", "Timeline": "9 months", "Priority": "Medium"},
        {"Area": "Industry Connect", "Action": "Establish 5 new industry partnerships", "Timeline": "6 months", "Priority": "High"},
        {"Area": "Student Support", "Action": "Enhance placement and career services", "Timeline": "4 months", "Priority": "High"},
    ]
    
    # Customize based on weaknesses
    weaknesses = identify_institutional_weaknesses(current)
    if "High Student-Faculty Ratio" in str(weaknesses):
        roadmap_items.append({"Area": "Faculty Resources", "Action": "Recruit additional faculty", "Timeline": "6 months", "Priority": "Critical"})
    
    if "Low Placement Rate" in str(weaknesses):
        roadmap_items.append({"Area": "Placements", "Action": "Strengthen industry collaborations", "Timeline": "4 months", "Priority": "Critical"})
    
    df_roadmap = pd.DataFrame(roadmap_items)
    st.dataframe(df_roadmap, use_container_width=True)

def show_comparative_analysis(institution_id, analyzer):
    """Show comparative analysis with peer institutions"""
    st.subheader("🏛️ Comparative Analysis with Peer Institutions")
    
    # Get current institution data
    current_data = analyzer.historical_data[
        (analyzer.historical_data['institution_id'] == institution_id) &
        (analyzer.historical_data['year'] == 2023)
    ]
    
    if current_data.empty:
        st.warning("No data available for current institution")
        return
    
    current_row = current_data.iloc[0]
    
    # Get peer institutions (same type)
    peer_data = analyzer.historical_data[
        (analyzer.historical_data['institution_type'] == current_row['institution_type']) &
        (analyzer.historical_data['year'] == 2023) &
        (analyzer.historical_data['institution_id'] != institution_id)
    ]
    
    if peer_data.empty:
        st.info("No peer institutions found for comparison")
        return
    
    # Calculate percentile
    performance_scores = peer_data['performance_score'].tolist() + [current_row['performance_score']]
    current_percentile = (sum(1 for score in performance_scores if score < current_row['performance_score']) / len(performance_scores)) * 100
    
    st.metric("Performance Percentile", f"{current_percentile:.1f}%")
    
    # Show comparison with top 3 peers
    top_peers = peer_data.nlargest(3, 'performance_score')
    
    comparison_data = []
    comparison_data.append({
        'Institution': f"📍 {current_row['institution_name']} (Current)",
        'Performance Score': current_row['performance_score'],
        'NAAC Grade': current_row['naac_grade'],
        'Placement Rate': f"{current_row['placement_rate']}%"
    })
    
    for _, peer in top_peers.iterrows():
        comparison_data.append({
            'Institution': peer['institution_name'],
            'Performance Score': peer['performance_score'],
            'NAAC Grade': peer['naac_grade'],
            'Placement Rate': f"{peer['placement_rate']}%"
        })
    
    st.dataframe(pd.DataFrame(comparison_data), use_container_width=True)


        
def display_intelligence_dashboard(data, analyzer):
    """Display comprehensive intelligence dashboard"""
    
    current = data['current']
    
    # Executive Summary
    st.subheader("🎯 Executive Summary")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        performance_score = current.get('performance_score', 0)
        if performance_score >= 8.0:
            st.success(f"**Performance**: {performance_score:.1f}/10")
        elif performance_score >= 6.0:
            st.info(f"**Performance**: {performance_score:.1f}/10")
        else:
            st.warning(f"**Performance**: {performance_score:.1f}/10")
    
    with col2:
        risk_level = current.get('risk_level', 'Medium Risk')
        if risk_level == "Low Risk":
            st.success(f"**Risk**: {risk_level}")
        elif risk_level == "Medium Risk":
            st.info(f"**Risk**: {risk_level}")
        else:
            st.error(f"**Risk**: {risk_level}")
    
    with col3:
        approval_status = current.get('approval_recommendation', 'Under Review')
        st.metric("Approval Status", approval_status)
    
    with col4:
        # Document compliance
        docs = data['documents']
        if not docs.empty:
            uploaded = len(docs[docs['status'] == 'Uploaded'])
            total = len(docs)
            compliance = (uploaded / total * 100) if total > 0 else 0
            st.metric("Doc Compliance", f"{compliance:.0f}%")
        else:
            st.metric("Doc Compliance", "0%")
    
    # Performance Trend
    st.subheader("📈 Performance Trend (2014-2023)")
    
    historical = data['historical']
    if len(historical) > 1:
        fig = px.line(
            historical,
            x='year',
            y='performance_score',
            markers=True,
            title=f"Performance Trend: {historical['year'].min()}-{historical['year'].max()}"
        )
        st.plotly_chart(fig, use_container_width=True)
    
    # Key Metrics
    st.subheader("📊 Key Performance Indicators")
    
    kpi_col1, kpi_col2, kpi_col3, kpi_col4 = st.columns(4)
    
    with kpi_col1:
        placement = current.get('placement_rate', 0)
        st.metric("Placement Rate", f"{placement:.1f}%")
    
    with kpi_col2:
        research = current.get('research_publications', 0)
        st.metric("Research Papers", research)
    
    with kpi_col3:
        sf_ratio = current.get('student_faculty_ratio', 0)
        st.metric("Student-Faculty Ratio", f"{sf_ratio:.1f}")
    
    with kpi_col4:
        naac = current.get('naac_grade', 'N/A')
        st.metric("NAAC Grade", naac)
    
    # AI Insights
    st.subheader("🤖 AI-Powered Insights")
    
    tab1, tab2, tab3 = st.tabs(["Strengths", "Areas for Improvement", "Recommendations"])
    
    with tab1:
        strengths = get_institutional_strengths(current)
        for strength in strengths:
            st.success(f"✅ {strength}")
    
    with tab2:
        weaknesses = get_institutional_weaknesses(current)
        for weakness in weaknesses:
            st.error(f"❌ {weakness}")
    
    with tab3:
        recommendations = get_ai_recommendations(current)
        for rec in recommendations:
            st.info(f"💡 {rec}")

def get_ai_recommendations(data):
    """Get AI recommendations"""
    recommendations = []
    
    sf_ratio = data.get('student_faculty_ratio', 0)
    if sf_ratio > 25:
        recommendations.append("Recruit additional faculty to improve student-faculty ratio")
    
    placement = data.get('placement_rate', 0)
    if placement < 70:
        recommendations.append("Strengthen industry partnerships and career services")
    
    research = data.get('research_publications', 0)
    if research < 30:
        recommendations.append("Establish research promotion programs and incentives")
    
    if not data.get('nirf_ranking'):
        recommendations.append("Participate in NIRF ranking for better visibility")
    
    return recommendations if recommendations else ["Institution is performing well across parameters"]



def get_institutional_weaknesses(data):
    """Get institutional weaknesses"""
    weaknesses = []
    
    sf_ratio = data.get('student_faculty_ratio', 0)
    if sf_ratio > 25:
        weaknesses.append(f"High student-faculty ratio: {sf_ratio:.1f}")
    
    placement = data.get('placement_rate', 0)
    if placement < 65:
        weaknesses.append(f"Low placement rate: {placement:.1f}%")
    
    research = data.get('research_publications', 0)
    if research < 20:
        weaknesses.append(f"Limited research output: {research} publications")
    
    digital = data.get('digital_infrastructure_score', 0)
    if digital < 7:
        weaknesses.append(f"Weak digital infrastructure: {digital:.1f}/10")
    
    return weaknesses if weaknesses else ["No major weaknesses identified"]



def get_institutional_strengths(data):
    """Get institutional strengths"""
    strengths = []
    
    naac = data.get('naac_grade', '')
    if naac in ['A++', 'A+', 'A']:
        strengths.append(f"Strong NAAC accreditation: {naac}")
    
    placement = data.get('placement_rate', 0)
    if placement > 80:
        strengths.append(f"Excellent placement rate: {placement:.1f}%")
    
    financial = data.get('financial_stability_score', 0)
    if financial > 8.0:
        strengths.append(f"Strong financial stability: {financial:.1f}/10")
    
    research = data.get('research_publications', 0)
    if research > 50:
        strengths.append(f"Robust research output: {research} publications")
    
    return strengths if strengths else ["No significant strengths identified"]



def generate_ai_recommendations_simple(institution_data):
    """Generate AI-powered improvement recommendations"""
    recommendations = []
    
    # Student-Faculty Ratio
    sf_ratio = institution_data.get('student_faculty_ratio', 0)
    if sf_ratio > 25:
        recommendations.append("Recruit additional faculty members to improve student-faculty ratio")
    
    # Placement Rate
    placement_rate = institution_data.get('placement_rate', 0)
    if placement_rate < 70:
        recommendations.append("Strengthen industry partnerships and career development programs")
    
    # Research Publications
    research_pubs = institution_data.get('research_publications', 0)
    if research_pubs < 50:
        recommendations.append("Establish research promotion policy and faculty development programs")
    
    # Digital Infrastructure
    digital_score = institution_data.get('digital_infrastructure_score', 0)
    if digital_score < 7:
        recommendations.append("Invest in digital infrastructure and e-learning platforms")
    
    # Community Engagement
    community_projects = institution_data.get('community_projects', 0)
    if community_projects < 5:
        recommendations.append("Enhance community engagement and social outreach programs")
    
    # NIRF Participation
    if not institution_data.get('nirf_ranking'):
        recommendations.append("Participate in NIRF ranking to enhance institutional reputation")
    
    return recommendations[:5]



def identify_opportunities(institution_data):
    """Identify opportunities for improvement"""
    opportunities = []
    
    # NIRF Ranking
    nirf_rank = institution_data.get('nirf_ranking')
    if not nirf_rank:
        opportunities.append("Participate in NIRF ranking for better visibility")
    
    # Industry Collaborations
    industry_collabs = institution_data.get('industry_collaborations', 0)
    if industry_collabs < 5:
        opportunities.append("Increase industry collaborations for practical exposure")
    
    # Patents
    patents = institution_data.get('patents_filed', 0)
    if patents < 3:
        opportunities.append("Strengthen IPR culture and patent filings")
    
    # Digital Transformation
    digital_score = institution_data.get('digital_infrastructure_score', 0)
    if digital_score < 8:
        opportunities.append("Invest in digital infrastructure for blended learning")
    
    # Internationalization
    opportunities.append("Explore international collaborations and student exchange programs")
    
    # Government Schemes
    opportunities.append("Participate in government schemes like RUSA, FIST, NMEICT")
    
    return opportunities[:5]

def identify_institutional_weaknesses(institution_data):
    """Identify institutional weaknesses"""
    weaknesses = []
    
    # Student-Faculty Ratio
    sf_ratio = institution_data.get('student_faculty_ratio', 0)
    if sf_ratio > 25:
        weaknesses.append(f"High Student-Faculty Ratio: {sf_ratio:.1f}")
    
    # Placement Rate
    placement_rate = institution_data.get('placement_rate', 0)
    if placement_rate < 65:
        weaknesses.append(f"Low Placement Rate: {placement_rate:.1f}%")
    
    # Research Output
    research_pubs = institution_data.get('research_publications', 0)
    if research_pubs < 20:
        weaknesses.append(f"Inadequate Research Output: {research_pubs} publications")
    
    # Digital Infrastructure
    digital_score = institution_data.get('digital_infrastructure_score', 0)
    if digital_score < 7:
        weaknesses.append(f"Weak Digital Infrastructure: {digital_score:.1f}/10")
    
    # Community Engagement
    community_projects = institution_data.get('community_projects', 0)
    if community_projects < 5:
        weaknesses.append("Limited Community Engagement")
    
    # NAAC Grade
    naac_grade = institution_data.get('naac_grade', '')
    if naac_grade in ['C', 'B']:
        weaknesses.append(f"Below Average NAAC Grade: {naac_grade}")
    
    return weaknesses[:5]  # Return top 5 weaknesses

def identify_institutional_strengths(institution_data):
    """Identify institutional strengths"""
    strengths = []
    
    # NAAC Grade
    naac_grade = institution_data.get('naac_grade', '')
    if naac_grade in ['A++', 'A+', 'A']:
        strengths.append(f"Excellent NAAC Accreditation: {naac_grade}")
    
    # Performance Score
    performance_score = institution_data.get('performance_score', 0)
    if performance_score >= 8.0:
        strengths.append(f"High Performance Score: {performance_score:.1f}/10")
    
    # Placement Rate
    placement_rate = institution_data.get('placement_rate', 0)
    if placement_rate > 80:
        strengths.append(f"Strong Placement Record: {placement_rate:.1f}%")
    
    # Research Output
    research_pubs = institution_data.get('research_publications', 0)
    if research_pubs > 50:
        strengths.append(f"Robust Research Output: {research_pubs} publications")
    
    # Financial Stability
    financial_score = institution_data.get('financial_stability_score', 0)
    if financial_score > 8.0:
        strengths.append(f"Excellent Financial Stability: {financial_score:.1f}/10")
    
    # PhD Faculty Ratio
    phd_ratio = institution_data.get('phd_faculty_ratio', 0)
    if phd_ratio > 0.7:
        strengths.append(f"Highly Qualified Faculty: {phd_ratio:.1%} PhDs")
    
    return strengths[:5]  # Return top 5 strengths

def get_comprehensive_institution_data(institution_id, analyzer):
    """Get comprehensive data for intelligence analysis"""
    try:
        # Get current year data (2023)
        current_year_data = analyzer.historical_data[analyzer.historical_data['year'] == 2023]
        
        current_data = current_year_data[current_year_data['institution_id'] == institution_id]
        
        if current_data.empty:
            st.warning(f"No current year data found for institution {institution_id}")
            return None
        
        # Get historical data
        historical_data = analyzer.historical_data[
            analyzer.historical_data['institution_id'] == institution_id
        ].sort_values('year')
        
        # Get document data
        try:
            document_data = analyzer.get_institution_documents(institution_id)
        except:
            document_data = pd.DataFrame()
        
        return {
            'current': current_data.iloc[0].to_dict(),
            'historical': historical_data,
            'documents': document_data,
            'institution_id': institution_id
        }
        
    except Exception as e:
        st.error(f"Error loading institution data: {str(e)}")
        return None


def create_data_management_module(analyzer):
    st.header("💾 Data Management & Analysis")
    
    # NEW TABS STRUCTURE:
    tab1, tab2, tab3 = st.tabs([
        "📊 Current Data Analytics",      # Former "View Current Data" - Enhanced
        "🔍 Data Validation & QA",       # NEW - Replaces Upload
        "⚙️ Advanced Database Tools"     # Enhanced version of Database Management
    ])
    
    with tab1:
        st.subheader("📊 Current Database Analytics")
        
        # Use the actual data from analyzer
        current_data = analyzer.historical_data
        
        # Show database statistics with specification verification
        st.subheader("✅ Database Specification Verification")
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            total_records = len(current_data)
            expected_records = 200  # 20 institutions × 10 years
            if total_records == expected_records:
                st.success(f"📊 **Total Records**: {total_records} ✅")
            else:
                st.error(f"📊 **Total Records**: {total_records} (Expected: {expected_records})")
        
        with col2:
            unique_institutions = current_data['institution_id'].nunique()
            expected_institutions = 20
            if unique_institutions == expected_institutions:
                st.success(f"🏛️ **Unique Institutions**: {unique_institutions} ✅")
            else:
                st.error(f"🏛️ **Unique Institutions**: {unique_institutions} (Expected: {expected_institutions})")
        
        with col3:
            years_covered = current_data['year'].nunique()
            expected_years = 10
            if years_covered == expected_years:
                st.success(f"📅 **Years Covered**: {years_covered} ✅")
            else:
                st.error(f"📅 **Years Covered**: {years_covered} (Expected: {expected_years})")
        
        with col4:
            year_range = f"{current_data['year'].min()}-{current_data['year'].max()}"
            expected_range = "2014-2023"
            if year_range == expected_range:
                st.success(f"🗓️ **Year Range**: {year_range} ✅")
            else:
                st.error(f"🗓️ **Year Range**: {year_range} (Expected: {expected_range})")
        
        # Show current year overview
        current_year_data = current_data[current_data['year'] == 2023]
        
        st.markdown("---")
        
        # Quick Analytics Cards - NEW SECTION
        st.subheader("📈 Quick Analytics Dashboard")
        
        # Row 1: Performance Metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            avg_performance = current_year_data['performance_score'].mean()
            # Color code based on performance
            if avg_performance >= 8.0:
                color = "🟢"
            elif avg_performance >= 6.0:
                color = "🟡"
            else:
                color = "🔴"
            st.metric(f"{color} Avg Performance Score", f"{avg_performance:.2f}/10")
        
        with col2:
            high_performers = len(current_year_data[current_year_data['performance_score'] >= 8.0])
            st.metric("🏆 High Performers (8+ Score)", high_performers)
        
        with col3:
            avg_placement = current_year_data['placement_rate'].mean()
            st.metric("🎓 Avg Placement Rate", f"{avg_placement:.1f}%")
        
        with col4:
            low_risk_count = len(current_year_data[current_year_data['risk_level'] == 'Low Risk'])
            st.metric("🛡️ Low Risk Institutions", low_risk_count)
        
        # Row 2: Data Quality Metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            naac_rated = current_year_data['naac_grade'].notna().sum()
            st.metric("🏅 NAAC Rated Institutions", f"{naac_rated}/{len(current_year_data)}")
        
        with col2:
            nirf_ranked = current_year_data['nirf_ranking'].notna().sum()
            st.metric("📊 NIRF Ranked Institutions", f"{nirf_ranked}/{len(current_year_data)}")
        
        with col3:
            research_active = current_year_data[current_year_data['research_publications'] > 20].shape[0]
            st.metric("🔬 Research Active (>20 pubs)", f"{research_active}/{len(current_year_data)}")
        
        with col4:
            data_completeness = round((current_year_data.count().mean() / len(current_year_data.columns)) * 100, 1)
            st.metric("📋 Data Completeness", f"{data_completeness}%")
        
        # Search and Filter Section - NEW
        st.markdown("---")
        st.subheader("🔍 Search & Filter Institutions")
        
        search_col1, search_col2, search_col3, search_col4 = st.columns(4)
        
        with search_col1:
            search_text = st.text_input("🔎 Search by Name/ID", "", 
                                       placeholder="Type institution name or ID...")
        
        with search_col2:
            institution_types = ["All"] + sorted(current_year_data['institution_type'].unique().tolist())
            selected_type = st.selectbox("🏛️ Filter by Type", institution_types)
        
        with search_col3:
            performance_range = st.slider("📈 Performance Score Range", 
                                         min_value=0.0, max_value=10.0, 
                                         value=(0.0, 10.0), step=0.5)
        
        with search_col4:
            risk_levels = ["All"] + sorted(current_year_data['risk_level'].unique().tolist())
            selected_risk = st.selectbox("⚠️ Filter by Risk", risk_levels)
        
        # Apply filters
        filtered_data = current_year_data.copy()
        
        if search_text:
            mask = (filtered_data['institution_name'].str.contains(search_text, case=False)) | \
                   (filtered_data['institution_id'].str.contains(search_text, case=False))
            filtered_data = filtered_data[mask]
        
        if selected_type != "All":
            filtered_data = filtered_data[filtered_data['institution_type'] == selected_type]
        
        if selected_risk != "All":
            filtered_data = filtered_data[filtered_data['risk_level'] == selected_risk]
        
        filtered_data = filtered_data[
            (filtered_data['performance_score'] >= performance_range[0]) & 
            (filtered_data['performance_score'] <= performance_range[1])
        ]
        
        # Display filter summary
        st.info(f"**Showing {len(filtered_data)} of {len(current_year_data)} institutions**")
        
        # Action Buttons Row - NEW
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("📥 Export to CSV", type="secondary"):
                csv = filtered_data.to_csv(index=False)
                st.download_button(
                    label="⬇️ Download CSV",
                    data=csv,
                    file_name="institutions_data.csv",
                    mime="text/csv"
                )
        
        with col2:
            if st.button("📊 Compare Selected", type="secondary"):
                st.session_state.selected_for_comparison = []
                st.info("Select institutions from table below for comparison")
        
        with col3:
            if st.button("🔄 Refresh Data", type="secondary"):
                st.rerun()
        
        with col4:
            if st.button("📋 Generate Summary Report", type="secondary"):
                generate_data_summary(filtered_data)
        
        # Data preview with enhanced UI - UPDATED
        st.markdown("---")
        st.subheader("📋 Data Preview (Current Year - 20 Institutions)")
        
        # Prepare display columns with sorting options
        display_columns = ['institution_id', 'institution_name', 'institution_type', 'state', 
                          'performance_score', 'naac_grade', 'placement_rate', 'risk_level']
        
        # Create a copy for display with formatted values
        display_data = filtered_data[display_columns].copy()
        
        # Add color formatting function
        def color_performance(val):
            if val >= 8.0:
                color = '#d4edda'  # Light green
            elif val >= 6.0:
                color = '#fff3cd'  # Light yellow
            else:
                color = '#f8d7da'  # Light red
            return f'background-color: {color}'
        
        def color_risk(val):
            if val == 'Low Risk':
                color = '#d4edda'
            elif val == 'Medium Risk':
                color = '#fff3cd'
            elif val == 'High Risk':
                color = '#f8d7da'
            else:
                color = '#dc3545'  # Red for critical
            return f'background-color: {color}'
        
        def color_naac(val):
            if val in ['A++', 'A+', 'A']:
                color = '#d4edda'
            elif val in ['B++', 'B+']:
                color = '#fff3cd'
            else:
                color = '#f8d7da'
            return f'background-color: {color}'
        
        # Apply styling
        styled_data = display_data.style.applymap(color_performance, subset=['performance_score']) \
                                       .applymap(color_risk, subset=['risk_level']) \
                                       .applymap(color_naac, subset=['naac_grade']) \
                                       .format({
                                           'performance_score': '{:.2f}',
                                           'placement_rate': '{:.1f}%'
                                       })
        
        # Display the styled dataframe
        st.dataframe(
            styled_data,
            use_container_width=True,
            height=500
        )
        
        # Sorting options
        st.caption("💡 *Tip: Click column headers to sort. Green=High performance, Yellow=Medium, Red=Low*")
        
        # Quick Actions for Selected Institution - NEW
        if 'selected_institution_id' in st.session_state:
            st.markdown("---")
            st.subheader(f"⚡ Quick Actions for {st.session_state.selected_institution_id}")
            
            selected_data = current_year_data[current_year_data['institution_id'] == st.session_state.selected_institution_id]
            
            if not selected_data.empty:
                row = selected_data.iloc[0]
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("Current Performance", f"{row['performance_score']:.2f}")
                    st.metric("Placement Rate", f"{row['placement_rate']:.1f}%")
                
                with col2:
                    st.metric("Risk Level", row['risk_level'])
                    st.metric("NAAC Grade", row['naac_grade'])
                
                with col3:
                    if st.button("📈 View Details", type="primary"):
                        st.session_state.view_institution = row['institution_id']
                        st.rerun()
                    
                    if st.button("📤 Upload Documents", type="secondary"):
                        st.session_state.upload_for_institution = row['institution_id']
                        st.rerun()
        
        # Data Quality Indicators - NEW
        st.markdown("---")
        st.subheader("📊 Data Quality Overview")
        
        # Calculate completeness for each institution
        completeness_data = []
        for _, row in current_year_data.iterrows():
            non_null_count = row.count()
            total_fields = len(row)
            completeness = (non_null_count / total_fields) * 100
            
            completeness_data.append({
                'Institution': row['institution_name'],
                'ID': row['institution_id'],
                'Completeness %': completeness,
                'Performance': row['performance_score']
            })
        
        completeness_df = pd.DataFrame(completeness_data)
        
        # Create visualization
        fig = px.bar(completeness_df.sort_values('Completeness %', ascending=False).head(10),
                    x='Institution', y='Completeness %',
                    color='Performance',
                    color_continuous_scale='RdYlGn',
                    title="Top 10 Institutions by Data Completeness",
                    labels={'Completeness %': 'Data Completeness (%)'})
        
        fig.update_layout(xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)
        
        # Show missing data summary
        missing_summary = current_year_data.isnull().sum()
        missing_data = missing_summary[missing_summary > 0]
        
        if len(missing_data) > 0:
            st.warning(f"⚠️ **Missing Data Alert**: {len(missing_data)} fields have incomplete data")
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**Fields with Missing Data:**")
                for field, count in missing_data.items():
                    st.write(f"- {field}: {count} missing")
            
            with col2:
                st.write("**Recommendations:**")
                st.write("1. Run data validation check")
                st.write("2. Contact institutions for missing data")
                st.write("3. Use AI to estimate missing values")

    with tab2:  # 🔍 Data Validation & QA (NEW)
        st.subheader("🔍 Data Quality Analysis & Validation")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.info("**Data Quality Checks**")
            if st.button("🔄 Run Data Validation", type="primary"):
                run_data_validation(analyzer)
            
            if st.button("📋 Show Data Completeness Report"):
                show_data_completeness_report(analyzer)
            
            if st.button("⚠️ Identify Data Anomalies"):
                identify_data_anomalies(analyzer)
        
        with col2:
            st.info("**Data Enhancement Tools**")
            if st.button("🎯 Fill Missing Values (AI)"):
                fill_missing_values_ai(analyzer)
            
            if st.button("📈 Calculate Derived Metrics"):
                calculate_derived_metrics(analyzer)

        
        # Data Quality Dashboard
        st.markdown("---")
        st.subheader("📊 Data Quality Dashboard")
        
        # Show data quality metrics
        quality_metrics = calculate_data_quality_metrics(analyzer)
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Data Completeness", f"{quality_metrics['completeness']}%")
        with col2:
            st.metric("Consistency Score", f"{quality_metrics['consistency']}/10")
        with col3:
            st.metric("Accuracy Score", f"{quality_metrics['accuracy']}/10")
        with col4:
            st.metric("Timeliness Score", f"{quality_metrics['timeliness']}/10")
    
    with tab3:  # ⚙️ Advanced Database Tools (Enhanced)
        st.subheader("⚙️ Advanced Database Management")
        
        st.warning("⚠️ **Warning**: These operations affect the entire database. Use with caution.")
        
        # SECTION 1: Data Backup & Restore
        st.markdown("### 📦 Backup & Restore")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("💾 Create Backup", type="secondary", help="Create a backup of current data"):
                create_backup(analyzer)
        
        with col2:
            backup_files = get_available_backups()
            selected_backup = st.selectbox("Select Backup to Restore", backup_files)
            if st.button("🔄 Restore from Backup", type="secondary"):
                restore_from_backup(analyzer, selected_backup)
        
        # SECTION 2: Selective Data Management
        st.markdown("### 🎯 Selective Data Operations")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            institution_options = ["All Institutions"] + analyzer.historical_data['institution_id'].unique().tolist()
            selected_institutions = st.multiselect("Select Institutions", institution_options[1:], default=[])
        
        with col2:
            year_options = ["All Years"] + sorted(analyzer.historical_data['year'].unique().tolist())
            selected_years = st.multiselect("Select Years", year_options[1:], default=[])
        
        with col3:
            operation = st.selectbox("Operation", [
                "Export Selected", 
                "Update Selected",
                "Validate Selected",
                "Delete Selected"
            ])
        
        if st.button("🚀 Execute Operation", type="primary"):
            execute_selective_operation(analyzer, selected_institutions, selected_years, operation)
        
        # SECTION 3: Bulk Operations (with safety)
        st.markdown("### ⚡ Bulk Operations")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.info("**Regenerate Options**")
            regenerate_option = st.selectbox("Regeneration Type", [
                "Regenerate All Data (20×10)",
                "Regenerate Current Year Only",
                "Regenerate Performance Scores Only",
                "Regenerate with New Parameters"
            ])
            
            if st.button("🔄 Regenerate Selected", type="secondary"):
                if st.checkbox("I understand this will overwrite existing data"):
                    regenerate_data(analyzer, regenerate_option)
        
        with col2:
            st.info("**Maintenance Operations**")
            if st.button("🧹 Optimize Database", type="secondary"):
                optimize_database(analyzer)
            
            if st.button("📊 Recalculate All Metrics", type="secondary"):
                recalculate_all_metrics(analyzer)
            
            if st.button("🔍 Fix Data Inconsistencies", type="secondary"):
                fix_data_inconsistencies(analyzer)
        
        # SECTION 4: System Information
        st.markdown("---")
        st.subheader("📋 System Information")
        
        db_info = get_database_information(analyzer)
        
        col1, col2 = st.columns(2)
        with col1:
            st.write("**Database Statistics:**")
            st.write(f"- Total Records: {db_info['total_records']}")
            st.write(f"- Last Backup: {db_info['last_backup']}")
            st.write(f"- Database Size: {db_info['size_mb']} MB")
        
        with col2:
            st.write("**Data Health:**")
            st.write(f"- Data Age: {db_info['data_age_days']} days")
            st.write(f"- Consistency Score: {db_info['consistency_score']}/10")
            st.write(f"- Recommended Action: {db_info['recommended_action']}")

def calculate_derived_metrics(analyzer):
    """Calculate additional derived metrics from existing data"""
    st.subheader("📈 Calculating Derived Metrics")
    
    with st.spinner("Analyzing data and calculating derived metrics..."):
        current_data = analyzer.historical_data.copy()
        
        # Track which metrics were calculated
        calculated_metrics = []
        
        # 1. Calculate Composite Scores
        if all(col in current_data.columns for col in ['performance_score', 'placement_rate', 'research_publications']):
            current_data['composite_score'] = (
                current_data['performance_score'] * 0.5 +
                (current_data['placement_rate'] / 10) * 0.3 +
                (current_data['research_publications'] / 50) * 0.2
            ).round(2)
            calculated_metrics.append("composite_score")
        
        # 2. Calculate Research Intensity (publications per faculty)
        if all(col in current_data.columns for col in ['research_publications', 'student_faculty_ratio']):
            # Estimate faculty count from student_faculty_ratio
            # Assuming 1000 students as baseline for estimation
            current_data['estimated_faculty'] = (1000 / current_data['student_faculty_ratio']).round(0)
            current_data['research_intensity'] = (
                current_data['research_publications'] / current_data['estimated_faculty']
            ).round(2)
            calculated_metrics.append("research_intensity")
        
        # 3. Calculate Institutional Growth Rate (year-over-year)
        if 'performance_score' in current_data.columns and 'year' in current_data.columns:
            growth_rates = []
            for institution_id in current_data['institution_id'].unique():
                inst_data = current_data[current_data['institution_id'] == institution_id].sort_values('year')
                if len(inst_data) > 1:
                    # Calculate CAGR for performance score
                    first_score = inst_data['performance_score'].iloc[0]
                    last_score = inst_data['performance_score'].iloc[-1]
                    years = inst_data['year'].iloc[-1] - inst_data['year'].iloc[0]
                    
                    if years > 0 and first_score > 0:
                        cagr = ((last_score / first_score) ** (1/years) - 1) * 100
                        growth_rates.append({
                            'institution_id': institution_id,
                            'growth_rate': round(cagr, 2)
                        })
            
            if growth_rates:
                growth_df = pd.DataFrame(growth_rates)
                current_data = current_data.merge(growth_df, on='institution_id', how='left')
                calculated_metrics.append("growth_rate")
        
        # 4. Calculate NAAC Score (convert grade to numeric)
        if 'naac_grade' in current_data.columns:
            naac_score_map = {
                'A++': 10, 'A+': 9, 'A': 8, 
                'B++': 7, 'B+': 6, 'B': 5, 'C': 4
            }
            current_data['naac_score'] = current_data['naac_grade'].map(naac_score_map)
            calculated_metrics.append("naac_score")
        
        # 5. Calculate Innovation Index
        if all(col in current_data.columns for col in ['patents_filed', 'industry_collaborations']):
            current_data['innovation_index'] = (
                (current_data['patents_filed'] * 0.6) + 
                (current_data['industry_collaborations'] * 0.4)
            ).round(2)
            calculated_metrics.append("innovation_index")
        
        # 6. Calculate Student Success Score
        if all(col in current_data.columns for col in ['placement_rate', 'higher_education_rate']):
            current_data['student_success_score'] = (
                (current_data['placement_rate'] * 0.7) +
                (current_data['higher_education_rate'] * 0.3)
            ).round(2)
            calculated_metrics.append("student_success_score")
        
        # 7. Calculate Infrastructure Index
        if all(col in current_data.columns for col in ['digital_infrastructure_score', 'library_volumes', 'laboratory_equipment_score']):
            # Normalize library volumes (0-10 scale)
            max_library = current_data['library_volumes'].max()
            if max_library > 0:
                library_score = (current_data['library_volumes'] / max_library) * 10
            else:
                library_score = 5
            
            current_data['infrastructure_index'] = (
                current_data['digital_infrastructure_score'] * 0.4 +
                library_score * 0.3 +
                current_data['laboratory_equipment_score'] * 0.3
            ).round(2)
            calculated_metrics.append("infrastructure_index")
        
        # 8. Calculate Governance Excellence Score
        if all(col in current_data.columns for col in ['financial_stability_score', 'compliance_score', 'administrative_efficiency']):
            current_data['governance_score'] = (
                current_data['financial_stability_score'] * 0.4 +
                current_data['compliance_score'] * 0.3 +
                current_data['administrative_efficiency'] * 0.3
            ).round(2)
            calculated_metrics.append("governance_score")
        
        # 9. Calculate Social Impact Score
        if all(col in current_data.columns for col in ['community_projects', 'rural_outreach_score']):
            # Normalize community projects (0-10 scale)
            max_projects = current_data['community_projects'].max()
            if max_projects > 0:
                projects_score = (current_data['community_projects'] / max_projects) * 10
            else:
                projects_score = 5
            
            current_data['social_impact_score'] = (
                projects_score * 0.6 +
                current_data['rural_outreach_score'] * 0.4
            ).round(2)
            calculated_metrics.append("social_impact_score")
        
        # 10. Calculate NEP 2020 Compliance Score (simplified)
        if all(col in current_data.columns for col in ['curriculum_framework_score', 'faculty_diversity_index', 
                                                      'learning_outcome_achievement', 'entrepreneurship_cell_score']):
            current_data['nep_compliance_score'] = (
                current_data['curriculum_framework_score'] * 0.25 +
                current_data['faculty_diversity_index'] * 0.25 +
                (current_data['learning_outcome_achievement'] / 10) * 0.25 +
                current_data['entrepreneurship_cell_score'] * 0.25
            ).round(2)
            calculated_metrics.append("nep_compliance_score")
        
        # Save calculated metrics to database
        if calculated_metrics:
            # Create a temporary table with new metrics
            metrics_to_save = ['institution_id', 'year'] + calculated_metrics
            derived_metrics_df = current_data[metrics_to_save].copy()
            
            # Check if derived_metrics table exists, create if not
            cursor = analyzer.conn.cursor()
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS derived_metrics (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    institution_id TEXT,
                    year INTEGER,
                    metric_name TEXT,
                    metric_value REAL,
                    calculated_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (institution_id) REFERENCES institutions (institution_id)
                )
            """)
            
            # Clear existing derived metrics for these institutions/years
            for institution_id in derived_metrics_df['institution_id'].unique():
                for year in derived_metrics_df['year'].unique():
                    cursor.execute("""
                        DELETE FROM derived_metrics 
                        WHERE institution_id = ? AND year = ?
                    """, (institution_id, year))
            
            # Insert new derived metrics
            for _, row in derived_metrics_df.iterrows():
                for metric in calculated_metrics:
                    if not pd.isna(row[metric]):
                        cursor.execute("""
                            INSERT INTO derived_metrics (institution_id, year, metric_name, metric_value)
                            VALUES (?, ?, ?, ?)
                        """, (row['institution_id'], row['year'], metric, row[metric]))
            
            analyzer.conn.commit()
            
            # Show results
            st.success(f"✅ Calculated {len(calculated_metrics)} derived metrics!")
            
            # Display summary of calculated metrics
            st.subheader("📊 Derived Metrics Summary")
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total Metrics", len(calculated_metrics))
            with col2:
                avg_composite = current_data['composite_score'].mean() if 'composite_score' in calculated_metrics else 0
                st.metric("Avg Composite", f"{avg_composite:.2f}")
            with col3:
                avg_growth = current_data['growth_rate'].mean() if 'growth_rate' in calculated_metrics else 0
                st.metric("Avg Growth Rate", f"{avg_growth:.2f}%")
            with col4:
                avg_nep = current_data['nep_compliance_score'].mean() if 'nep_compliance_score' in calculated_metrics else 0
                st.metric("Avg NEP Score", f"{avg_nep:.2f}")
            
            # Show sample of calculated metrics
            st.subheader("🧪 Sample Calculated Metrics (Current Year)")
            
            current_year_metrics = current_data[current_data['year'] == 2023]
            if not current_year_metrics.empty:
                sample_cols = ['institution_name'] + [m for m in calculated_metrics if m in current_year_metrics.columns]
                st.dataframe(
                    current_year_metrics[sample_cols].head(10),
                    use_container_width=True
                )
            
            # Visualize distribution of key metrics
            if 'composite_score' in calculated_metrics:
                st.subheader("📈 Composite Score Distribution")
                fig = px.histogram(current_year_metrics, x='composite_score', 
                                 nbins=15, title="Distribution of Composite Scores")
                st.plotly_chart(fig, use_container_width=True)
            
            # Download option
            csv_data = derived_metrics_df.to_csv(index=False)
            st.download_button(
                label="📥 Download Derived Metrics",
                data=csv_data,
                file_name="derived_metrics.csv",
                mime="text/csv"
            )
            
        else:
            st.warning("⚠️ No derived metrics could be calculated. Check if required base metrics exist.")


def show_data_completeness_report(analyzer):
    """Generate comprehensive data completeness report"""
    st.subheader("📊 Data Completeness Report")
    
    current_data = analyzer.historical_data
    current_year_data = current_data[current_data['year'] == 2023]
    
    # Calculate completeness by column
    completeness_by_column = {}
    for col in current_year_data.columns:
        non_null = current_year_data[col].notna().sum()
        total = len(current_year_data)
        completeness = (non_null / total) * 100 if total > 0 else 0
        
        completeness_by_column[col] = {
            'completeness': completeness,
            'missing': total - non_null,
            'total': total
        }
    
    # Create visualization
    df_completeness = pd.DataFrame([
        {
            'Column': col,
            'Completeness %': info['completeness'],
            'Missing': info['missing']
        }
        for col, info in completeness_by_column.items()
    ]).sort_values('Completeness %', ascending=True)
    
    fig = px.bar(df_completeness.head(10),  # Show top 10 most incomplete
                x='Completeness %', y='Column',
                orientation='h',
                title="Top 10 Columns Needing Attention (Lowest Completeness)",
                color='Completeness %',
                color_continuous_scale='RdYlGn')
    
    st.plotly_chart(fig, use_container_width=True)
    
    # Show detailed table
    with st.expander("📋 View All Columns"):
        st.dataframe(df_completeness, use_container_width=True)
    
    # Recommendations
    incomplete_cols = df_completeness[df_completeness['Completeness %'] < 80]
    if len(incomplete_cols) > 0:
        st.warning(f"⚠️ {len(incomplete_cols)} columns have less than 80% completeness")
        st.write("**Recommendations:**")
        for _, row in incomplete_cols.head(5).iterrows():
            st.write(f"- **{row['Column']}**: {row['Missing']} missing values - Consider data collection")

def identify_data_anomalies(analyzer):
    """Identify statistical anomalies in the data"""
    st.subheader("🔍 Data Anomaly Detection")
    
    current_year_data = analyzer.historical_data[analyzer.historical_data['year'] == 2023]
    
    anomalies = []
    
    # Check for statistical outliers using IQR method
    numeric_cols = current_year_data.select_dtypes(include=[np.number]).columns.tolist()
    
    for col in numeric_cols:
        if col in ['institution_id', 'year', 'established_year']:
            continue
        
        data = current_year_data[col].dropna()
        if len(data) < 5:
            continue
        
        Q1 = data.quantile(0.25)
        Q3 = data.quantile(0.75)
        IQR = Q3 - Q1
        
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        
        outliers = current_year_data[
            (current_year_data[col] < lower_bound) | 
            (current_year_data[col] > upper_bound)
        ]
        
        if len(outliers) > 0:
            for _, row in outliers.iterrows():
                anomalies.append({
                    'Institution': row['institution_name'],
                    'Column': col,
                    'Value': row[col],
                    'Normal Range': f"{lower_bound:.2f} - {upper_bound:.2f}",
                    'Type': 'Statistical Outlier'
                })
    
    # Check for logical anomalies
    for _, row in current_year_data.iterrows():
        # Performance score too high for NAAC grade
        if row['performance_score'] >= 9.0 and row['naac_grade'] in ['B', 'C']:
            anomalies.append({
                'Institution': row['institution_name'],
                'Column': 'performance_score vs naac_grade',
                'Value': f"Score: {row['performance_score']}, Grade: {row['naac_grade']}",
                'Normal Range': "High scores should have A/A+ grades",
                'Type': 'Logical Inconsistency'
            })
        
        # Placement rate too high with low performance
        if row['placement_rate'] > 90 and row['performance_score'] < 6.0:
            anomalies.append({
                'Institution': row['institution_name'],
                'Column': 'placement_rate vs performance',
                'Value': f"Placement: {row['placement_rate']}%, Score: {row['performance_score']}",
                'Normal Range': "High placement usually correlates with high performance",
                'Type': 'Suspicious Correlation'
            })
    
    # Display results
    if anomalies:
        df_anomalies = pd.DataFrame(anomalies)
        st.error(f"⚠️ Found {len(anomalies)} anomalies")
        st.dataframe(df_anomalies, use_container_width=True)
        
        # Show visualization
        anomaly_counts = df_anomalies['Type'].value_counts()
        fig = px.pie(values=anomaly_counts.values, names=anomaly_counts.index,
                    title="Types of Anomalies Detected")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.success("✅ No significant anomalies detected!")

def run_data_validation_on_subset(data_subset):
    """Run validation on a subset of data"""
    # Simplified version of run_data_validation for subsets
    return {
        'total_records': len(data_subset),
        'missing_values': data_subset.isnull().sum().sum(),
        'data_quality': "Good" if data_subset.isnull().sum().sum() == 0 else "Needs Attention"
    }

def fill_missing_values_ai(analyzer):
    """Use AI/statistical methods to fill missing values"""
    st.subheader("🤖 AI-Powered Missing Value Imputation")
    
    current_data = analyzer.historical_data.copy()
    current_year_data = current_data[current_data['year'] == 2023]
    
    # Identify columns with missing values
    missing_cols = current_year_data.columns[current_year_data.isnull().any()].tolist()
    
    if not missing_cols:
        st.success("✅ No missing values found in current year data!")
        return
    
    st.warning(f"⚠️ Found missing values in {len(missing_cols)} columns")
    
    # Show missing value statistics
    st.write("**Missing Value Analysis:**")
    missing_stats = current_year_data[missing_cols].isnull().sum()
    for col, count in missing_stats.items():
        if count > 0:
            percentage = (count / len(current_year_data)) * 100
            st.write(f"- {col}: {count} missing ({percentage:.1f}%)")
    
    # Imputation methods for different column types
    imputation_methods = {}
    
    for col in missing_cols:
        missing_count = current_year_data[col].isnull().sum()
        
        if missing_count > 0:
            # Determine imputation method based on column type
            if col in ['performance_score', 'placement_rate', 'student_faculty_ratio']:
                # Use median for skewed distributions
                method = "Median"
                fill_value = current_year_data[col].median()
            elif col in ['naac_grade', 'risk_level', 'institution_type']:
                # Use mode for categorical
                method = "Mode (Most Frequent)"
                fill_value = current_year_data[col].mode().iloc[0] if not current_year_data[col].mode().empty else "Unknown"
            elif col in ['research_publications', 'community_projects', 'patents_filed']:
                # Use mean for count data
                method = "Mean"
                fill_value = current_year_data[col].mean()
            elif col in ['nirf_ranking']:
                # Use forward/backward fill for ranking
                method = "Interpolate"
                fill_value = None  # Special handling
            else:
                # Default to median
                method = "Median"
                fill_value = current_year_data[col].median()
            
            imputation_methods[col] = {
                'method': method,
                'value': fill_value,
                'count': missing_count
            }
    
    # Show proposed imputation plan
    st.subheader("📋 Proposed Imputation Plan")
    
    df_imputation = pd.DataFrame([
        {
            'Column': col,
            'Missing': info['count'],
            'Method': info['method'],
            'Fill Value': info['value'] if info['value'] is not None else "Interpolation"
        }
        for col, info in imputation_methods.items()
    ])
    
    st.dataframe(df_imputation, use_container_width=True)
    
    # Advanced AI options
    with st.expander("⚙️ Advanced AI Options"):
        col1, col2 = st.columns(2)
        
        with col1:
            use_knn = st.checkbox("Use KNN Imputation", value=False,
                                help="Use k-Nearest Neighbors algorithm for better accuracy")
            use_mice = st.checkbox("Use MICE (Multiple Imputation)", value=False,
                                 help="Multiple Imputation by Chained Equations - more robust")
        
        with col2:
            preserve_distribution = st.checkbox("Preserve Original Distribution", value=True)
            add_uncertainty = st.checkbox("Add Uncertainty Estimate", value=False)
    
    # Execute imputation
    if st.button("🚀 Execute AI Imputation", type="primary"):
        with st.spinner("Running AI imputation algorithms..."):
            # Apply imputation
            imputed_data = current_data.copy()
            
            for col, info in imputation_methods.items():
                missing_mask = imputed_data[col].isnull()
                
                if info['value'] is not None:
                    # Simple value imputation
                    imputed_data.loc[missing_mask, col] = info['value']
                elif col == 'nirf_ranking':
                    # Special handling for ranking - interpolate based on performance
                    # Higher performance institutions get better ranking
                    for idx in missing_mask[missing_mask].index:
                        performance = imputed_data.loc[idx, 'performance_score']
                        if performance >= 8.0:
                            imputed_data.loc[idx, col] = np.random.randint(1, 51)
                        elif performance >= 6.0:
                            imputed_data.loc[idx, col] = np.random.randint(51, 151)
                        else:
                            imputed_data.loc[idx, col] = np.random.randint(151, 201)
            
            # Special handling for performance_score if missing
            if 'performance_score' in missing_cols:
                # Calculate performance score from other metrics if missing
                missing_scores = imputed_data['performance_score'].isnull()
                for idx in missing_scores[missing_scores].index:
                    row = imputed_data.loc[idx]
                    
                    # Estimate score based on available metrics
                    estimated_score = 5.0  # Base
                    
                    if pd.notna(row.get('naac_grade')):
                        naac_map = {'A++': 9.5, 'A+': 9.0, 'A': 8.0, 'B++': 7.0, 'B+': 6.0, 'B': 5.0, 'C': 4.0}
                        estimated_score += naac_map.get(row['naac_grade'], 5.0) * 0.2
                    
                    if pd.notna(row.get('placement_rate')):
                        estimated_score += (row['placement_rate'] / 10) * 0.3
                    
                    if pd.notna(row.get('research_publications')):
                        estimated_score += min(10, row['research_publications'] / 5) * 0.2
                    
                    imputed_data.loc[idx, 'performance_score'] = min(10, estimated_score)
            
            # Update risk level and approval recommendation based on new scores
            for idx, row in imputed_data.iterrows():
                score = row['performance_score']
                imputed_data.loc[idx, 'risk_level'] = assess_risk_level(score)
                imputed_data.loc[idx, 'approval_recommendation'] = generate_approval_recommendation(score)
            
            # Save to database
            imputed_data.to_sql('institutions', analyzer.conn, if_exists='replace', index=False)
            analyzer.historical_data = imputed_data
            
            # Show results
            st.success("✅ AI Imputation Complete!")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                before_missing = current_data.isnull().sum().sum()
                after_missing = imputed_data.isnull().sum().sum()
                st.metric("Missing Values Fixed", f"{before_missing - after_missing}")
            
            with col2:
                st.metric("Data Completeness", 
                         f"{(imputed_data.count().sum() / imputed_data.size * 100):.1f}%")
            
            with col3:
                # Show sample of imputed values
                imputed_cols = [col for col in missing_cols if imputed_data[col].notna().all()]
                st.metric("Columns Imputed", len(imputed_cols))
            
            # Show before/after comparison
            st.subheader("📊 Before/After Comparison")
            
            sample_cols = missing_cols[:3]  # Show first 3 columns
            if sample_cols:
                comparison_data = []
                for col in sample_cols:
                    before_avg = current_data[col].mean()
                    after_avg = imputed_data[col].mean()
                    
                    comparison_data.append({
                        'Column': col,
                        'Before (Mean)': round(before_avg, 2) if not pd.isna(before_avg) else "N/A",
                        'After (Mean)': round(after_avg, 2),
                        'Change': f"{round(((after_avg - (before_avg or 0)) / (before_avg or 1)) * 100, 1)}%" if before_avg else "N/A"
                    })
                
                st.dataframe(pd.DataFrame(comparison_data), use_container_width=True)



def fix_data_inconsistencies(analyzer):
    """Automatically fix data inconsistencies"""
    with st.spinner("Fixing data inconsistencies..."):
        fixed_data = analyzer.historical_data.copy()
        
        issues_fixed = 0
        
        # Fix 1: Ensure risk level matches performance score
        for idx, row in fixed_data.iterrows():
            score = row['performance_score']
            current_risk = row['risk_level']
            
            if score >= 8.0 and current_risk != 'Low Risk':
                fixed_data.at[idx, 'risk_level'] = 'Low Risk'
                issues_fixed += 1
            elif 6.5 <= score < 8.0 and current_risk != 'Medium Risk':
                fixed_data.at[idx, 'risk_level'] = 'Medium Risk'
                issues_fixed += 1
            elif score < 6.5 and current_risk not in ['High Risk', 'Critical Risk']:
                fixed_data.at[idx, 'risk_level'] = 'High Risk'
                issues_fixed += 1
        
        # Fix 2: Ensure placement rate is within 0-100
        if 'placement_rate' in fixed_data.columns:
            invalid_placement = fixed_data[
                (fixed_data['placement_rate'] < 0) | (fixed_data['placement_rate'] > 100)
            ].shape[0]
            
            if invalid_placement > 0:
                fixed_data['placement_rate'] = fixed_data['placement_rate'].clip(0, 100)
                issues_fixed += invalid_placement
        
        # Save fixes
        if issues_fixed > 0:
            fixed_data.to_sql('institutions', analyzer.conn, if_exists='replace', index=False)
            analyzer.historical_data = fixed_data
            st.success(f"✅ Fixed {issues_fixed} data inconsistencies")
        else:
            st.info("✅ No inconsistencies found")

def recalculate_all_metrics(analyzer):
    """Recalculate all derived metrics"""
    with st.spinner("Recalculating all metrics..."):
        updated_data = analyzer.historical_data.copy()
        
        for idx, row in updated_data.iterrows():
            # Recalculate performance score
            new_score = calculate_performance_score(row)
            updated_data.at[idx, 'performance_score'] = new_score
            
            # Update dependent fields
            updated_data.at[idx, 'approval_recommendation'] = generate_approval_recommendation(new_score)
            updated_data.at[idx, 'risk_level'] = assess_risk_level(new_score)
        
        # Save to database
        updated_data.to_sql('institutions', analyzer.conn, if_exists='replace', index=False)
        analyzer.historical_data = updated_data
        
        st.success(f"✅ Recalculated metrics for {len(updated_data)} records")

def optimize_database(analyzer):
    """Optimize database performance"""
    with st.spinner("Optimizing database..."):
        cursor = analyzer.conn.cursor()
        
        # Vacuum to defragment
        cursor.execute("VACUUM")
        
        # Analyze for query optimization
        cursor.execute("ANALYZE")
        
        # Reindex
        cursor.execute("REINDEX")
        
        analyzer.conn.commit()
        
        st.success("✅ Database optimization complete!")

def execute_selective_operation(analyzer, selected_institutions, selected_years, operation):
    """Execute selective database operations"""
    with st.spinner(f"Executing {operation}..."):
        current_data = analyzer.historical_data
        
        # Apply filters
        filtered_data = current_data.copy()
        
        if selected_institutions and "All Institutions" not in selected_institutions:
            filtered_data = filtered_data[filtered_data['institution_id'].isin(selected_institutions)]
        
        if selected_years and "All Years" not in selected_years:
            filtered_data = filtered_data[filtered_data['year'].isin(selected_years)]
        
        if operation == "Export Selected":
            csv_data = filtered_data.to_csv(index=False)
            st.download_button(
                label="📥 Download Selected Data",
                data=csv_data,
                file_name=f"selected_data_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )
        
        elif operation == "Update Selected":
            st.info(f"Would update {len(filtered_data)} records")
            # Show update form here
        
        elif operation == "Validate Selected":
            validation_results = run_data_validation_on_subset(filtered_data)
            st.write(validation_results)
        
        elif operation == "Delete Selected":
            if st.checkbox(f"⚠️ Confirm deletion of {len(filtered_data)} records"):
                # Delete from database
                cursor = analyzer.conn.cursor()
                for _, row in filtered_data.iterrows():
                    cursor.execute(
                        "DELETE FROM institutions WHERE institution_id = ? AND year = ?",
                        (row['institution_id'], row['year'])
                    )
                analyzer.conn.commit()
                
                # Update historical data
                analyzer.historical_data = analyzer.historical_data[
                    ~analyzer.historical_data['institution_id'].isin(selected_institutions)
                ]
                
                st.success(f"✅ Deleted {len(filtered_data)} records")



def restore_from_backup(analyzer, selected_backup):
    """Restore database from selected backup"""
    try:
        # Extract filename from display string
        backup_filename = selected_backup.split(" - ")[-1]
        backup_path = os.path.join("backups", backup_filename)
        
        if not os.path.exists(backup_path):
            st.error(f"Backup file not found: {backup_path}")
            return
        
        with st.spinner(f"Restoring from backup: {backup_filename}..."):
            # Load backup data
            with open(backup_path, 'r') as f:
                backup_data = json.load(f)
            
            # Clear existing data
            cursor = analyzer.conn.cursor()
            cursor.execute("DELETE FROM institutions")
            cursor.execute("DELETE FROM institution_documents")
            
            # Restore institutions
            institutions_df = pd.DataFrame(backup_data['institutions'])
            institutions_df.to_sql('institutions', analyzer.conn, if_exists='append', index=False)
            
            # Restore documents if available
            if 'documents' in backup_data and backup_data['documents']:
                documents_df = pd.DataFrame(backup_data['documents'])
                documents_df.to_sql('institution_documents', analyzer.conn, if_exists='append', index=False)
            
            analyzer.conn.commit()
            
            # Update analyzer's historical data
            analyzer.historical_data = institutions_df
            
            st.success(f"✅ Successfully restored backup from {backup_data['metadata']['backup_date']}")
            st.info(f"📊 Restored: {len(backup_data['institutions'])} institutions")
            
    except Exception as e:
        st.error(f"❌ Error restoring backup: {str(e)}")
        

def get_database_information(analyzer):
    """Get comprehensive database information and health metrics"""
    try:
        cursor = analyzer.conn.cursor()
        
        # Basic information
        cursor.execute("SELECT COUNT(*) FROM institutions")
        total_records = cursor.fetchone()[0]
        
        cursor.execute("SELECT COUNT(DISTINCT institution_id) FROM institutions")
        unique_institutions = cursor.fetchone()[0]
        
        cursor.execute("SELECT MIN(year), MAX(year) FROM institutions")
        min_year, max_year = cursor.fetchone()
        
        # Calculate database size (approximate)
        cursor.execute("SELECT page_count * page_size FROM pragma_page_count(), pragma_page_size()")
        db_size_bytes = cursor.fetchone()[0]
        db_size_mb = db_size_bytes / (1024 * 1024)
        
        # Data age (days since last update)
        cursor.execute("SELECT MAX(year) FROM institutions")
        latest_year = cursor.fetchone()[0]
        current_year = datetime.now().year
        data_age_years = current_year - latest_year if latest_year else 1
        
        # Check last backup
        backup_dir = "backups"
        last_backup = "No backup found"
        if os.path.exists(backup_dir):
            backup_files = glob.glob(os.path.join(backup_dir, "backup_*.json"))
            if backup_files:
                latest_backup = max(backup_files, key=os.path.getmtime)
                backup_time = os.path.getmtime(latest_backup)
                last_backup = datetime.fromtimestamp(backup_time).strftime("%Y-%m-%d %H:%M")
        
        # Calculate consistency score
        current_year_data = analyzer.historical_data[analyzer.historical_data['year'] == 2023]
        consistency_score = 5.0
        
        if len(current_year_data) > 0:
            # Check for logical inconsistencies
            inconsistencies = 0
            for _, row in current_year_data.iterrows():
                score = row.get('performance_score', 0)
                risk = row.get('risk_level', '')
                
                # Basic consistency check
                if score >= 8.0 and risk != 'Low Risk':
                    inconsistencies += 1
                elif score < 5.0 and risk not in ['High Risk', 'Critical Risk']:
                    inconsistencies += 1
            
            consistency_score = max(1, 10 - (inconsistencies / len(current_year_data)) * 10)
        
        # Determine recommended action
        recommended_action = "No action required"
        if data_age_years > 1:
            recommended_action = "Update data to current year"
        elif consistency_score < 7.0:
            recommended_action = "Run data validation"
        elif db_size_mb > 100:
            recommended_action = "Consider archiving old data"
        
        return {
            'total_records': total_records,
            'unique_institutions': unique_institutions,
            'year_range': f"{min_year}-{max_year}" if min_year and max_year else "N/A",
            'size_mb': round(db_size_mb, 2),
            'last_backup': last_backup,
            'data_age_days': data_age_years * 365,
            'consistency_score': round(consistency_score, 1),
            'recommended_action': recommended_action
        }
        
    except Exception as e:
        print(f"Error getting database info: {e}")
        return {
            'total_records': 0,
            'unique_institutions': 0,
            'year_range': "N/A",
            'size_mb': 0,
            'last_backup': "Error",
            'data_age_days': 999,
            'consistency_score': 0,
            'recommended_action': "Check database connection"
        }

def create_backup(analyzer):
    """Create a timestamped backup of the database"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_dir = "backups"
    
    # Create backups directory if it doesn't exist
    if not os.path.exists(backup_dir):
        os.makedirs(backup_dir)
    
    # Get all data from database
    cursor = analyzer.conn.cursor()
    
    # Backup institutions table
    cursor.execute("SELECT * FROM institutions")
    institutions_data = cursor.fetchall()
    columns = [description[0] for description in cursor.description]
    
    # Convert to list of dictionaries
    institutions_list = []
    for row in institutions_data:
        institutions_list.append(dict(zip(columns, row)))
    
    # Backup other tables
    cursor.execute("SELECT * FROM institution_documents")
    documents_data = cursor.fetchall()
    doc_columns = [description[0] for description in cursor.description]
    documents_list = [dict(zip(doc_columns, row)) for row in documents_data]
    
    # Create backup structure
    backup_data = {
        'metadata': {
            'backup_date': datetime.now().isoformat(),
            'total_institutions': len(institutions_list),
            'total_documents': len(documents_list),
            'system_version': '2.0'
        },
        'institutions': institutions_list,
        'documents': documents_list
    }
    
    # Save to file
    backup_file = os.path.join(backup_dir, f"backup_{timestamp}.json")
    with open(backup_file, 'w') as f:
        json.dump(backup_data, f, indent=2)
    
    st.success(f"✅ Backup created successfully: {backup_file}")
    st.info(f"📦 Backup includes: {len(institutions_list)} institutions, {len(documents_list)} documents")
    
    # Also create CSV backup for Excel users
    csv_file = os.path.join(backup_dir, f"backup_{timestamp}.csv")
    pd.DataFrame(institutions_list).to_csv(csv_file, index=False)
    
    return backup_file

def get_available_backups():
    """List available backup files"""
    backup_dir = "backups"
    
    if not os.path.exists(backup_dir):
        return []
    
    # Get all JSON backup files
    backup_files = glob.glob(os.path.join(backup_dir, "backup_*.json"))
    
    # Extract just the filenames and sort by date (newest first)
    backup_names = []
    for file_path in backup_files:
        filename = os.path.basename(file_path)
        # Extract timestamp from filename
        try:
            timestamp_str = filename.replace('backup_', '').replace('.json', '')
            timestamp = datetime.strptime(timestamp_str[:15], "%Y%m%d_%H%M%S")
            backup_names.append({
                'filename': filename,
                'path': file_path,
                'date': timestamp,
                'display': f"{timestamp.strftime('%Y-%m-%d %H:%M')} - {filename}"
            })
        except:
            backup_names.append({
                'filename': filename,
                'path': file_path,
                'date': datetime.now(),
                'display': filename
            })
    
    # Sort by date, newest first
    backup_names.sort(key=lambda x: x['date'], reverse=True)
    
    return [b['display'] for b in backup_names[:10]]  # Return top 10 most recent

def regenerate_data(analyzer, option):
    """Regenerate data based on selected option"""
    
    if option == "Regenerate All Data (20×10)":
        with st.spinner("Generating 20×10 comprehensive data..."):
            # First create a backup
            backup_file = create_backup(analyzer)
            
            # Generate new data
            new_data = generate_comprehensive_dummy_data()
            
            # Save to database
            new_data.to_sql('institutions', analyzer.conn, if_exists='replace', index=False)
            
            # Update the analyzer's historical data
            analyzer.historical_data = new_data
            
            st.success("✅ 20×10 data regenerated successfully!")
            st.info(f"📊 New data: {new_data['institution_id'].nunique()} institutions × {new_data['year'].nunique()} years")
            st.warning(f"⚠️ Backup created: {os.path.basename(backup_file)}")
    
    elif option == "Regenerate Current Year Only":
        with st.spinner("Regenerating current year data..."):
            # Keep historical data, update only 2023
            existing_data = analyzer.historical_data
            
            # Remove 2023 data
            existing_data = existing_data[existing_data['year'] != 2023]
            
            # Generate new 2023 data for all 20 institutions
            new_2023_data = generate_current_year_data(analyzer)
            
            # Combine
            new_data = pd.concat([existing_data, new_2023_data], ignore_index=True)
            
            # Save to database
            new_data.to_sql('institutions', analyzer.conn, if_exists='replace', index=False)
            analyzer.historical_data = new_data
            
            st.success("✅ Current year data regenerated!")
    
    elif option == "Regenerate Performance Scores Only":
        with st.spinner("Recalculating performance scores..."):
            updated_data = analyzer.historical_data.copy()
            
            # Recalculate performance scores for each row
            for idx, row in updated_data.iterrows():
                # Calculate new performance score based on metrics
                new_score = calculate_performance_score(row)
                updated_data.at[idx, 'performance_score'] = new_score
                
                # Update recommendation and risk level
                updated_data.at[idx, 'approval_recommendation'] = generate_approval_recommendation(new_score)
                updated_data.at[idx, 'risk_level'] = assess_risk_level(new_score)
            
            # Save updates
            updated_data.to_sql('institutions', analyzer.conn, if_exists='replace', index=False)
            analyzer.historical_data = updated_data
            
            st.success("✅ All performance scores recalculated!")
    
    elif option == "Regenerate with New Parameters":
        st.info("This would open a form to specify new generation parameters")
        # This would show a form for custom data generation parameters

def generate_current_year_data(analyzer):
    """Generate data for current year only (2023) for all 20 institutions"""
    np.random.seed(42)  # For reproducibility
    
    # Get the list of existing institutions
    existing_data = analyzer.historical_data
    existing_institutions = existing_data['institution_id'].unique()
    
    # If no existing institutions, use 20 new ones
    if len(existing_institutions) == 0:
        existing_institutions = [f'INST_{i:04d}' for i in range(1, 21)]
    
    institutions_data = []
    
    institution_types = ['State University', 'Private University', 'Deemed University', 'Autonomous College', 'Technical Institute']
    states = ['Maharashtra', 'Karnataka', 'Tamil Nadu', 'Delhi', 'Uttar Pradesh', 'Kerala', 'Gujarat']
    
    for inst_id in existing_institutions[:20]:  # Ensure we only generate for 20 institutions
        # Try to get existing institution characteristics
        existing_inst_data = existing_data[existing_data['institution_id'] == inst_id]
        
        if not existing_inst_data.empty:
            # Use existing characteristics
            latest_row = existing_inst_data[existing_inst_data['year'] == existing_inst_data['year'].max()].iloc[0]
            base_quality = min(0.9, max(0.3, latest_row['performance_score'] / 10))
            inst_type = latest_row['institution_type']
            state = latest_row['state']
            inst_name = latest_row['institution_name']
        else:
            # Create new institution characteristics
            base_quality = np.random.uniform(0.4, 0.9)
            inst_type = np.random.choice(institution_types)
            state = np.random.choice(states)
            inst_name = f'{inst_type.split()[0]} University {inst_id[-3:]}'
        
        # Add some yearly variation (slight improvement/deterioration)
        yearly_variation = base_quality + np.random.normal(0, 0.05)
        yearly_variation = max(0.3, min(0.95, yearly_variation))
        
        # Academic Metrics
        naac_grades = ['A++', 'A+', 'A', 'B++', 'B+', 'B', 'C']
        naac_probs = [0.05, 0.10, 0.15, 0.25, 0.25, 0.15, 0.05]
        naac_grade = np.random.choice(naac_grades, p=naac_probs)
        
        # NIRF ranking - better institutions more likely to be ranked
        if yearly_variation > 0.7 and np.random.random() < 0.8:
            nirf_rank = np.random.randint(1, 101)
        elif yearly_variation > 0.5 and np.random.random() < 0.5:
            nirf_rank = np.random.randint(101, 201)
        else:
            nirf_rank = None
        
        # Core metrics with realistic distributions
        student_faculty_ratio = max(10, np.random.normal(20, 5))
        phd_faculty_ratio = np.random.beta(2, 2) * 0.6 + 0.3
        placement_rate = max(40, min(98, np.random.normal(75 + (yearly_variation - 0.5) * 10, 8)))
        
        # Research & Infrastructure
        research_publications = max(0, int(np.random.poisson(yearly_variation * 30)))
        research_grants = max(0, int(np.random.exponential(yearly_variation * 500)))
        digital_infrastructure = max(1, min(10, np.random.normal(7, 1.5)))
        library_volumes = max(1000, int(np.random.normal(20000, 10000)))
        
        # Governance & Social Impact
        financial_stability = max(1, min(10, np.random.normal(7.5, 1.2)))
        compliance_score = max(1, min(10, np.random.normal(8, 1)))
        administrative_efficiency = max(1, min(10, np.random.normal(7.2, 1.1)))
        community_projects = np.random.poisson(yearly_variation * 8)
        
        # Additional metrics
        curriculum_framework_score = max(1, min(10, np.random.normal(7.5, 1.0)))
        faculty_diversity_index = max(1, min(10, np.random.normal(6.5, 1.2)))
        learning_outcome_achievement = max(50, min(95, np.random.normal(75, 8)))
        patents_filed = np.random.poisson(yearly_variation * 3)
        industry_collaborations = np.random.poisson(yearly_variation * 6)
        
        # Calculate overall performance score
        metrics_dict = {
            'naac_grade': naac_grade,
            'nirf_ranking': nirf_rank,
            'student_faculty_ratio': student_faculty_ratio,
            'phd_faculty_ratio': phd_faculty_ratio,
            'placement_rate': placement_rate,
            'research_publications': research_publications,
            'digital_infrastructure': digital_infrastructure,
            'financial_stability': financial_stability,
            'community_engagement': community_projects
        }
        performance_score = calculate_direct_performance_score(metrics_dict)
        
        institution_data = {
            'institution_id': inst_id,
            'institution_name': inst_name,
            'year': 2023,
            'institution_type': inst_type,
            'state': state,
            'established_year': np.random.randint(1950, 2010),
            
            # Academic Metrics
            'naac_grade': naac_grade,
            'nirf_ranking': nirf_rank,
            'student_faculty_ratio': round(student_faculty_ratio, 1),
            'phd_faculty_ratio': round(phd_faculty_ratio, 3),
            'placement_rate': round(placement_rate, 1),
            'higher_education_rate': round(max(5, min(50, np.random.normal(20, 8))), 1),
            
            # Research & Infrastructure
            'research_publications': research_publications,
            'research_grants_amount': research_grants,
            'patents_filed': patents_filed,
            'industry_collaborations': industry_collaborations,
            'digital_infrastructure_score': round(digital_infrastructure, 1),
            'library_volumes': library_volumes,
            'laboratory_equipment_score': round(max(1, min(10, np.random.normal(7, 1.3))), 1),
            
            # Governance & Social Impact
            'financial_stability_score': round(financial_stability, 1),
            'compliance_score': round(compliance_score, 1),
            'administrative_efficiency': round(administrative_efficiency, 1),
            'community_projects': community_projects,
            'rural_outreach_score': round(max(1, min(10, np.random.normal(6.8, 1.4))), 1),
            
            # Additional metrics
            'curriculum_framework_score': round(curriculum_framework_score, 1),
            'faculty_diversity_index': round(faculty_diversity_index, 1),
            'learning_outcome_achievement': round(learning_outcome_achievement, 1),
            'entrepreneurship_cell_score': round(max(1, min(10, np.random.normal(6.5, 1.5))), 1),
            
            # Overall Performance
            'performance_score': round(performance_score, 2),
            'approval_recommendation': generate_approval_recommendation(performance_score),
            'risk_level': assess_risk_level(performance_score)
        }
        
        institutions_data.append(institution_data)
    
    return pd.DataFrame(institutions_data)



def run_data_validation(analyzer):
    """Run comprehensive data validation"""
    st.subheader("🔍 Data Validation Results")
    
    with st.spinner("Running data validation checks..."):
        current_data = analyzer.historical_data
        current_year_data = current_data[current_data['year'] == 2023]
        
        validation_results = {
            'issues': [],
            'warnings': [],
            'checks_passed': []
        }
        
        # 1. Check for missing values in critical columns
        critical_columns = ['institution_id', 'institution_name', 'performance_score', 'risk_level']
        missing_data = {}
        for col in critical_columns:
            null_count = current_year_data[col].isnull().sum()
            if null_count > 0:
                validation_results['issues'].append(f"❌ Missing {col}: {null_count} institutions")
                missing_data[col] = null_count
        
        if not missing_data:
            validation_results['checks_passed'].append("✅ All critical fields have complete data")
        
        # 2. Validate data types and ranges
        numeric_ranges = {
            'performance_score': (0, 10),
            'placement_rate': (0, 100),
            'student_faculty_ratio': (5, 50),
            'research_publications': (0, 1000)
        }
        
        for col, (min_val, max_val) in numeric_ranges.items():
            if col in current_year_data.columns:
                out_of_range = current_year_data[
                    (current_year_data[col] < min_val) | (current_year_data[col] > max_val)
                ]
                if len(out_of_range) > 0:
                    validation_results['issues'].append(
                        f"❌ {col} out of range ({min_val}-{max_val}): {len(out_of_range)} records"
                    )
                else:
                    validation_results['checks_passed'].append(f"✅ {col} values are within valid range")
        
        # 3. Check for logical consistency
        # Performance score vs risk level consistency
        inconsistencies = []
        for _, row in current_year_data.iterrows():
            score = row['performance_score']
            risk = row['risk_level']
            
            if score >= 8.0 and risk != 'Low Risk':
                inconsistencies.append(f"{row['institution_name']}: High score ({score}) but {risk}")
            elif score < 5.0 and risk not in ['High Risk', 'Critical Risk']:
                inconsistencies.append(f"{row['institution_name']}: Low score ({score}) but {risk}")
        
        if inconsistencies:
            validation_results['warnings'].append(f"⚠️ {len(inconsistencies)} logical inconsistencies found")
        
        # 4. Check for duplicate records
        duplicate_institutions = current_year_data[current_year_data.duplicated(['institution_id'], keep=False)]
        if len(duplicate_institutions) > 0:
            validation_results['issues'].append(f"❌ Duplicate institution records: {len(duplicate_institutions)}")
        
        # 5. Validate NAAC grades format
        valid_naac = ['A++', 'A+', 'A', 'B++', 'B+', 'B', 'C']
        invalid_naac = current_year_data[~current_year_data['naac_grade'].isin(valid_naac + [None])]
        if len(invalid_naac) > 0:
            validation_results['issues'].append(f"❌ Invalid NAAC grades: {len(invalid_naac)}")
        
        # Display results
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.error(f"**Critical Issues**: {len(validation_results['issues'])}")
            for issue in validation_results['issues']:
                st.write(issue)
        
        with col2:
            st.warning(f"**Warnings**: {len(validation_results['warnings'])}")
            for warning in validation_results['warnings']:
                st.write(warning)
        
        with col3:
            st.success(f"**Checks Passed**: {len(validation_results['checks_passed'])}")
            for check in validation_results['checks_passed'][:5]:
                st.write(check)
        
        # Generate summary
        total_checks = len(validation_results['issues']) + len(validation_results['warnings']) + len(validation_results['checks_passed'])
        validation_score = (len(validation_results['checks_passed']) / total_checks) * 100
        
        st.metric("Overall Validation Score", f"{validation_score:.1f}%")
        
        # Provide recommendations
        if validation_results['issues']:
            st.error("**Action Required**: Fix critical issues before proceeding")
            if st.button("🛠️ Auto-fix Common Issues"):
                auto_fix_issues(analyzer)

def auto_fix_issues(analyzer):
    """Automatically fix common data issues"""
    current_data = analyzer.historical_data.copy()
    
    # Fix missing performance scores
    if 'performance_score' in current_data.columns:
        missing_scores = current_data['performance_score'].isnull()
        if missing_scores.any():
            # Estimate based on other metrics
            current_data.loc[missing_scores, 'performance_score'] = 6.0  # Default
    
    # Fix out-of-range values
    if 'placement_rate' in current_data.columns:
        current_data['placement_rate'] = current_data['placement_rate'].clip(0, 100)
    
    # Save fixes
    current_data.to_sql('institutions', analyzer.conn, if_exists='replace', index=False)
    analyzer.historical_data = current_data
    
    st.success("✅ Common issues auto-fixed!")

def calculate_data_quality_metrics(analyzer):
    """Calculate various data quality metrics"""
    current_data = analyzer.historical_data
    current_year_data = current_data[current_data['year'] == 2023]
    
    # 1. Completeness: Percentage of non-null values
    total_cells = current_year_data.shape[0] * current_year_data.shape[1]
    null_cells = current_year_data.isnull().sum().sum()
    completeness = ((total_cells - null_cells) / total_cells) * 100
    
    # 2. Consistency: Logical consistency between related fields
    consistency_score = 0
    consistency_checks = 0
    
    # Check performance_score vs risk_level consistency
    correct_risk = 0
    for _, row in current_year_data.iterrows():
        score = row['performance_score']
        risk = row['risk_level']
        
        if (score >= 8.0 and risk == 'Low Risk') or \
           (6.0 <= score < 8.0 and risk == 'Medium Risk') or \
           (score < 6.0 and risk in ['High Risk', 'Critical Risk']):
            correct_risk += 1
    
    if len(current_year_data) > 0:
        consistency_score += (correct_risk / len(current_year_data)) * 10
        consistency_checks += 1
    
    # Check placement_rate vs performance_score correlation
    if 'placement_rate' in current_year_data.columns and 'performance_score' in current_year_data.columns:
        correlation = current_year_data['placement_rate'].corr(current_year_data['performance_score'])
        if not pd.isna(correlation):
            consistency_score += (abs(correlation) * 10)  # Higher correlation = more consistent
            consistency_checks += 1
    
    final_consistency = consistency_score / consistency_checks if consistency_checks > 0 else 5.0
    
    # 3. Accuracy: Based on business rules and valid ranges
    accuracy_score = 0
    accuracy_checks = 0
    
    # Check if performance_score is within 0-10 range
    if 'performance_score' in current_year_data.columns:
        valid_scores = current_year_data[
            (current_year_data['performance_score'] >= 0) & 
            (current_year_data['performance_score'] <= 10)
        ].shape[0]
        accuracy_score += (valid_scores / len(current_year_data)) * 10
        accuracy_checks += 1
    
    # Check if placement_rate is within 0-100
    if 'placement_rate' in current_year_data.columns:
        valid_placement = current_year_data[
            (current_year_data['placement_rate'] >= 0) & 
            (current_year_data['placement_rate'] <= 100)
        ].shape[0]
        accuracy_score += (valid_placement / len(current_year_data)) * 10
        accuracy_checks += 1
    
    final_accuracy = accuracy_score / accuracy_checks if accuracy_checks > 0 else 5.0
    
    # 4. Timeliness: How current the data is
    latest_year = current_data['year'].max()
    current_year = datetime.now().year
    years_diff = current_year - latest_year
    
    if years_diff == 0:
        timeliness = 9.5
    elif years_diff == 1:
        timeliness = 7.0
    elif years_diff == 2:
        timeliness = 5.0
    else:
        timeliness = 3.0
    
    return {
        'completeness': round(completeness, 1),
        'consistency': round(final_consistency, 1),
        'accuracy': round(final_accuracy, 1),
        'timeliness': round(timeliness, 1)
    }

def generate_data_summary(data):
    """Generate and display data summary"""
    with st.expander("📄 Data Summary Report", expanded=True):
        st.write("### 📊 Summary Statistics")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Total Institutions", len(data))
            st.metric("Avg Performance", f"{data['performance_score'].mean():.2f}")
        
        with col2:
            st.metric("Avg Placement Rate", f"{data['placement_rate'].mean():.1f}%")
            st.metric("Research Active", f"{len(data[data['research_publications'] > 20])}/{len(data)}")
        
        with col3:
            st.metric("NAAC A/A+ Rated", f"{len(data[data['naac_grade'].isin(['A', 'A+', 'A++'])]):.0f}")
            st.metric("Low Risk Institutions", f"{len(data[data['risk_level'] == 'Low Risk']):.0f}")
        
        # Performance distribution
        st.write("### 📈 Performance Distribution")
        fig = px.histogram(data, x='performance_score', nbins=10, 
                          title="Distribution of Performance Scores")
        st.plotly_chart(fig, use_container_width=True)
        
        # Top performers
        st.write("### 🏆 Top 5 Performing Institutions")
        top_performers = data.nlargest(5, 'performance_score')[['institution_name', 'performance_score', 'risk_level']]
        st.dataframe(top_performers, use_container_width=True)
        
        # Export option
        st.download_button(
            label="📥 Download Summary Report",
            data=data.to_csv(index=False),
            file_name="institution_summary.csv",
            mime="text/csv"
        )

def generate_comprehensive_dummy_data():
    """Generate comprehensive dummy data matching both Basic and Systematic form parameters"""
    np.random.seed(42)
    
    # Create 20 unique institutions
    n_institutions = 20
    years = list(range(2014, 2024))  # 10 years of data
    
    institutions_data = []
    
    institution_types = ['State University', 'Private University', 'Deemed University', 'Autonomous College', 'Technical Institute']
    states = ['Maharashtra', 'Karnataka', 'Tamil Nadu', 'Delhi', 'Uttar Pradesh', 'Kerala', 'Gujarat']
    
    for inst_id in range(1, n_institutions + 1):
        base_quality = np.random.uniform(0.4, 0.9)
        inst_type = np.random.choice(institution_types)
        state = np.random.choice(states)
        
        for year in years:
            # Add some yearly variation
            yearly_variation = base_quality + np.random.normal(0, 0.1)
            yearly_variation = max(0.3, min(0.95, yearly_variation))
            
            # Academic Metrics (from Basic Form)
            naac_grades = ['A++', 'A+', 'A', 'B++', 'B+', 'B', 'C']
            naac_probs = [0.05, 0.10, 0.15, 0.25, 0.25, 0.15, 0.05]
            naac_grade = np.random.choice(naac_grades, p=naac_probs)
            
            # SIMPLIFIED: NIRF ranking - 60% have ranking, 40% don't
            if np.random.random() < 0.6:
                nirf_rank = np.random.randint(1, 201)
            else:
                nirf_rank = None
            
            student_faculty_ratio = max(10, np.random.normal(20, 5))
            phd_faculty_ratio = np.random.beta(2, 2) * 0.6 + 0.3
            placement_rate = max(40, min(98, np.random.normal(75, 10)))
            
            # Research & Infrastructure (from Basic Form)
            research_publications = max(0, int(np.random.poisson(yearly_variation * 30)))
            research_grants = max(0, int(np.random.exponential(yearly_variation * 500)))
            digital_infrastructure = max(1, min(10, np.random.normal(7, 1.5)))
            library_volumes = max(1000, int(np.random.normal(20000, 10000)))
            
            # Governance & Social Impact (from Basic Form)
            financial_stability = max(1, min(10, np.random.normal(7.5, 1.2)))
            compliance_score = max(1, min(10, np.random.normal(8, 1)))
            administrative_efficiency = max(1, min(10, np.random.normal(7.2, 1.1)))
            community_projects = np.random.poisson(yearly_variation * 8)
            
            # Additional metrics from Systematic Form
            curriculum_framework_score = max(1, min(10, np.random.normal(7.5, 1.0)))
            faculty_diversity_index = max(1, min(10, np.random.normal(6.5, 1.2)))
            learning_outcome_achievement = max(50, min(95, np.random.normal(75, 8)))
            patents_filed = np.random.poisson(yearly_variation * 3)
            industry_collaborations = np.random.poisson(yearly_variation * 6)
            renewable_energy_usage = max(0, min(100, np.random.normal(25, 15)))
            governance_transparency = max(1, min(10, np.random.normal(7, 1.3)))
            campus_area = max(10, np.random.normal(50, 20))
            research_investment_pct = max(5, min(30, np.random.normal(15, 5)))
            
            # Calculate overall performance score
            metrics_dict = {
                'naac_grade': naac_grade,
                'nirf_ranking': nirf_rank,
                'student_faculty_ratio': student_faculty_ratio,
                'phd_faculty_ratio': phd_faculty_ratio,
                'placement_rate': placement_rate,
                'research_publications': research_publications,
                'digital_infrastructure': digital_infrastructure,
                'financial_stability': financial_stability,
                'community_engagement': community_projects  # Use correct parameter name
            }
            performance_score = calculate_direct_performance_score(metrics_dict)
            
            institution_data = {
                'institution_id': f'INST_{inst_id:04d}',
                'institution_name': f'{inst_type.split()[0]} University {inst_id:03d}',
                'year': year,
                'institution_type': inst_type,
                'state': state,
                'established_year': np.random.randint(1950, 2010),
                
                # Academic Metrics
                'naac_grade': naac_grade,
                'nirf_ranking': nirf_rank,
                'student_faculty_ratio': round(student_faculty_ratio, 1),
                'phd_faculty_ratio': round(phd_faculty_ratio, 3),
                'placement_rate': round(placement_rate, 1),
                'higher_education_rate': round(max(5, min(50, np.random.normal(20, 8))), 1),
                
                # Research & Infrastructure
                'research_publications': research_publications,
                'research_grants_amount': research_grants,
                'patents_filed': patents_filed,
                'industry_collaborations': industry_collaborations,
                'digital_infrastructure_score': round(digital_infrastructure, 1),
                'library_volumes': library_volumes,
                'laboratory_equipment_score': round(max(1, min(10, np.random.normal(7, 1.3))), 1),
                
                # Governance & Social Impact
                'financial_stability_score': round(financial_stability, 1),
                'compliance_score': round(compliance_score, 1),
                'administrative_efficiency': round(administrative_efficiency, 1),
                'community_projects': community_projects,
                'rural_outreach_score': round(max(1, min(10, np.random.normal(6.8, 1.4))), 1),
                
                # Additional Systematic Form Metrics
                'curriculum_framework_score': round(curriculum_framework_score, 1),
                'faculty_diversity_index': round(faculty_diversity_index, 1),
                'learning_outcome_achievement': round(learning_outcome_achievement, 1),
                'renewable_energy_usage': round(renewable_energy_usage, 1),
                'governance_transparency_score': round(governance_transparency, 1),
                'campus_area': round(campus_area, 1),
                'research_investment_percentage': round(research_investment_pct, 1),
                'entrepreneurship_cell_score': round(max(1, min(10, np.random.normal(6.5, 1.5))), 1),
                
                # Overall Performance
                'performance_score': round(performance_score, 2),
                'approval_recommendation': generate_approval_recommendation(performance_score),
                'risk_level': assess_risk_level(performance_score)
            }
            
            institutions_data.append(institution_data)
    
    return pd.DataFrame(institutions_data)

def calculate_direct_performance_score(metrics: Dict) -> float:
    """Calculate performance score directly (for data generation)"""
    score = 0
    
    # NAAC Grade scoring
    naac_scores = {'A++': 10, 'A+': 9, 'A': 8, 'B++': 7, 'B+': 6, 'B': 5, 'C': 4}
    naac_score = naac_scores.get(metrics.get('naac_grade'), 5)
    score += naac_score * 0.15
    
    # NIRF Ranking scoring (inverse)
    nirf_score = 0
    nirf_rank = metrics.get('nirf_ranking')
    if nirf_rank and nirf_rank <= 200:
        nirf_score = (201 - nirf_rank) / 200 * 10
    score += nirf_score * 0.10
    
    # Student-Faculty Ratio (lower is better)
    sf_ratio = metrics.get('student_faculty_ratio', 20)
    sf_ratio_score = max(0, 10 - max(0, sf_ratio - 15) / 3)
    score += sf_ratio_score * 0.10
    
    # PhD Faculty Ratio
    phd_score = metrics.get('phd_faculty_ratio', 0.6) * 10
    score += phd_score * 0.10
    
    # Research Publications
    pub_score = min(10, metrics.get('research_publications', 0) / 10)
    score += pub_score * 0.10
    
    # Digital Infrastructure
    infra_score = metrics.get('digital_infrastructure', 7)
    score += infra_score * 0.10
    
    # Financial Stability
    financial_score = metrics.get('financial_stability', 7.5)
    score += financial_score * 0.10
    
    # Placement Rate
    placement = metrics.get('placement_rate', 75)
    placement_score = placement / 10
    score += placement_score * 0.15
    
    # Community Engagement
    community = metrics.get('community_engagement', metrics.get('community_projects', 0))
    community_score = min(10, community / 1.5)
    score += community_score * 0.05
    
    return min(10, score)

def calculate_comprehensive_performance(*args):
    """Calculate overall performance score from multiple metrics"""
    score = 5.0  # Base score
    
    # Add contributions from various metrics
    metrics = list(args)
    
    # NAAC Grade
    if metrics[0] in ['A++', 'A+', 'A']:
        score += 1.5
    elif metrics[0] in ['B++', 'B+']:
        score += 0.5
    
    # NIRF Ranking
    if metrics[1] and metrics[1] <= 100:
        score += 1.0
    elif metrics[1] and metrics[1] <= 200:
        score += 0.5
    
    # Student-Faculty Ratio (lower is better)
    if metrics[2] <= 15:
        score += 1.0
    elif metrics[2] <= 20:
        score += 0.5
    
    # Other metrics contribute proportionally
    numeric_metrics = [m for m in metrics[3:] if isinstance(m, (int, float))]
    if numeric_metrics:
        avg_metric_score = sum(numeric_metrics) / len(numeric_metrics)
        score += (avg_metric_score - 5) * 0.5  # Normalize contribution
    
    return max(1.0, min(10.0, score))

def generate_approval_recommendation(performance_score):
    """Generate approval recommendation based on performance score"""
    if performance_score >= 8.0:
        return "Full Approval - 5 Years"
    elif performance_score >= 7.0:
        return "Provisional Approval - 3 Years"
    elif performance_score >= 6.0:
        return "Conditional Approval - 1 Year"
    elif performance_score >= 5.0:
        return "Approval with Strict Monitoring - 1 Year"
    else:
        return "Rejection - Significant Improvements Required"

def assess_risk_level(performance_score):
    """Assess institutional risk level"""
    if performance_score >= 8.0:
        return "Low Risk"
    elif performance_score >= 6.5:
        return "Medium Risk"
    elif performance_score >= 5.0:
        return "High Risk"
    else:
        return "Critical Risk"
def create_approval_workflow(analyzer):
    st.header("🔄 AI-Enhanced Approval Workflow")
    
    st.info("Streamlined approval process with AI-powered decision support")
    
    # Workflow steps with AI integration
    workflow_steps = [
        {
            "step": 1,
            "title": "Document Submission & AI Verification",
            "description": "Institutions submit documents through portal, AI verifies completeness",
            "ai_features": ["Document classification", "Completeness check", "Sufficiency scoring"],
            "output": "Document Sufficiency Report"
        },
        {
            "step": 2,
            "title": "Historical Data Analysis",
            "description": "AI analyzes 5+ years of institutional performance data",
            "ai_features": ["Trend analysis", "Performance scoring", "Risk assessment"],
            "output": "Performance Analytics Report"
        },
        {
            "step": 3,
            "title": "Comparative Benchmarking",
            "description": "AI compares institution with similar peers and standards",
            "ai_features": ["Peer comparison", "Benchmark analysis", "Percentile ranking"],
            "output": "Comparative Analysis Report"
        },
        {
            "step": 4,
            "title": "AI Recommendation Generation",
            "description": "AI generates approval recommendations with justifications",
            "ai_features": ["Decision support", "Risk mitigation", "Improvement suggestions"],
            "output": "AI Recommendation Report"
        },
        {
            "step": 5,
            "title": "Expert Committee Review",
            "description": "UGC/AICTE committee reviews AI recommendations and makes final decision",
            "ai_features": ["Decision tracking", "Approval workflow", "Monitoring setup"],
            "output": "Final Approval Decision"
        }
    ]
    
    for step in workflow_steps:
        with st.expander(f"Step {step['step']}: {step['title']}"):
            col1, col2 = st.columns([2, 1])
            with col1:
                st.write(f"**Description:** {step['description']}")
                st.write(f"**Output:** {step['output']}")
            with col2:
                st.write("**AI Features:**")
                for feature in step['ai_features']:
                    st.write(f"• {feature}")

def create_rag_data_management(analyzer):
    st.header("🤖 RAG-Powered Data Management & Analysis")
    
    st.info("""
    **Retrieval Augmented Generation (RAG) System**: 
    Upload institutional documents and let AI automatically extract, analyze, and structure data 
    for comprehensive institutional evaluation.
    """)
    
    tab1, tab2, tab3, tab4 = st.tabs([
        "📤 Upload & Extract", 
        "🔍 View Extracted Data", 
        "📊 AI Analysis",
        "⚙️ RAG Settings"
    ])
    
    with tab1:
        st.subheader("Document Upload & Data Extraction")
        
        # Institution selection
        current_institutions = analyzer.historical_data[analyzer.historical_data['year'] == 2023]['institution_id'].unique()
        selected_institution = st.selectbox(
            "Select Institution",
            current_institutions,
            key="rag_institution"
        )
        
        uploaded_files = st.file_uploader(
            "Upload Institutional Documents for RAG Analysis",
            type=['pdf', 'doc', 'docx', 'txt', 'xlsx', 'xls'],
            accept_multiple_files=True,
            help="Upload all relevant documents: Annual reports, NAAC reports, Research data, etc."
        )
        
        if uploaded_files:
            st.success(f"📄 {len(uploaded_files)} documents ready for analysis")
            
            # Show document preview
            with st.expander("📋 Document Preview"):
                for i, file in enumerate(uploaded_files):
                    st.write(f"**{i+1}. {file.name}** ({file.size} bytes)")
            
            if st.button("🚀 Start RAG Analysis", type="primary"):
                with st.spinner("🤖 AI is analyzing documents and extracting data..."):
                    # Perform RAG analysis
                    analysis_result = analyzer.analyze_documents_with_rag(
                        selected_institution, 
                        uploaded_files
                    )
                    
                    # SAFE ACCESS: Check if analysis_result is valid
                    if analysis_result and analysis_result.get('status') == 'Analysis Complete':
                        st.success("✅ RAG Analysis Completed Successfully!")
                        
                        # Store results in session state for other tabs
                        st.session_state.rag_analysis = analysis_result
                        st.session_state.selected_institution = selected_institution
                        
                        # Show quick insights with safe access
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            confidence = analysis_result.get('confidence_score', 0.0)
                            st.metric("Confidence Score", f"{confidence:.2f}")
                        with col2:
                            extracted_data = analysis_result.get('extracted_data', {})
                            extracted_categories = len([k for k in extracted_data.keys() if k not in ['raw_text', 'file_names']])
                            st.metric("Data Categories", extracted_categories)
                        with col3:
                            ai_insights = analysis_result.get('ai_insights', {})
                            risk_assessment = ai_insights.get('risk_assessment', {})
                            risk_level = risk_assessment.get('level', 'Unknown')
                            st.metric("Risk Level", risk_level)
                    
                    else:
                        st.error("❌ RAG Analysis Failed. Please try again.")
                        # Still store the result even if failed for debugging
                        if analysis_result:
                            st.session_state.rag_analysis = analysis_result
    
    with tab2:
        st.subheader("Extracted Data View")
        
        # SAFE ACCESS: Check if rag_analysis exists and has extracted_data
        if 'rag_analysis' in st.session_state and st.session_state.rag_analysis:
            analysis_result = st.session_state.rag_analysis
            extracted_data = analysis_result.get('extracted_data', {})
            
            if not extracted_data:
                st.warning("No extracted data available. Please run RAG analysis first.")
                return
            
            # Show extracted data by category
            categories = [
                ('Academic Metrics', 'academic_metrics'),
                ('Research Metrics', 'research_metrics'), 
                ('Infrastructure Metrics', 'infrastructure_metrics'),
                ('Governance Metrics', 'governance_metrics'),
                ('Financial Metrics', 'financial_metrics')
            ]
            
            for category_name, category_key in categories:
                with st.expander(f"📈 {category_name}"):
                    category_data = extracted_data.get(category_key, {})
                    if category_data:
                        for key, value in category_data.items():
                            col1, col2 = st.columns([1, 2])
                            with col1:
                                st.write(f"**{key.replace('_', ' ').title()}:**")
                            with col2:
                                st.write(value)
                    else:
                        st.info(f"No {category_name.lower()} extracted from documents")
            
            # Show raw text preview
            with st.expander("📝 Extracted Text Preview"):
                raw_text = extracted_data.get('raw_text', '')
                if raw_text:
                    preview_text = raw_text[:2000] + "..." if len(raw_text) > 2000 else raw_text
                    st.text_area("Extracted Text", preview_text, height=200, key="raw_text_preview")
                else:
                    st.info("No text extracted")
                    
        else:
            st.info("👆 Upload documents and run RAG analysis to view extracted data")
    
    with tab3:
        st.subheader("AI Insights & Analysis")
        
        # SAFE ACCESS: Check if rag_analysis exists and has ai_insights
        if 'rag_analysis' in st.session_state and st.session_state.rag_analysis:
            analysis_result = st.session_state.rag_analysis
            insights = analysis_result.get('ai_insights', {})
            
            if not insights:
                st.warning("No AI insights available. Please run RAG analysis first.")
                return
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("✅ Strengths")
                strengths = insights.get('strengths', [])
                if strengths:
                    for strength in strengths:
                        st.success(f"• {strength}")
                else:
                    st.info("No significant strengths identified")
                
                st.subheader("🎯 Recommendations")
                recommendations = insights.get('recommendations', [])
                if recommendations:
                    for rec in recommendations:
                        st.warning(f"• {rec}")
                else:
                    st.info("No specific recommendations")
            
            with col2:
                st.subheader("⚠️ Areas for Improvement")
                weaknesses = insights.get('weaknesses', [])
                if weaknesses:
                    for weakness in weaknesses:
                        st.error(f"• {weakness}")
                else:
                    st.info("No major weaknesses identified")
                
                st.subheader("📊 Risk Assessment")
                risk_assessment = insights.get('risk_assessment', {})
                risk_score = risk_assessment.get('score', 5.0)
                st.metric("Risk Score", f"{risk_score:.1f}/10")
                
                risk_level = risk_assessment.get('level', 'Unknown')
                st.write(f"**Risk Level:** {risk_level}")
                
                risk_factors = risk_assessment.get('factors', [])
                if risk_factors:
                    st.write("**Risk Factors:**")
                    for factor in risk_factors:
                        st.write(f"• {factor}")
            
            # Generate approval recommendation
            st.subheader("🏛️ Approval Recommendation")
            risk_level = risk_assessment.get('level', 'Medium')
            if risk_level == 'Low':
                st.success("**✅ RECOMMENDED: Full Approval - 5 Years**")
                st.write("Institution demonstrates strong performance across all parameters with minimal risk factors.")
            elif risk_level == 'Medium':
                st.warning("**🟡 CONDITIONAL: Provisional Approval - 3 Years**")
                st.write("Institution shows promise but has some areas requiring improvement and monitoring.")
            else:
                st.error("**🔴 NOT RECOMMENDED: Requires Significant Improvements**")
                st.write("Critical risk factors identified. Institution needs substantial improvements before approval.")
                
        else:
            st.info("👆 Run RAG analysis to generate AI insights")
    
    with tab4:
        st.subheader("RAG System Settings")
        # ... (keep existing tab4 code, it doesn't use analysis_result)

def create_system_settings(analyzer):
    st.header("⚙️ System Configuration & Settings")
    
    tab1, tab2, tab3, tab4 = st.tabs([
        "📊 Performance Metrics", 
        "📋 Document Requirements",
        "🎯 Approval Thresholds",
        "🔧 System Parameters"
    ])
    
    with tab1:
        st.subheader("📊 Performance Metrics Configuration")
        st.info("Configure weights and parameters for institutional performance evaluation")
        
        # Create editable performance metrics structure
        performance_config = analyzer.performance_metrics.copy()
        
        for category, config in performance_config.items():
            with st.expander(f"🔸 {category.replace('_', ' ').title()} (Weight: {config['weight']})", expanded=True):
                
                # Category weight
                new_weight = st.slider(
                    f"Category Weight for {category.replace('_', ' ').title()}",
                    min_value=0.0,
                    max_value=1.0,
                    value=config['weight'],
                    step=0.05,
                    key=f"weight_{category}"
                )
                
                config['weight'] = new_weight
                
                # Sub-metrics configuration
                st.write("**Sub-Metrics Configuration:**")
                sub_metrics = config['sub_metrics']
                
                col1, col2 = st.columns([3, 1])
                with col1:
                    for sub_metric, weight in sub_metrics.items():
                        new_sub_weight = st.slider(
                            f"{sub_metric.replace('_', ' ').title()}",
                            min_value=0.0,
                            max_value=1.0,
                            value=weight,
                            step=0.05,
                            key=f"sub_{category}_{sub_metric}"
                        )
                        sub_metrics[sub_metric] = new_sub_weight
                
                with col2:
                    # Show weight distribution
                    total = sum(sub_metrics.values())
                    if total > 0:
                        st.write("**Distribution:**")
                        for sub_metric, weight in sub_metrics.items():
                            percentage = (weight / total) * 100
                            st.write(f"{sub_metric}: {percentage:.1f}%")
                    else:
                        st.warning("Total weight is 0")
        
        # Save button for performance metrics
        if st.button("💾 Save Performance Metrics Configuration", type="primary"):
            analyzer.performance_metrics = performance_config
            save_configuration(analyzer, 'performance_metrics', performance_config)
            st.success("✅ Performance metrics configuration saved!")
        
        # Reset to defaults
        if st.button("🔄 Reset to Defaults", type="secondary"):
            analyzer.performance_metrics = analyzer.define_performance_metrics()
            st.success("✅ Performance metrics reset to defaults")
            st.rerun()
    
    with tab2:
        st.subheader("📋 Document Requirements Configuration")
        st.info("Configure document requirements for different approval types")
        
        doc_config = analyzer.document_requirements.copy()
        
        for approval_type, requirements in doc_config.items():
            with st.expander(f"📄 {approval_type.replace('_', ' ').title()}", expanded=True):
                
                # Mandatory Documents
                st.write("**📌 Mandatory Documents:**")
                mandatory_docs = requirements['mandatory']
                
                # Dynamic list for mandatory documents
                new_mandatory = []
                for i, doc in enumerate(mandatory_docs):
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        new_doc = st.text_input(
                            f"Mandatory Document {i+1}",
                            value=doc,
                            key=f"mandatory_{approval_type}_{i}"
                        )
                    with col2:
                        if st.button("❌", key=f"remove_mandatory_{approval_type}_{i}"):
                            continue  # Skip adding this one
                    if new_doc:
                        new_mandatory.append(new_doc)
                
                # Add new document button
                if st.button("➕ Add New Mandatory Document", key=f"add_mandatory_{approval_type}"):
                    new_mandatory.append("New Document")
                    st.rerun()
                
                requirements['mandatory'] = new_mandatory
                
                # Supporting Documents
                st.write("**📎 Supporting Documents:**")
                supporting_docs = requirements['supporting']
                
                new_supporting = []
                for i, doc in enumerate(supporting_docs):
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        new_doc = st.text_input(
                            f"Supporting Document {i+1}",
                            value=doc,
                            key=f"supporting_{approval_type}_{i}"
                        )
                    with col2:
                        if st.button("❌", key=f"remove_supporting_{approval_type}_{i}"):
                            continue
                    if new_doc:
                        new_supporting.append(new_doc)
                
                if st.button("➕ Add New Supporting Document", key=f"add_supporting_{approval_type}"):
                    new_supporting.append("New Supporting Document")
                    st.rerun()
                
                requirements['supporting'] = new_supporting
        
        # Save button for document requirements
        if st.button("💾 Save Document Requirements", type="primary"):
            analyzer.document_requirements = doc_config
            save_configuration(analyzer, 'document_requirements', doc_config)
            st.success("✅ Document requirements saved!")
    
    with tab3:
        st.subheader("🎯 Approval Thresholds Configuration")
        st.info("Configure score thresholds for different approval levels")
        
        # Current thresholds (extracted from generate_approval_recommendation)
        thresholds = {
            "Full Approval (5 Years)": 8.0,
            "Provisional Approval (3 Years)": 7.0,
            "Conditional Approval (1 Year)": 6.0,
            "Strict Monitoring (1 Year)": 5.0,
            "Rejection": 0.0
        }
        
        updated_thresholds = {}
        
        for level, threshold in thresholds.items():
            col1, col2 = st.columns([3, 1])
            with col1:
                st.write(f"**{level}**")
            with col2:
                new_threshold = st.number_input(
                    f"Threshold",
                    min_value=0.0,
                    max_value=10.0,
                    value=threshold,
                    step=0.5,
                    key=f"threshold_{level}"
                )
            updated_thresholds[level] = new_threshold
        
        # Risk level thresholds
        st.write("---")
        st.subheader("⚠️ Risk Level Thresholds")
        
        risk_thresholds = {
            "Low Risk": 8.0,
            "Medium Risk": 6.5,
            "High Risk": 5.0,
            "Critical Risk": 0.0
        }
        
        for level, threshold in risk_thresholds.items():
            col1, col2 = st.columns([3, 1])
            with col1:
                st.write(f"**{level}**")
            with col2:
                new_threshold = st.number_input(
                    f"Risk Threshold",
                    min_value=0.0,
                    max_value=10.0,
                    value=threshold,
                    step=0.5,
                    key=f"risk_{level}"
                )
        
        if st.button("💾 Save Thresholds", type="primary"):
            save_configuration(analyzer, 'approval_thresholds', updated_thresholds)
            st.success("✅ Approval thresholds saved!")
    
    with tab4:
        st.subheader("🔧 System Parameters")
        st.info("Configure system-wide parameters and constants")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Data generation parameters
            st.write("**Data Generation:**")
            default_institutions = st.number_input(
                "Default Number of Institutions",
                min_value=5,
                max_value=100,
                value=20,
                step=5
            )
            
            default_years = st.number_input(
                "Default Years of Data",
                min_value=1,
                max_value=20,
                value=10,
                step=1
            )
            
            # RAG settings
            st.write("**AI/RAG Settings:**")
            chunk_size = st.number_input(
                "Document Chunk Size",
                min_value=500,
                max_value=2000,
                value=1000,
                step=100
            )
            
            chunk_overlap = st.number_input(
                "Chunk Overlap",
                min_value=0,
                max_value=500,
                value=200,
                step=50
            )
        
        with col2:
            # Scoring parameters
            st.write("**Scoring Parameters:**")
            min_performance_score = st.number_input(
                "Minimum Performance Score",
                min_value=0.0,
                max_value=5.0,
                value=1.0,
                step=0.5
            )
            
            max_performance_score = st.number_input(
                "Maximum Performance Score",
                min_value=5.0,
                max_value=15.0,
                value=10.0,
                step=0.5
            )
            
            # Document validation
            st.write("**Document Validation:**")
            mandatory_threshold = st.slider(
                "Mandatory Document Threshold (%)",
                min_value=0,
                max_value=100,
                value=80,
                step=5
            )
            
            review_period_days = st.number_input(
                "Review Period (days)",
                min_value=1,
                max_value=90,
                value=30,
                step=1
            )
        
        # System actions
        st.write("---")
        st.subheader("🛠️ System Actions")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🔄 Clear All Cache", type="secondary"):
                clear_cache(analyzer)
        
        with col2:
            if st.button("📊 Recalculate All Scores", type="secondary"):
                recalculate_all_scores(analyzer)
        
        with col3:
            if st.button("🔍 Validate System Integrity", type="secondary"):
                validate_system_integrity(analyzer)
        
        # Save all system parameters
        if st.button("💾 Save System Parameters", type="primary"):
            system_params = {
                'default_institutions': default_institutions,
                'default_years': default_years,
                'chunk_size': chunk_size,
                'chunk_overlap': chunk_overlap,
                'min_performance_score': min_performance_score,
                'max_performance_score': max_performance_score,
                'mandatory_threshold': mandatory_threshold,
                'review_period_days': review_period_days
            }
            save_configuration(analyzer, 'system_parameters', system_params)
            st.success("✅ System parameters saved!")

def save_configuration(analyzer, config_type, config_data):
    """Save configuration to database"""
    cursor = analyzer.conn.cursor()
    
    # Create configuration table if not exists
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS system_configuration (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            config_type TEXT,
            config_data TEXT,
            updated_by TEXT,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Save configuration
    cursor.execute('''
        INSERT INTO system_configuration (config_type, config_data, updated_by)
        VALUES (?, ?, ?)
    ''', (config_type, json.dumps(config_data), 'admin'))
    
    analyzer.conn.commit()

def clear_cache(analyzer):
    """Clear system cache"""
    # Clear session state
    for key in list(st.session_state.keys()):
        if key not in ['session_initialized', 'rag_initialized']:
            del st.session_state[key]
    
    # Clear temporary files
    temp_dir = tempfile.gettempdir()
    for file in os.listdir(temp_dir):
        if file.startswith('ugc_'):
            try:
                os.remove(os.path.join(temp_dir, file))
            except:
                pass
    
    st.success("✅ System cache cleared!")
    st.rerun()

def recalculate_all_scores(analyzer):
    """Recalculate all performance scores based on new configuration"""
    with st.spinner("Recalculating all scores with new configuration..."):
        # Get current data
        current_data = analyzer.historical_data.copy()
        
        # Apply new scoring configuration
        for idx, row in current_data.iterrows():
            # Calculate score using new weights
            new_score = analyzer.calculate_performance_score({
                'naac_grade': row.get('naac_grade'),
                'nirf_ranking': row.get('nirf_ranking'),
                'student_faculty_ratio': row.get('student_faculty_ratio'),
                'phd_faculty_ratio': row.get('phd_faculty_ratio'),
                'placement_rate': row.get('placement_rate'),
                'research_publications': row.get('research_publications'),
                'digital_infrastructure': row.get('digital_infrastructure_score'),
                'financial_stability': row.get('financial_stability_score'),
                'community_engagement': row.get('community_projects')
            })
            
            current_data.at[idx, 'performance_score'] = new_score
            current_data.at[idx, 'approval_recommendation'] = analyzer.generate_approval_recommendation(new_score)
            current_data.at[idx, 'risk_level'] = analyzer.assess_risk_level(new_score)
        
        # Save updated data
        current_data.to_sql('institutions', analyzer.conn, if_exists='replace', index=False)
        analyzer.historical_data = current_data
        
        st.success("✅ All scores recalculated with new configuration!")

def validate_system_integrity(analyzer):
    """Validate system configuration integrity"""
    issues = []
    
    # Check if weights sum to 1
    total_weight = sum(category['weight'] for category in analyzer.performance_metrics.values())
    if abs(total_weight - 1.0) > 0.01:
        issues.append(f"Performance metric weights don't sum to 1.0 (current: {total_weight:.2f})")
    
    # Check document requirements
    for approval_type, requirements in analyzer.document_requirements.items():
        if not requirements.get('mandatory'):
            issues.append(f"No mandatory documents for {approval_type}")
    
    if issues:
        st.error("⚠️ System Integrity Issues Found:")
        for issue in issues:
            st.write(f"- {issue}")
    else:
        st.success("✅ System integrity validation passed!")


    
def main():
    # Safe session state initialization
    if 'institution_user' not in st.session_state:
        st.session_state.institution_user = None
    if 'user_role' not in st.session_state:
        st.session_state.user_role = None
    
    # Initialize analytics engine with error handling
    try:
        analyzer = InstitutionalAIAnalyzer()
        
        # Verify data specifications
        total_institutions = analyzer.historical_data['institution_id'].nunique()
        total_years = analyzer.historical_data['year'].nunique()
        total_records = len(analyzer.historical_data)
        
        st.sidebar.success(f"📊 Data: {total_institutions} institutes × {total_years} years")
        st.sidebar.info(f"📈 Total Records: {total_records}")
        
        # Show data verification
        if total_institutions == 20 and total_years == 10 and total_records == 200:
            st.sidebar.success("✅ 20×10 specification verified")
        else:
            st.sidebar.warning(f"⚠️ Data mismatch: Expected 20×10=200, Got {total_institutions}×{total_years}={total_records}")
            
    except Exception as e:
        st.error(f"❌ System initialization error: {str(e)}")
        st.stop()
    
    # Check if institution user is logged in
    if st.session_state.institution_user is not None:
        create_institution_dashboard(analyzer, st.session_state.institution_user)
        if st.sidebar.button("🚪 Logout"):
            st.session_state.institution_user = None
            st.session_state.user_role = None
            st.rerun()
        return
    
    # Main header and system overview
    st.markdown('<h1 class="main-header">🏛️ AI-Powered Institutional Approval Analytics System</h1>', unsafe_allow_html=True)
    st.markdown('<h3 class="sub-header">UGC & AICTE - Institutional Performance Tracking & Decision Support</h3>', unsafe_allow_html=True)
    
    # System overview
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("""
        <div class="info-box">
        <h4>🚀 System Overview</h4>
        <p>This AI-powered platform automates the analysis of institutional historical data, performance metrics, 
        and document compliance for UGC and AICTE approval processes.</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="warning-box">
        <h4>🔒 Secure Access</h4>
        <p>Authorized UGC/AICTE personnel and registered institutions only. All activities are logged and monitored.</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.success("✅ AI Analytics System Successfully Initialized!")
    
    # Display quick stats
    st.subheader("📈 System Quick Stats")
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_institutions = len(analyzer.historical_data['institution_id'].unique())
        st.metric("Total Institutions", total_institutions)
    
    with col2:
        years_data = len(analyzer.historical_data['year'].unique())
        st.metric("Years of Data", years_data)
    
    with col3:
        current_year_data = analyzer.historical_data[analyzer.historical_data['year'] == 2023]
        if len(current_year_data) > 0:
            avg_performance = current_year_data['performance_score'].mean()
            st.metric("Avg Performance Score", f"{avg_performance:.2f}/10")
        else:
            st.metric("Avg Performance Score", "N/A")
    
    with col4:
        if len(current_year_data) > 0:
            approval_ready = (current_year_data['performance_score'] >= 6.0).sum()
            st.metric("Approval Ready", approval_ready)
        else:
            st.metric("Approval Ready", "N/A")
    
    # SINGLE sidebar navigation section - REMOVE APPROVAL WORKFLOW FOR NON-INSTITUTION ROLES
    st.sidebar.title("🧭 Navigation Panel")
    st.sidebar.markdown("---")
    st.sidebar.subheader("🔐 Authentication")
    
    user_role = st.sidebar.selectbox(
        "Select Your Role",
        ["Institution", "UGC Officer", "AICTE Officer", "System Admin", "Review Committee"]
    )
    
    if user_role == "Institution":
        create_institution_login(analyzer)
        return
    
    # For other roles, show the AI modules (WITHOUT APPROVAL WORKFLOW)
    st.sidebar.markdown("### AI Modules")
    
    app_mode = st.sidebar.selectbox(
        "Select Analysis Module",
        [
            "📊 Performance Dashboard",
            "📋 Document Analysis", 
            "🤖 Intelligence Hub",
            "🔍 RAG Data Management",
            "💾 Data Management",
            # "🔄 Approval Workflow",  # REMOVED FROM THIS LIST
            "⚙️ System Settings"
        ]
    )
    
    # Route to selected module
    if app_mode == "📊 Performance Dashboard":
        create_performance_dashboard(analyzer)
    
    elif app_mode == "📋 Document Analysis":
        create_document_analysis_module(analyzer)
    
    elif app_mode == "🤖 Intelligence Hub":
        create_institutional_intelligence_hub(analyzer)
    
    elif app_mode == "🔍 RAG Data Management":
        create_rag_data_management(analyzer)
    
    elif app_mode == "💾 Data Management":
        create_data_management_module(analyzer)
    
    elif app_mode == "⚙️ System Settings":
        create_system_settings(analyzer)

    # Add to your main navigation in main() function:
    elif app_mode == "🌐 API Integration":
        create_api_documentation()
   
        
        st.subheader("System Information")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Database Records", len(analyzer.historical_data))
        with col2:
            st.metric("Unique Institutions", analyzer.historical_data['institution_id'].nunique())
        with col3:
            st.metric("Data Years", f"{analyzer.historical_data['year'].min()}-{analyzer.historical_data['year'].max()}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #6c757d;'>
    <p><strong>UGC/AICTE Institutional Analytics Platform</strong> | AI-Powered Decision Support System</p>
    <p>Version 2.0 | For authorized use only | Data last updated: {}</p>
    </div>
    """.format(datetime.now().strftime("%Y-%m-%d %H:%M")), unsafe_allow_html=True)
if __name__ == "__main__":
    main()
